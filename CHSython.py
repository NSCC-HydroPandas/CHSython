from tkinter import *
from idlelib.ToolTip import *
from tkinter import ttk
from tkinter import messagebox
from tkcalendar import Calendar, DateEntry
import pandas as pd
import geopandas as gpd
import numpy as np
from tkinter import filedialog
from os import mkdir, chdir, listdir, path, walk, startfile, getcwd, rename, startfile, remove
import win32com.client as com
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
import datetime as DATES
import subprocess as S
import matplotlib.pyplot as plt
from math import sqrt, pi
import openpyxl
import re
import time
from shapely import wkt
from shapely.geometry import Polygon
from shapely.ops import cascaded_union
import shapefile
import pygeoif
import fileinput
from xml.dom import minidom
import xml.etree.cElementTree as ET
import sys
sys.path.append('C:/Program Files/CARIS/HIPS and SIPS/11.4/python/3.5')
import caris.coverage as cov
import caris
from xml.dom import minidom
import warnings
warnings.simplefilter("ignore")
from hips_project import *


owd = getcwd()
Caris = ('C:/Program Files/CARIS/HIPS and SIPS/11.4/bin')
BASE4 = ('C:/Program Files/CARIS/BASE Editor/4.4/bin')
BASE5 = ('C:/Program Files/CARIS/BASE Editor/5.5/bin')
PosPac = ('C:/Program Files/Applanix/POSPac MMS 8.3')
QCTools = ('C:/Users/LegerMi/Downloads/QCTools.3.6.1')
Python = ('C:/Users/LegerMI/AppData/Local/Programs/Python/Python35')
##CATools = ('C:/Tools/CATools.2.3.0')

def CONV_DOY():
    """ Opens and Run Julian Day Convertor"""
    chdir(owd)
    p = S.check_call(Python + "/python.exe " + owd + "/JD.py", shell=True)
    #p = S.Popen(['JD.exe'])
    #p.communicate()


def IWLStoHIPSTIDE():
    chdir(owd)
    p = S.check_call(Python + "/python.exe " + owd + "/Convert_IWLS.py", shell=True)


def CSARtoGEOTIFF():
    chdir(owd)
    p = S.check_call(Python + "/python.exe " + owd + "/ExportGeotiffs.py", shell=True)

def RefractionEditor():
    chdir(owd)
    p = S.check_call(Python + "/python.exe " + owd + "/Refract.py", shell=True)

    
def DMS_to_DD(coords_DMS):
    """Converts DMS to DD"""

    Sep_DMS = coords_DMS.split('-')
    coords_DD = (float(Sep_DMS[2])/3600) + (float(Sep_DMS[1])/60) + float(Sep_DMS[0])
    return(coords_DD)

def DD_to_Rads(coords_DD):
    """Converts DD to Radians"""
    coords_rads = coords_DD*(pi/180)
    return(coords_rads)


def TPU(order, Depth):
    """Allowable Vertical and Horizontal Uncertainty Calculation for each
    IHO and CHS order"""

    if order == 'EXCLUSIVE':
        a = 0.15
        b = 0.0075
        THU_v = 1

    elif order =='SPECIAL':
        a = 0.25
        b = 0.0075
        THU_v = 2

    elif order == '1A' or order == '1B':
        a = 0.5
        b = 0.013
        THU_v = 5 + 0.05 * Depth

    elif order == '2' or order == '3':
        a = 1.0
        b = 0.023
        THU_v = 20 + 0.10 * Depth


    TVU_v = round(sqrt(a**2 + (b * Depth)**2),3)
    return(TVU_v, THU_v) ## Return Expected Total Vertical Uncertainty SOUACC and Postional Uncertainty POSACC

##def TransformCoordinates(WKT, Sur):
##
##    InCRS = caris.CoordinateReferenceSystem('epsg', '4326')
##    in_geom = caris.Geometry(InCRS, WKT)
##
##    Sur_CRS = caris.CoordinateReferenceSystem('WKT', Sur.wkt_cosys)
##    Outgeo = in_geom.transform(Sur_CRS)
##    CRSTransformed = Outgeo.wkt
##    return(CRSTransformed)


class Application(Frame):


    def __init__(self, master):
        """ Initialize the Frames for CHSython"""

        Frame.__init__(self, master)
        self.grid()
        self.general_hips_options()
        self.Load_Auxiliary_Par()
        self.Load_Hips_Project_Par()
        self.Load_SVP_Par()
        self.Sub_Rep()
        self.app_widgets()
        self.Load_GRID_Par()
        self.POSPAC_Par()
        #self.Copy_HIPS()


    def app_widgets(self):
        """Defines GUI Widgets"""

        ## Create Main Menu Bar
        menu.add_cascade(label = "File", menu = submenu)

        ## Create Submission Menu Bar
        menu.add_cascade(label = "Additional Tools", menu = submenu2)

        ## Create Project Dir
        submenu2.add_command(label = "Create Proj Dir", command = self.Create_Project_Dir)

        ## Convert DOY to JD
        submenu2.add_command(label = "DOY to JD", command = CONV_DOY)

        ## Convert IWLS to CARIS Tide
        submenu2.add_command(label = "IWLS to CARIS Tide", command = IWLStoHIPSTIDE)

        ## Convert CSAR to GEOTIFF
        submenu2.add_command(label = "CSAR Surface to GEOTIFF", command = CSARtoGEOTIFF)

        ## Convert CSAR to Refraction Editor
        submenu2.add_command(label = "Refraction Editor", command = RefractionEditor)

        ## Save User Parameters
        submenu.add_command(label = "Save Parameters", command = self.Save_Par)

        ## Help Submenu
        submenu.add_command(label = "Help", command = self.Help)

        ## Close Submenu
        submenu.add_command(label = "Close Application", command = self.close)

        ## Process Data
        self.Button_P = Button(self, text="Process Data", height=0,
                               command=self.CHS_Proccessing)
        self.Button_P.grid(row=2, column=2, sticky=W, padx=2)


    def CHS_Proccessing(self):
        """Processing steps based on user inputs"""

        chdir(str(self.OUT_F.get()))
        if path.exists(str(self.JULIAN_D.get())):
            pass
        else:
            mkdir(str(self.JULIAN_D.get())) ## Create a Julian Day dump folder for CHSython Output

        chdir(owd)

        if self.S_T.get()==1 or self.S_T.get()==2 or self.S_T.get()==3 or self.S_T.get()==4 or self.S_T.get()==5:
            self.IMPORT_TO_HIPS() ## Imports RAW data through Import HIPS Process

        if self.A_T.get()==1 or self.A_T.get()==2:
            self.Import_Auxiliary() ## Imports POSMV or SBET Data through Import Applanix Data Proccess

        if (self.T_T.get()==1 or self.T_T.get()==2) or self.COMP_TPU.get()==1 or self.APPLY_SVP.get() or self.MERGE_TRACK.get()==1:
            self.GEOREFERENCE_HIPS() ## Runs Georeferencing steps through Georeferenceing Process

        if self.GRID.get()==1 or self.GRID.get()==2:
            self.Create_Addto_Hips_Grid() ## Creates or adds Hips data to surfaces using HIPS Gridding

        self.Combine_Caris_Output() ## Combines all Output logs into 1 File and saves to Julian Day dump folder

        if self.D_R.get() == 1:
            self.Run_Daily_Report() ## Run Reporting Script for Daily and Weekly Reports


    def Search_RAW_Data(self):
        """Allows the user to choose the
        dir where RAW sonar files live, then updates the
        Raw file entry box with the selected dir path"""

        raw = self.RAW_F.get()

        RAW_Filedir = filedialog.askdirectory(title='Select Raw Sensor File ' +
                                              'Directory', initialdir=raw)
        self.RAW_F.set(RAW_Filedir)
        tip_RAW = ToolTip(self.RAW_f, (self.RAW_F.get()))


    def Search_HDCS_Data(self):
        """Allows the user to choose the
        dir with the HDCS Data (Proccessing Folder), then updates the
        Proccessing entry box with the selected dir path"""

        hdcs = self.HDCS_D.get()

        HDCS_Filedir = filedialog.askdirectory(title='Select Proccessing folder' +
                                               'Directory', initialdir=hdcs)
        self.HDCS_D.set(HDCS_Filedir)
        tip_HDCS = ToolTip(self.HDCS_d, (self.HDCS_D.get()))


    def Search_VesselFile(self):
        """Allows the user to choose the
        Vessel Config file, then updates the
        Vessel file entry box with the selected Vessel file path"""

        vessel = self.VESSEL_N.get()
        V = path.split(vessel)

        VESSEL_File = filedialog.askopenfilename(initialdir = V[0],
                                       title = 'Select Vessel File',
                                       filetypes = (("Vessel Config","*.hvf"),("all files","*.*")))
        self.VESSEL_N.set(VESSEL_File)
        tip_Vessel = ToolTip(self.VESSEL_n, (self.VESSEL_N.get()))


    def Search_Aux_Data(self):
        """Allows the user to choose the
        dir with the POS files"""


        Aux = self.AUX_F.get()

        AUX_Filedir = filedialog.askdirectory(title='Select Folder ' +
                                                'Directory', initialdir=Aux)
        self.AUX_F.set(AUX_Filedir)
        tip_AUX_F = ToolTip(self.AUX_f, (self.AUX_F.get()))
        self.POSDIR.set(AUX_Filedir)
        tip_POSDir = ToolTip(self.POSDir, (self.POSDIR.get()))



    def Search_Aux_Data2(self):
        """Allows the user to choose the
        dir with the RMS File, then updates the
        RMS file entry box with the selected dir path"""

        fdir2 = self.AUX_F2.get()
        fd2 = path.split(fdir2)

        AUX_Filedir = filedialog.askopenfilename(initialdir = fd2[0],
                                       title = 'Select RMS File',
                                       filetypes = (("RMS","*.out"),("all files","*.*")))
        self.AUX_F2.set(AUX_Filedir)
        tip_AUX_F2 = ToolTip(self.AUX_f2, (self.AUX_F2.get()))


    def Search_Aux_Data3(self):
        """Allows the user to choose the
        dir with the SBET File, then updates the
        SBET file entry box with the selected dir path"""

        fdir3 = self.AUX_F3.get()
        fd3 = path.split(fdir3)


        AUX_Filedir = filedialog.askopenfilename(initialdir = fd3[0],
                                    title = 'Select SBET File',
                                    filetypes = (("SBET","*.out"),("all files","*.*")))
        self.AUX_F3.set(AUX_Filedir)
        tip_AUX_F3 = ToolTip(self.AUX_f3, (self.AUX_F3.get()))


    def Search_GNSS_Obs(self):
        """Allows the user to choose the
        GNSS observable file for the GPS Station to be used in
        User Defined Single Base POSPAC Processing"""

        gnssobs = self.GNSSFile.get()
        gnssobsf = path.split(gnssobs)


        Gnss_File = filedialog.askopenfilename(initialdir = gnssobsf[0],
                                       title = 'Select Observation File',
                                       filetypes = (("Obs","*.*o"),("all files","*.*")))

        self.GNSSFile.set(Gnss_File)
        tip_GNSSFile = ToolTip(self.GnssFile, (self.GNSSFile.get()))


    def Search_OUTPUT(self):
        """Allows the user to choose an Output
        dir for the Caris log information to be saved as text files,
        then updates the Output entry box with the selected dir path"""

        out = self.OUT_F.get()

        OUTPUT_Filedir = filedialog.askdirectory(title='Select Output Folder ', initialdir=out)
        tip_Out = ToolTip(self.OUT_f, (self.OUT_F.get()))
        self.OUT_F.set(str(OUTPUT_Filedir))


    def Search_SVP(self):
        """Allows the user to choose a SVP dir
        for runing Caris SVP in Georeferencing"""

        svp = self.SVPDir.get()

        #svp = path.split(sv)

        SVP_Filedir = filedialog.askdirectory(initialdir = svp ,
                                              title='Select SVP File ' +
                                              'Directory')
        self.SVPDir.set(SVP_Filedir)
        #tip_Out = ToolTip(self.SVPdir,(self.SVPDir.get()))



        ##Steps for Combining SVP
        ##Project = self.PROJECT_n.get()
        ##chdir(SVP_Filedir)
        ##with open('Combined_SVP.bat', 'w') as Svp_C:
        ##     Svp_C.write('cd ' + SVP_Filedir + '\n \b')
        ##     Svp_C.write('copy *.svp ' + str(Project)+ '.svp')
        ##self.SVP_F.set(str(SVP_Filedir) + '/' + str(Project) + '.svp') ## No longer required as Caris now concatinates files


    def Search_Sub_dir_file(self):
        """Allows the user to choose the
        ATL ISO Submission File"""


        Sub_dir = filedialog.askdirectory(initialdir = "/", title='Select Submission Directory ')
        self.SUB_D.set(str(Sub_dir))


    def Search_dir(self):
        """Allows the user to choose a Project
        dir for creating Project folder structure"""

        self.PF = filedialog.askdirectory(initialdir = "/", title='Select Project directory ')


    def Search_TIDE_File(self):
        """Allows the user to choose the
        Tide file, then updates the
        Tide file entry box with the selected Tide file path"""

        tf = self.T_f.get()
        tfp = path.split(tf)

        Tide_File = filedialog.askopenfilename(initialdir = tfp[0],
                                       title = 'Select Tide File',
                                       filetypes = (("Tide Files","*.tid"),("all files","*.*")))
        self.T_F.set(Tide_File)
        tip_MODEL = ToolTip(self.T_f, (self.T_F.get()))


    def Search_Model_File(self):
        """Allows the user to choose the
        Tide Model file, then updates the
        Tide file entry box with the selected Tide Model file path"""

        mf = self.M_F.get()
        mfp = path.split(mf)

        Model_File = filedialog.askopenfilename(initialdir = mfp[0],
                                       title = 'Select Model File',
                                       filetypes = (("all files","*.*"), ("Text Model File","*.txt"),("CSV","*.csv"), ("XYZ","*.xyz"),
                                                    ("Raster Model File","*.csar")))
        self.M_F.set(Model_File)
        tip_MODEL = ToolTip(self.M_f, (self.M_F.get()))


    def Search_Info_File(self):
        """Allows the user to choose the
        Tide Model info file, then updates the
        Tide file entry box with the selected Tide Model info file path"""

        info = self.INFO_F.get()
        infop = path.split(info)

        INFO_File = filedialog.askopenfilename(initialdir = infop[0],
                                       title = 'Select Model Info File',
                                       filetypes = (("Tide Info Files","*.info"),("all files","*.*")))
        self.INFO_F.set(INFO_File)
        tip_INFO = ToolTip(self.Info_f, (self.INFO_F.get()))


    def Search_Grid_Dir(self):
        """Allows the user to choose the
        Surfaces dir, then updates the
        Surfaces dir entry box with the selected path"""

        grid = self.GRID_DIR.get()

        Grid_dir = filedialog.askdirectory(initialdir = grid, title='Select Surface directory ')
        self.GRID_DIR.set(str(Grid_dir))
        tip_GRID = ToolTip(self.GRID_dir, (self.GRID_DIR.get()))


    def Search_CSAR_File(self):
        """Allows the user to choose the
        Csar file for input into QCTools"""

        CSAR_File = filedialog.askopenfilename(initialdir = "/",
                                       title = 'Select CSAR File',
                                       filetypes = (("Csar File","*.csar"),
                                                    ("all files","*.*")))
        CSAR_File = CSAR_File.replace("/", "\\")
        self.CSAR_F.set(CSAR_File)
        tip_CSAR_F = ToolTip(self.CSAR_f, (self.CSAR_F.get()))


    def Search_GEOTIFF_File(self):  ## Not running Code
        """Allows the user to choose the
        Geotiff or Geotif file for input into CATools"""

        GEOTIFF_File = filedialog.askopenfilename(initialdir = "/",
                                       title = 'Select Geotiff File',
                                       filetypes = (("Geotiff File","*.tiff"),
                                                    ("Geotif File","*.tif"),
                                                    ("all files","*.*")))
        GEOTIFF_File = GEOTIFF_File.replace("/", "\\")
        self.GEOTIFF_F.set(GEOTIFF_File)
        tip_GEOTIFF_F = ToolTip(self.GEOTIFF_f, (self.GEOTIFF_F.get())) ## May not need once command line is availible


    def Search_LINE_File(self):
        """Allows the user to choose the
        Caris Line Report - Created using HIPS"""

        line = self.LINE_F.get()
        linep = path.split(line)

        LINE_File = filedialog.askopenfilename(initialdir = linep[0],
                                       title = 'Select Caris Line Report',
                                       filetypes = (("ASCII Text","*.txt"),
                                                    ("all files","*.*")))
        self.LINE_F.set(LINE_File)
        tip_LINE_F = ToolTip(self.LINE_f, (self.LINE_F.get()))


    def Search_SpreadSheet_File(self):
        """Allows the user to choose the
        Daily Report Spreadsheet"""

        rf = self.REP_F.get()
        rfp = path.split(rf)

        REP_File = filedialog.askopenfilename(initialdir = "/",
                                       title = 'Select Line Report Spreadsheet',
                                       filetypes = (("Spread Sheet","*.xlsx"),
                                                    ("all files","*.*")))
        self.REP_F.set(REP_File)
        tip_REP_F = ToolTip(self.REP_f, (self.REP_F.get()))


    def Search_SpreadSheet_File2(self):
        """Allows the user to choose the
        Weekly Report Spreadsheet"""

        Wrf = self.WREP_F.get()
        Wrfp = path.split(Wrf)

        WREP_File = filedialog.askopenfilename(initialdir = "/",
                                       title = 'Select Line Report Spreadsheet',
                                       filetypes = (("Spread Sheet","*.xlsx"),
                                                    ("all files","*.*")))
        self.WREP_F.set(WREP_File)
        tip_WREP_F = ToolTip(self.WREP_f, (self.WREP_F.get()))


    def Search_VALSRC_Folder(self):
        """Allows the user to choose the
        dir with the VALSRC surfaces, then updates the
        VALSRC entry box with the selected dir path"""

        Valsrc = self.VALSRC_F.get()
        VALSRC_Filedir = filedialog.askdirectory(title='Select VALSRC folder' +
                                               'Directory', initialdir=Valsrc)
        self.VALSRC_F.set(VALSRC_Filedir)
        tip_HDCS = ToolTip(self.VALSRC_f, (self.VALSRC_F.get()))

    def Search_DTMFolder(self):
        Valsrc = self.DTM_DIR.get()
        VALSRC_Filedir = filedialog.askdirectory(title='Select VALSRC folder' +
                                               'Directory', initialdir=Valsrc)
        self.DTM_DIR.set(VALSRC_Filedir)
        tip_HDCS = ToolTip(self.DTM_dir, (self.DTM_DIR.get()))


    def Search_QC_OUT(self): ## Not running Code
        """Allows the user to set the
        Output dir for QCTools and CATools"""

        OUTPUT_QC = filedialog.askdirectory(title='Select QC Tools Output Folder')
        OUTPUT_QC = OUTPUT_QC.replace("/", "\\")
        self.QC_OUT.set(OUTPUT_QC)
        tip_Out = ToolTip(self.OUT_f, (self.OUT_F.get()))


    def Search_ENC_Dir(self):  ## Not running Code
        """Allows the user to choose the
        ENC Directory for input into CATools"""

        ENC_DIR = filedialog.askdirectory(title='Select QC Tools Output Folder')
        ENC_DIR = ENC_DIR.replace("/", "\\")
        self.ENC_DIR.set(ENC_DIR)
        tip_Out = ToolTip(self.ENC_Dir, (self.ENC_DIR.get()))


    def Search_TrackLines (self):
        """Search for Trackline Folder"""
        Dir = filedialog.askdirectory(initialdir = "/", title='Select Track Line Directory')
        self.TL_Dir.set(Dir)

        TL_Dir_list = listdir(Dir)

        self.listbox.delete(0,'end')
        for item in TL_Dir_list:
            if not item.startswith('JD') and not item.endswith(".rawdataindex"): 
                self.listbox.insert(END, item)


    def Save_Par(self):
        """"Allows user to save all CHSython Parameters to a CSV file to be
        loaded on application startup"""

        chdir(owd) ## Application Dir

        ## General Parameters
        RAW = self.RAW_f.get() ## RAW File Dir
        HDCS = self.HDCS_d.get() ## HDCS_Data Dir
        PROJECT = self.PROJECT_n.get() ## Project Name
        VESSEL = self.VESSEL_n.get() ## Vessel Config File
        CRS = self.CRS_op.get() ## Project CRS
        JULIAN = self.JULIAN_d.get() ## Julian Day of survey RAW Files
        YEAR = self.Year.get() ## Year of Survey
        OUT = self.OUT_F.get() ## Output folder for Caris Batch cmd Output
        if self.CONVERT_N.get()==1: ## Convert Navigation check box from Import to HIPS Proccess
            CONNAV = '1'
        else:
            CONNAV='0'

        ## Save General Parameters to Parameters.txt
        GenPar_List = [RAW, HDCS, PROJECT, VESSEL, CRS, JULIAN, YEAR, CONNAV, OUT,'N/A','N/A','N/A']
        Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
        Parameters.iloc[0] = GenPar_List
        Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        ## Auxillairy Parameters
        if self.ALLOW_P.get()==1: ## Allow Partially Covered
            ApC = '1'
        else:
            ApC = '0'
        MaG = self.MAG.get() ## Maximum Allowable Gap
        CrsAux = self.CRS_POS.get()

        if self.A_T.get()==2:
            AUX_F3 = self.AUX_F3.get()
            AUX_F2 = self.AUX_F2.get()
            App_List1 = [ApC, MaG,'N/A',AUX_F2,AUX_F3,CrsAux,'N/A','N/A','N/A','N/A','N/A','N/A']
        else:
            AUX_F = self.AUX_F.get() ## Aux File Directory
            App_List1 = [ApC, MaG, AUX_F, 'N/A','N/A',CrsAux,'N/A','N/A','N/A','N/A','N/A','N/A']
        
        ## Save General POS/SBET/RMS Parameters to Parameters.txt
        Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
        Parameters.iloc[1] = App_List1
        Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        if self.A_T.get()==1 or self.A_T.get()==2:
            if self.NAV.get()==1: ## Navigation
                N = '1'
            else:
                N ='0'
            G = self.GYRO.get() ## Gyro
            P = self.PITCH.get() ## Pitch
            R = self.ROLL.get() ## Roll
            GPSH = self.GPS_H.get() ## GPSH
            DH = self.D_H.get() ## Delayed Heave
            N_RMS = self.NAV_RMS.get() ## Naviagtion RMS
            G_RMS = self.GYRO_RMS.get() ## Gyro RMS
            P_RMS = self.PITCH_RMS.get() ## Pitch RMS
            R_RMS = self.ROLL_RMS.get() ## Roll RMS
            GPS_RMS = self.GPSH_RMS.get() ## GPSH RMS
            D_RMS = self.DH_RMS.get() ## Delayed Heave RMS

            ## Save POS/SBET/RMS Parameters to Parameters.txt
            App_List2 = [N, G, P, R, GPSH, DH, N_RMS, G_RMS, P_RMS, R_RMS, GPS_RMS, D_RMS]
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[2] = App_List2
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

            ## Save type of postions data to apply in Import to Applanix Proccess
            if self.POS_GYRO.get()==1:
                Gc = 1
            else:
                Gc = 0

            if  self.POS_PITCH.get==1:
                Pc = 1
            else:
                Pc = 0

            if self.POS_ROLL.get()==1:
                Rc = 1
            else:
                Rc = 0

            if self.POS_GPSH.get()==1:
                GPSHc = 1
            else:
                GPSHc = 0

            if self.POS_DH.get()==1:
                DHc = 1
            else:
                DHc = 0

            if self.POS_NRMS.get()==1:
                NRMSc = 1
            else:
                NRMSc = 0

            if self.POS_GRMS.get()==1:
                GRMS = 1
            else:
                GRMSc= 0

            if self.POS_PRMS.get()==1:
                PRMSc = 1
            else:
                PRMSc = 0

            if self.POS_RRMS.get()==1:
                RRMSc = 1
            else:
                RRMSc = 0

            if self.POS_GPSHRMS.get()==1:
               GPSHRMSc = 1
            else:
                GPSHRMSc = 0

            if self.POS_DHRMS.get()==1:
                DHRMSc = 1
            else:
                DHRMSc = 0

            App_List3 = [Gc, Pc, Rc, GPSHc, DHc, NRMSc, GRMSc, PRMSc, RRMSc, GPSHRMSc, DHRMSc, 'N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[12] = App_List3
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)


        ## RAW Sonar Parameters
        ## Kongsberg .all
        if self.S_T.get()==1:
            NAV_D = self.Nav_D.get() ## Navigation Device
            GPSH_D = self.GPSH_D.get() ## GPSH Device
            Heave_D = self.Heave_D.get() ## Heave Device
            Heading_D = self.Heading_D.get() ## Heading Device
            GPS_T = self.GPS_T.get() ## GPS Time
            Pitch_D = self.Pitch_D.get() ## Pitch Device
            Roll_D = self.Roll_D.get() ## Roll Device
            SSP_D = self.SSP_D.get() ## Surface Sound Speed Device

            ## Save Kongsberg .all Parameters to Parameters.txt
            Raw_List = [NAV_D, GPSH_D, Heave_D, Heading_D, GPS_T, Pitch_D, Roll_D, SSP_D,'N/A','N/A','N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[3] = Raw_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        elif self.S_T.get()==2:
            ## R2Sonic .gsf
            D_S = self.D_S.get() ## Depth Source
            IN_OFF = self.IN_OFF.get() ## Include Offline
            REJ_OFF = self.REJ_OFF.get() ## Reject Offline

            ## Save R2Sonic .gsf Parameters to Parameters.txt
            Raw_List = [D_S, IN_OFF, REJ_OFF,'N/A','N/A','N/A','N/A','N/A','N/A','N/A','N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[4] = Raw_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        elif self.S_T.get()==3:
            ## Triton .xtf
            NAV_DX = self.Nav_DX.get() ## Navigation Device
            GPSH_DX = self.GPSH_DX.get() ## GPS Height Device
            M_D = self.M_D.get() ## Motion Device
            C_B = self.C_B.get() ## Convert Bathymetry
            Heading_DX = self.Heading_DX.get() ## Heading Device
            CONV_SS = self.CONV_SS.get() ## Convert Side Scan
            SSWF = self.SSWF.get() ## Side Scan Weighting Factor
            SS_NAV = saelf.SS_NAV.get() ## Side Scan Navigation Device
            SS_HEAD = self.SS_HEAD.get() ## Side Scan Heading Device
            TIME_S = self.TIME_S.get() ## Time Stamps

            ## Save Triton .xtf Parameters to Parameters.txt
            Raw_List = [NAV_DX, GPSH_DX, M_D, C_B, Heading_DX, CONV_SS, SSWF, SS_NAV, SS_HEAD, TIME_S,'N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[5] = Raw_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        elif self.S_T.get()==4:
            ## Teledyne .s7k
            CB = self.CB.get() ## Convert Bathymetry
            NAV_D = self.NAV_D.get() ## Navigation Device
            HEAD_D = self.HEAD_D.get() ## Heading Device
            MOTION_D = self.MOTION_D.get() ## Motion Device
            SWATH_D = self.SWATH_D.get() ## Swath Device

            ## Save Teledyne .s7k Parameters to Parameters.txt
            Raw_List = [CB, NAV_D, HEAD_D, MOTION_D, SWATH_D,'N/A','N/A','N/A','N/A','N/A','N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[15] = Raw_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        elif self.S_T.get()==5:
            NAV_D = self.Nav_D.get() ## Navigation Device
            GPSH_D = self.GPSH_D.get() ## GPSH Device
            Heave_D = self.Heave_D.get() ## Heave Device
            Heading_D = self.Heading_D.get() ## Heading Device
            Pitch_D = self.Pitch_D.get() ## Pitch Device
            Roll_D = self.Roll_D.get() ## Roll Device
            DelHeave_D = self.DelHeave_D.get() ## Delayed Heave Device
            GPS_T = self.GPS_T.get()

            ## Save Kongsberg .all Parameters to Parameters.txt
            Raw_List = [NAV_D, GPSH_D, Heave_D, Heading_D, Pitch_D, Roll_D, DelHeave_D, GPS_T, 'N/A','N/A','N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[16] = Raw_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        ## Tide Parameters
        if self.T_T.get()==1:
            ## Compute GPS Tide
            C_GPS_ADJ = self.C_GPS_ADJ.get() ## Compute GPS Tide Adjustment
            SD_OFF = self.SD_OFF.get() ##
            M_F = self.M_F.get() ## Tide Model File
            W_L = self.W_L.get() ## Water Level

            ## Check if Model File is Raster (.Csar) or Text File (.txt)
            if (M_F.endswith('.txt') or M_F.endswith('.csv')
                or M_F.endswith('.xyz')):
                INFO_F = self.INFO_F.get() ## Tide Model Info File
                INFO_CRS = self.INFO_CRS.get() ## Tide Model CRS
            else:
                INFO_F = ('N/A')
                INFO_CRS = ('N/A')

            ## Save Compute GPS Tide Parameters to Parameters.txt
            GPST_List = [C_GPS_ADJ, SD_OFF, M_F, INFO_F,  INFO_CRS, W_L,'N/A','N/A','N/A','N/A','N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[7] = GPST_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        elif self.T_T.get()==2:
            ## Observed/ Predicted Tides
            T_F = self.T_F.get() ## Tide File
            W_Ave = self.W_Ave.get() ## Wieghted Ave
            COMP_Errors = self.COMP_Errors.get() ## Compute Errors

            ## Save Tide Parameters to Parameters.txt
            Tides_List = [T_F, W_Ave, COMP_Errors,'N/A','N/A','N/A','N/A','N/A','N/A','N/A', 'N/A', 'N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[6] = Tides_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        ## TPU Parameters
        if self.COMP_TPU.get()==1:
            ## Compute TPU
            TIDE_M = self.TIDE_M.get() ## Measured Tide
            SV_M = self.SV_M.get() ## Measured Sound Velocity
            SS_V = self.SS_V.get() ## Surface Sound Velocity
            S_N = self.S_N.get() ## Navigation Source
            S_G = self.S_G.get() ## Gyro Source
            S_S = self.S_S.get() ## Sonar Source
            S_P = self.S_P.get() ## Pitch Source
            S_R = self.S_R.get() ## Roll Source
            S_H = self.S_H.get() ## Heave Source
            S_Tide = self.S_Tide.get() ## Tide Source
            ## Merge
            H_Merged = self.H_MERGED.get() ## Heave Type
            V_REF = self.VERT_REF.get() ## Vertical Reference Meta Data

            ## Save TPU Parameters to Parameters.txt
            TPU_List = [TIDE_M, SV_M, SS_V, S_N, S_G, S_S, S_P, S_R, S_H, S_Tide, 'N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[9] = TPU_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

            ## Save Merged & TPU Parameters to Parameters.txt
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[8,0] = H_Merged
            Parameters.iloc[8,1] = V_REF
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        ## Surface Gridding Parameters
        if self.GRID.get()==1 or self.GRID.get()==2:
            RES = self.RES.get()
            GRID_DIR = self.GRID_DIR.get()
            IHO_O = self.IHO_ORDER.get()

            ## Save Gridding Parameters to Parameters.txt
            GRID_List = [RES, IHO_O, GRID_DIR, 'N/A','N/A','N/A','N/A','N/A','N/A','N/A', 'N/A', 'N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[10] = GRID_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        ## SVP Parameters
        if self.APPLY_SVP.get():

            SVP_F = self.SVP_F.get() ## Directory of SVP
            PROFILE = self.PROFILE.get() ## Profile Selection Method
            ND_HOUR = self.ND_HOUR.get() ## Hour for Profile Time
            H_Merged2 = self.H_MERGED2.get() ## Heave Type
            V_REF2 = self.VERT_REF2.get() ## Vertical Reference Type

            ## Save SVP Parameters to Parameters.txt
            SVP_List = [SVP_F, PROFILE, ND_HOUR, 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A' , 'N/A' ]
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[11] = SVP_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

            ## Save Merged & SVP Parameters to Parameters.txt
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[8,0] = H_Merged2
            Parameters.iloc[8,1] = V_REF2
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        ## POSPAC Processing Parameters to Parameters.txt
        POSDir = self.POSDIR.get()
        StationFile = self.GNSSFile.get() ## Base Station File
        StationID = self.STID.get() ## Base Station ID/Name
        REF_FRAME = self.REF_FRAME.get() ## POSPAC Reference Frame for Proccessing

        POSPAC_List = [POSDir, StationFile, StationID, REF_FRAME,'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A', 'N/A']
        Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
        Parameters.iloc[13] = POSPAC_List
        Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        if self.D_R.get() == 1:
            ## Reporting Parameters to Parameters.txt
            Daily = self.REP_F.get() ## Daily Report Spreadsheet
            Weekly = self.WREP_F.get() ## Weekly Report Spreadsheet
            Week = self.WeekNO.get() ## Week number/name
            IHOOrder = self.IHO_ORDER2.get() ## IHO Order
            QC = self.TPUQC.get() ## QC Type (Surface or HIPS)
            
            ## Reporting Parameters to Parameters.txt
            Reporting_List = ['N/A',Daily,Weekly,IHOOrder,Week,QC,'N/A','N/A','N/A','N/A','N/A','N/A']
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Parameters.iloc[14] = Reporting_List
            Parameters.to_csv('Parameters.txt', mode='w', index=False, header=False)

        self.Exit = 'False'
        self.popup_SavePar()  ## Provide output window to User


    def popup_SavePar(self):
        """Creates a window alerting the user that all parameters
        have been saved"""

        msg = "Parameters Are Saved"
        popup= Tk()


        def leavemini():
            if self.Exit == 'True':
                popup.destroy()
                root.destroy()
            else:
                 popup.destroy()

        popup.wm_title("Save Parameteres")
        label = Label(popup, text=msg)
        label.grid(row=1, column=1)
        B1 = Button(popup, text="Okay", command = leavemini)
        B1.grid(row=2, column=1)
        popup.mainloop()


    def Load_SVP_Par(self):
        """Loads deafult user input options for SVP Proccessing"""

        chdir(owd) ## Application Dir

        ## Creating SVP Frame and User Input Options
        if self.APPLY_SVP.get()==1:

            ## Creating SVP Options
            self.SVP_op = LabelFrame(frame6, text="Apply SVP", foreground="blue")
            self.SVP_op.grid(row=1, column=0, padx=1, sticky=W)

            ## Sound Speed Profile Directory
            self.SVP_F = StringVar()
            self.SVP_f = Entry(self.SVP_op, width=38, textvariable=self.SVP_F)
            self.SVP_text = Label(self.SVP_op, text="SVP Dir")
            self.SVP_text.grid(row=1, column=0, sticky=W)
            self.SVP_f.grid(row=1, column=1, sticky=W)

            ## Sound Speed Profile Selection Method
            self.PROFILE = StringVar()
            self.Profile = Entry(self.SVP_op, width=32, textvariable=self.PROFILE, state='disabled')
            self.Profile_text = Label(self.SVP_op, text='Profile Selection\nMethod')
            self.Profile_text.grid(row=3, column=0, sticky=W)
            self.Profile.grid(row=3, column=1, sticky=W)

            ## Sound Speed Profile Time
            self.ND_HOUR = StringVar()
            self.ND_Hour = Entry(self.SVP_op, width=3, textvariable=self.ND_HOUR)
            self.ND_Hour_text = Label(self.SVP_op, text='Nearest Distance\nHours')
            self.ND_Hour_text.grid(row=4, column=0, sticky=W)
            self.ND_Hour.grid(row=4, column=1, sticky=W)

            ## Heave Type for Merge
            self.H_MERGED2 = StringVar()
            self.H_Merged2 = Entry(self.SVP_op, width=15, textvariable=self.H_MERGED2, state='disabled')
            self.H_Merged_text2 = Label(self.SVP_op, text="Heave Type")
            self.H_Merged_text2.grid(row=5, column=0, sticky=W)
            self.H_Merged2.grid(row=5, column=1, sticky=W, padx=1)

            ## Vertical Reference Type
            self.VERT_REF2 = StringVar()
            vert_ref2 = ['NONE',
                        'GPS',
                        'TIDE']

            self.VREF_op2 = ttk.Combobox(self.SVP_op, values=vert_ref2, width=7, textvariable=self.VERT_REF2)
            self.VREF_text2 = Label(self.SVP_op, text="Choose Vertical Reference")
            self.VREF_text2.grid(row=2, column=0, sticky=W)
            self.VREF_op2.grid(row=2, column=1, sticky=W+E, padx=0)

            self.Button16 = Button(self.SVP_op, text="...", height=0,
                               command=self.Search_and_Combine_SVP)
            self.Button16.grid(row=1, column=2, sticky=W, padx=2)

            ## Reading defaults or user saved inputs for SVP
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            SVP_F = Parameters.iloc[11,0]
            PROFILE = Parameters.iloc[11,1]
            ND_HOUR = int(Parameters.iloc[11,2])
            H_MERGE2 = Parameters.iloc[8,0]
            VERT_R2 = Parameters.iloc[8,1]

            ## Setting defaults from Parameter file for POSMV
            self.SVP_F.set(SVP_F)
            self.PROFILE.set(PROFILE)
            self.ND_HOUR.set(ND_HOUR)
            self.H_MERGED2.set(H_MERGE2)
            self.VERT_REF2.set(VERT_R2)

##            ## Sound Speed Profile Editor
##            self.EditSVP_op = LabelFrame(frame6, text="Edit SVP", foreground="blue")
##            self.EditSVP_op.grid(row=2, column=0, sticky=W)
##
##            self.COORD_F = StringVar()
##            self.COORD_f = Entry(self.EditSVP_op, width=38, textvariable=self.COORD_F)
##            self.COORD_text = Label(self.EditSVP_op, text="SVP Coordinate File")
##            self.COORD_text.grid(row=1, column=0, sticky=W)
##            self.COORD_f.grid(row=1, column=1, sticky=W)
##            self.Button_coord = Button(self.EditSVP_op, text="...", height=0,
##                              command=self.Search_Aux_Data)
##            self.Button_coord.grid(row=1, column=2, sticky=W, padx=2)
##
##
##            ## ToolTips for SVP
##            tip_SVP = ToolTip(self.SVP_f, (self.SVP_F.get()))

        if self.APPLY_SVP.get()==0:
            try:
                ## Forget the SVP Options
                ##self.EditSVP_op.grid_forget()
                self.SVP_op.grid_forget()
            except AttributeError:
                pass


    def Load_Auxiliary_Par(self):
        """Loads deafult user input options for Importing Auxilliary Data"""

        chdir(owd) ##Application Dir

        STATE = 'normal'

        ##Creating General Import Auxilliary Data User Input Options
        Aux_op = LabelFrame(frame3, text="Applanix Import", foreground="blue")
        Aux_op.grid(row=0, column=0, padx=1, sticky=N)

        ## Applanix Dir
        self.AUX_F = StringVar()
        self.AUX_f = Entry(Aux_op, width=38, textvariable=self.AUX_F)

        ## Allow Partial Covered
        self.ALLOW_P = IntVar()
        self.ALLOW_n = Checkbutton(Aux_op, variable=self.ALLOW_P,
                                   text= "Allow Partially Covered", state='disabled')
        self.ALLOW_n.grid(row=0, column=0, sticky=W)

        ## Max Allowable GAP (Default 2sec)
        self.MAG = StringVar()
        self.Mag = Entry(Aux_op, width=7, textvariable=self.MAG, state=STATE)
        self.MAG_text = Label(Aux_op, text="Maximum Allowable Gap")
        self.MAG_text.grid(row=1, column=0, sticky=W)
        self.Mag.grid(row=1, column=1, sticky=W, padx=2)

        self.Refweek = DateEntry(Aux_op, width=25, background= "magenta3", foreground= "white", bd=2)
        self.Refweek_text = Label(Aux_op, text="GPS Refference Week")
        self.Refweek_text.grid(row=2, column=0, sticky=W)
        self.Refweek.grid(row=2, column=1, sticky=W, padx=2)

        ##Project and Hips Data CSRS
        self.CRS_POS = StringVar()
        crs_POS = ['ITRF2014: EPSG:7912@2010',
                  'NAD83(CSRS)v6: EPSG:8252@2010']

        self.CRS_pos = ttk.Combobox(Aux_op, values=crs_POS, width=35, textvariable=self.CRS_POS)
        self.CRS_pos_text = Label(Aux_op, text="Choose CRS")
        self.CRS_pos_text.grid(row=6, column=0, sticky=W)
        self.CRS_pos.grid(row=6, column=1, sticky=W+E, padx=0)

        ## Setting default Auxilliary Parameters
        Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
        APC = int(Parameters.iloc[1,0])
        M_AG = Parameters.iloc[1,1]
        crs_pos = Parameters.iloc[1,5]
        self.ALLOW_P.set(APC)
        self.MAG.set(M_AG)
        self.CRS_POS.set(crs_pos)
        

        self.Aux_msg = LabelFrame(frame3, text="Applanix User Warning", foreground="blue")
        self.Aux_msg.grid(row=0, column=1, padx=1, sticky=N+W)

        msg = ('For regular processing\nleave default values 0sec\n' +
               'and an allowable gap of\n2sec only change when\n' +
               'errors or warnings arise.\n' +
               'For SBET ensure Navigation\n' +
               'is selected.\n') ## User Reminder for Proccessing

        self.User_Msg = Text(self.Aux_msg, width=28, height=7)
        self.User_Msg.insert(END, msg)
        self.User_Msg.config(state='disabled')
        self.User_Msg.grid(row=0, column=0, padx=1, sticky=W)

        msg2 = ('Please check POS or \n' +
                'SBET CRS \n' +
                'MarineStar - \n' +
                'ITRF 2014 Epoch 2010 \n' +
                'Cannet - NAD83(CSRS)v6 \n' +
                'Epoch 2010')

        self.User_Msg2 = Text(self.Aux_msg, width=28, height=6)
        self.User_Msg2.insert(END, msg2)
        self.User_Msg2.config(state='disabled')
        self.User_Msg2.grid(row=1, column=0, padx=1, sticky=W)


        ## User Inputs for POSMV Data
        if self.A_T.get()==1:

            self.AUX_F = StringVar()
            self.AUX_f = Entry(Aux_op, width=38, textvariable=self.AUX_F)
            self.AUX_text = Label(Aux_op, text="POS Files")
            self.AUX_text.grid(row=3, column=0, sticky=W)
            self.AUX_f.grid(row=3, column=1, sticky=W)
            self.Button_AUX = Button(Aux_op, text="...", height=0,
                                command=self.Search_Aux_Data)
            self.Button_AUX.grid(row=3, column=2, sticky=W, padx=2)


            ## Creating Import POSMV User Input Options
            self.POSMV_op = LabelFrame(frame3, text="Import POSMV", foreground="blue")
            self.POSMV_op.grid(row=2, column=0, padx=1, sticky=W)

            ## Navigation
            self.NAV = IntVar()
            self.Nav = Checkbutton(self.POSMV_op, variable=self.NAV,
                                   text= "Navigation")
            self.Nav.grid(row=1, column=0, sticky=W)

            ## Gyro
            self.GYRO = StringVar()
            self.Gyro = Entry(self.POSMV_op, width=7, textvariable=self.GYRO, state=STATE)
            self.Gyro.grid(row=2, column=1, sticky=W, padx=2)
            self.POS_GYRO = IntVar()
            self.POS_gyro = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_GYRO,
                                        command=None, text= "Gyro")
            self.POS_gyro.grid(row=2, column=0, sticky=W)

            ## Pitch
            self.PITCH = StringVar()
            self.Pitch = Entry(self.POSMV_op, width=7, textvariable=self.PITCH, state=STATE)
            self.Pitch.grid(row=3, column=1, sticky=W, padx=2)
            self.POS_PITCH = IntVar()
            self.POS_pitch = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_PITCH,
                                        command=None, text= "Pitch")
            self.POS_pitch.grid(row=3, column=0, sticky=W)

            ## Roll
            self.ROLL = StringVar()
            self.Roll = Entry(self.POSMV_op, width=7, textvariable=self.ROLL, state=STATE)
            self.Roll.grid(row=4, column=1, sticky=W, padx=2)
            self.POS_ROLL = IntVar()
            self.POS_roll = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_ROLL,
                                        command=None, text='Roll')
            self.POS_roll.grid(row=4, column=0, sticky=W)

            ## GPS Height
            self.GPS_H = StringVar()
            self.GPS_h = Entry(self.POSMV_op, width=7, textvariable=self.GPS_H, state=STATE)
            self.GPS_h.grid(row=5, column=1, sticky=W, padx=2)
            self.POS_GPSH = IntVar()
            self.POS_gpsh = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_GPSH,
                                        command=None, text='GPS Height')
            self.POS_gpsh.grid(row=5, column=0, sticky=W)

            ## Delayed Heave
            self.D_H = StringVar()
            self.D_h = Entry(self.POSMV_op, width=7, textvariable=self.D_H, state=STATE)
            self.D_h.grid(row=6, column=1, sticky=W, padx=2)
            self.POS_DH = IntVar()
            self.POS_dh = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_DH,
                                        command=None, text='Delayed Heave')
            self.POS_dh.grid(row=6, column=0, sticky=W)

            ## Navigation RMS
            self.NAV_RMS = StringVar()
            self.NAV_rms = Entry(self.POSMV_op, width=7, textvariable=self.NAV_RMS, state=STATE)
            self.NAV_rms.grid(row=7, column=1, sticky=W, padx=2)
            self.POS_NRMS = IntVar()
            self.POS_nrms = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_NRMS,
                                        command=None, text='Navigation RMS')
            self.POS_nrms.grid(row=7, column=0, sticky=W)

            ## Gyro RMS
            self.GYRO_RMS = StringVar()
            self.GYRO_rms = Entry(self.POSMV_op, width=7, textvariable=self.GYRO_RMS, state=STATE)
            self.GYRO_rms.grid(row=8, column=1, sticky=W, padx=2)
            self.POS_GRMS = IntVar()
            self.POS_grms = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_GRMS,
                                        command=None, text='Gyro RMS')
            self.POS_grms.grid(row=8, column=0, sticky=W)

            ## Pitch RMS
            self.PITCH_RMS = StringVar()
            self.PITCH_rms = Entry(self.POSMV_op, width=7, textvariable=self.PITCH_RMS, state=STATE)
            self.PITCH_rms.grid(row=9, column=1, sticky=W, padx=2)
            self.POS_PRMS = IntVar()
            self.POS_prms = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_PRMS,
                                        command=None, text='Pitch RMS')
            self.POS_prms.grid(row=9, column=0, sticky=W)

            ## Roll RMS
            self.ROLL_RMS = StringVar()
            self.ROLL_rms = Entry(self.POSMV_op, width=7, textvariable=self.ROLL_RMS, state=STATE)
            self.ROLL_rms.grid(row=10, column=1, sticky=W, padx=2)
            self.POS_RRMS = IntVar()
            self.POS_rrms = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_RRMS,
                                        command=None, text='Roll RMS')
            self.POS_rrms.grid(row=10, column=0, sticky=W)

            ## GPS Height RMS
            self.GPSH_RMS = StringVar()
            self.GPSH_rms = Entry(self.POSMV_op, width=7, textvariable=self.GPSH_RMS, state=STATE)
            self.GPSH_rms.grid(row=11, column=1, sticky=W, padx=2)
            self.POS_GPSHRMS = IntVar()
            self.POS_gpshrms = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_GPSHRMS,
                                        command=None, text='GPS Height RMS')
            self.POS_gpshrms.grid(row=11, column=0, sticky=W)

            ## Delayed Heave RMS
            self.DH_RMS = StringVar()
            self.DH_rms = Entry(self.POSMV_op, width=7, textvariable=self.DH_RMS, state=STATE)
            self.DH_rms.grid(row=12, column=1, sticky=W, padx=2)
            self.POS_DHRMS = IntVar()
            self.POS_dhrms = Checkbutton(self.POSMV_op, onvalue=1, offvalue=0, variable=self.POS_DHRMS,
                                        command=None, text='Delayed Heave RMS')
            self.POS_dhrms.grid(row=12, column=0, sticky=W)

            ## Reading defaults from saved inputs for POSMV
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Na = int(Parameters.iloc[2,0])
            G = Parameters.iloc[2,1]
            P = Parameters.iloc[2,2]
            R = Parameters.iloc[2,3]
            GPSH = Parameters.iloc[2,4]
            DH = Parameters.iloc[2,5]
            N_RMS = Parameters.iloc[2,6]
            G_RMS = Parameters.iloc[2,7]
            P_RMS = Parameters.iloc[2,8]
            R_RMS = Parameters.iloc[2,9]
            GPS_RMS = Parameters.iloc[2,10]
            D_RMS = Parameters.iloc[2,11]

            Gc = Parameters.iloc[12,0]
            Pc = Parameters.iloc[12,1]
            Rc = Parameters.iloc[12,2]
            GPSHc = Parameters.iloc[12,3]
            DHc = Parameters.iloc[12,4]
            N_RMSc = Parameters.iloc[12,5]
            G_RMSc = Parameters.iloc[12,6]
            P_RMSc = Parameters.iloc[12,7]
            R_RMSc = Parameters.iloc[12,8]
            GPS_RMSc = Parameters.iloc[12,9]
            D_RMSc = Parameters.iloc[12,10]

            ## Setting defaults from Parameters file for POSMV
            self.NAV.set(Na)
            self.GYRO.set(G)
            self.PITCH.set(P)
            self.ROLL.set(R)
            self.GPS_H.set(GPSH)
            self.D_H.set(DH)
            self.NAV_RMS.set(N_RMS)
            self.GYRO_RMS.set(G_RMS)
            self.PITCH_RMS.set(P_RMS)
            self.ROLL_RMS.set(R_RMS)
            self.GPSH_RMS.set(GPS_RMS)
            self.DH_RMS.set(D_RMS)

            self.POS_GYRO.set(Gc)
            self.POS_PITCH.set(Pc)
            self.POS_ROLL.set(Rc)
            self.POS_GPSH.set(GPSHc)
            self.POS_DH.set(DHc)
            self.POS_NRMS.set(N_RMSc)
            self.POS_GRMS.set(G_RMSc)
            self.POS_PRMS.set(P_RMSc)
            self.POS_RRMS.set(R_RMSc)
            self.POS_GPSHRMS.set(GPS_RMSc)
            self.POS_DHRMS.set(D_RMSc)

            A_F = Parameters.iloc[1,2]
            self.AUX_F.set(A_F)

            ## ToolTips For Applanix Data
            tip_AUX = ToolTip(self.AUX_f, str(self.AUX_F.get()))

            try:
                ## Forget the SBET & RMS Options
                self.SBET_RMS_op.grid_forget()
                self.AUX_f2.grid_forget()
                self.Button_AUX2.grid_forget()
                self.AUX_text2.grid_forget()
            except AttributeError:
                pass

        ## User Inputs for SBET/ RMS Data
        elif self.A_T.get()==2:

            self.AUX_F2 = StringVar()
            self.AUX_f2 = Entry(Aux_op, width=38, textvariable=self.AUX_F2)
            self.AUX_text2 = Label(Aux_op, text="RMS Files")
            self.AUX_text2.grid(row=4, column=0, sticky=W)
            self.AUX_f2.grid(row=4, column=1, sticky=W)
            self.Button_AUX2 = Button(Aux_op, text="...", height=0,
                              command=self.Search_Aux_Data2)
            self.Button_AUX2.grid(row=4, column=2, sticky=W, padx=2)

            self.AUX_F3 = StringVar()
            self.AUX_f3 = Entry(Aux_op, width=38, textvariable=self.AUX_F3)
            self.AUX3_text = Label(Aux_op, text="SBET Files")
            self.AUX3_text.grid(row=3, column=0, sticky=W)
            self.AUX_f3.grid(row=3, column=1, sticky=W)
            self.Button_AUX3 = Button(Aux_op, text="...", height=0,
                                  command=self.Search_Aux_Data3)
            self.Button_AUX3.grid(row=3, column=2, sticky=W, padx=2)


            ## Creating Import SBET User Input Options
            self.SBET_RMS_op = LabelFrame(frame3, text="Import SBET and RMS", foreground="blue")
            self.SBET_RMS_op.grid(row=2, column=0, padx=1, sticky=W)

            ## Navigation
            self.NAV = IntVar()
            self.Nav = Checkbutton(self.SBET_RMS_op, variable=self.NAV,
                                   text= "Navigation")
            self.Nav.grid(row=1, column=0, sticky=W)

            ## Gyro
            self.GYRO = StringVar()
            self.Gyro = Entry(self.SBET_RMS_op, width=7, textvariable=self.GYRO, state=STATE)
            self.Gyro.grid(row=2, column=1, sticky=W, padx=2)
            self.POS_GYRO = IntVar()
            self.POS_gyro = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_GYRO,
                                        command=None, text= "Gyro")
            self.POS_gyro.grid(row=2, column=0, sticky=W)

            ## Pitch
            self.PITCH = StringVar()
            self.Pitch = Entry(self.SBET_RMS_op, width=7, textvariable=self.PITCH, state=STATE)
            self.Pitch.grid(row=3, column=1, sticky=W, padx=2)
            self.POS_PITCH = IntVar()
            self.POS_pitch = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_PITCH,
                                        command=None, text= "Pitch")
            self.POS_pitch.grid(row=3, column=0, sticky=W)

            ## Roll
            self.ROLL = StringVar()
            self.Roll = Entry(self.SBET_RMS_op, width=7, textvariable=self.ROLL, state=STATE)
            self.Roll.grid(row=4, column=1, sticky=W, padx=2)
            self.POS_ROLL = IntVar()
            self.POS_roll = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_ROLL,
                                        command=None, text='Roll')
            self.POS_roll.grid(row=4, column=0, sticky=W)

            ## GPS Height
            self.GPS_H = StringVar()
            self.GPS_h = Entry(self.SBET_RMS_op, width=7, textvariable=self.GPS_H, state=STATE)
            self.GPS_h.grid(row=5, column=1, sticky=W, padx=2)
            self.POS_GPSH = IntVar()
            self.POS_gpsh = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_GPSH,
                                        command=None, text='GPS Height')
            self.POS_gpsh.grid(row=5, column=0, sticky=W)

            ## Delayed Heave
            self.D_H = StringVar()
            self.D_h = Entry(self.SBET_RMS_op, width=7, textvariable=self.D_H, state='normal')
            self.D_h.grid(row=6, column=1, sticky=W, padx=2)

            self.POS_DH = IntVar()
            self.POS_dh = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_DH,
                                        command=None, text='Delayed Heave')
            self.POS_dh.grid(row=6, column=0, sticky=W)

            ## Navigation RMS
            self.NAV_RMS = StringVar()
            self.NAV_rms = Entry(self.SBET_RMS_op, width=7, textvariable=self.NAV_RMS, state=STATE)
            self.NAV_rms.grid(row=7, column=1, sticky=W, padx=2)
            self.POS_NRMS = IntVar()
            self.POS_nrms = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_NRMS,
                                        command=None, text='Navigation RMS')
            self.POS_nrms.grid(row=7, column=0, sticky=W)

            ## Gyro RMS
            self.GYRO_RMS = StringVar()
            self.GYRO_rms = Entry(self.SBET_RMS_op, width=7, textvariable=self.GYRO_RMS, state=STATE)
            self.GYRO_rms.grid(row=8, column=1, sticky=W, padx=2)
            self.POS_GRMS = IntVar()
            self.POS_grms = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_GRMS,
                                        command=None, text='Gyro RMS')
            self.POS_grms.grid(row=8, column=0, sticky=W)

            ## Pitch RMS
            self.PITCH_RMS = StringVar()
            self.PITCH_rms = Entry(self.SBET_RMS_op, width=7, textvariable=self.PITCH_RMS, state=STATE)
            self.PITCH_rms.grid(row=9, column=1, sticky=W, padx=2)
            self.POS_PRMS = IntVar()
            self.POS_prms = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_PRMS,
                                        command=None, text='Navigation RMS')
            self.POS_prms.grid(row=9, column=0, sticky=W)

            ## Roll RMS
            self.ROLL_RMS = StringVar()
            self.ROLL_rms = Entry(self.SBET_RMS_op, width=7, textvariable=self.ROLL_RMS, state=STATE)
            self.ROLL_rms.grid(row=10, column=1, sticky=W, padx=2)
            self.POS_RRMS = IntVar()
            self.POS_rrms = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_RRMS,
                                        command=None, text='Roll RMS')
            self.POS_rrms.grid(row=10, column=0, sticky=W)

            ## GPS Height RMS
            self.GPSH_RMS = StringVar()
            self.GPSH_rms = Entry(self.SBET_RMS_op, width=7, textvariable=self.GPSH_RMS, state=STATE)
            self.GPSH_rms.grid(row=11, column=1, sticky=W, padx=2)
            self.POS_GPSHRMS = IntVar()
            self.POS_gpshrms = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_GPSHRMS,
                                        command=None, text='GPS Height RMS')
            self.POS_gpshrms.grid(row=11, column=0, sticky=W)

            ## Delayed Heave RMS
            self.DH_RMS = StringVar()
            self.DH_rms = Entry(self.SBET_RMS_op, width=7, textvariable=self.DH_RMS, state=STATE)
            self.DH_rms.grid(row=12, column=1, sticky=W, padx=2)
            self.POS_DHRMS = IntVar()
            self.POS_dhrms = Checkbutton(self.SBET_RMS_op, onvalue=1, offvalue=0, variable=self.POS_DHRMS,
                                        command=None, text='Delayed Heave RMS')
            self.POS_dhrms.grid(row=12, column=0, sticky=W)

            ## Reading defaults inputs for SBET
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            Na = int(Parameters.iloc[2,0])
            G = Parameters.iloc[2,1]
            P = Parameters.iloc[2,2]
            R = Parameters.iloc[2,3]
            GPSH = Parameters.iloc[2,4]
            DH = Parameters.iloc[2,5]

            Gc = Parameters.iloc[12,0]
            Pc = Parameters.iloc[12,1]
            Rc = Parameters.iloc[12,2]
            GPSHc = Parameters.iloc[12,3]
            DHc = Parameters.iloc[12,4]

            ## Setting defaults from Parameter file for SBET
            self.NAV.set(Na)
            self.GYRO.set(G)
            self.PITCH.set(P)
            self.ROLL.set(R)
            self.GPS_H.set(GPSH)
            self.D_H.set(DH)

            self.POS_GYRO.set(Gc)
            self.POS_PITCH.set(Pc)
            self.POS_ROLL.set(Rc)
            self.POS_GPSH.set(GPSHc)
            self.POS_DH.set(DHc)

            ## Reading the defaults inputs for RMS
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            N_RMS = Parameters.iloc[2,6]
            G_RMS = Parameters.iloc[2,7]
            P_RMS = Parameters.iloc[2,8]
            R_RMS = Parameters.iloc[2,9]
            GPS_RMS = Parameters.iloc[2,10]
            D_RMS = Parameters.iloc[2,11]

            N_RMSc = Parameters.iloc[12,5]
            G_RMSc = Parameters.iloc[12,6]
            P_RMSc = Parameters.iloc[12,7]
            R_RMSc = Parameters.iloc[12,8]
            GPS_RMSc = Parameters.iloc[12,9]
            D_RMSc = Parameters.iloc[12,10]

            ## Setting defaults from Parameter file for RMS
            self.NAV_RMS.set(N_RMS)
            self.GYRO_RMS.set(G_RMS)
            self.PITCH_RMS.set(P_RMS)
            self.ROLL_RMS.set(R_RMS)
            self.GPSH_RMS.set(GPS_RMS)
            self.DH_RMS.set(D_RMS)

            self.POS_NRMS.set(N_RMSc)
            self.POS_GRMS.set(G_RMSc)
            self.POS_PRMS.set(P_RMSc)
            self.POS_RRMS.set(R_RMSc)
            self.POS_GPSHRMS.set(GPS_RMSc)
            self.POS_DHRMS.set(D_RMSc)

            ## ToolTips For Applanix Data
            A_F2 = Parameters.iloc[1,3]
            self.AUX_F2.set(A_F2)
            A_F3 = Parameters.iloc[1,4]
            self.AUX_F3.set(A_F3)

            ## ToolTips For Applanix Data
            tip_AUX2 = ToolTip(self.AUX_f2, str(self.AUX_F2.get()))
            tip_AUX3 = ToolTip(self.AUX_f3, str(self.AUX_F3.get()))

            try:
                ## Forget the POSMV Options
                self.POSMV_op.grid_forget()

            except AttributeError:
                pass


    def split_Project_Name(self):
        """Splits the Project name into 2 character strings
        Projectno_Location_Year_Vessel_Sytstem
        1. Projectno_Location_Year 2. Vessel_System"""

        Project_N = self.PROJECT_n.get()
        P_split =  Project_N.split('_')
        Project_N = P_split[0] + ('_') + P_split[1] + ('_') + P_split[2]
        HIPSFILE = P_split[3] + ('_') + P_split[4]
        return(Project_N, HIPSFILE, P_split)


    def Import_Auxiliary(self):
        """"""

        ## Get Default Values from Use Input
        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        HDCS_Folder = self.HDCS_D.get() ## HDCS_Data Dir
        crs = self.CRS_POS.get() ## Coordinate Ref System
        CRS = crs.partition(": ")[2]
        Vessel_F = self.VESSEL_N.get() ## Vessel File
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = self.Year.get() ## Survey Year
        JD = self.JULIAN_D.get() ## Julian Day
        Out = self.OUT_F.get() ## Processing Window Output Dir

        Allow_P = self.ALLOW_P.get() ## Allow Partially Covered Data
        Maximum_Gap = self.MAG.get() ## Maximum Allowable Gap
        AUX_F = self.AUX_F.get() ## AUX Dir
        REFWEEK = pd.to_datetime(DATES.datetime.strptime((self.Refweek.get()), "%m/%d/%y").strftime("%Y-%m-%d")).date() ## GPS Refference Week

        ## Getting Import POSMV Parameters
        if self.A_T.get()==1:
            L_POSF = listdir(AUX_F)

            Nav = self.NAV.get() ## Navigation
            Gyro = self.GYRO.get() ## Gyro
            Pitch = self.PITCH.get() ## Pitch
            Roll = self.ROLL.get() ## Roll
            GPS_h = self.GPS_H.get() ## GPS Height
            D_h = self.D_H.get() ## Delayed
            N_rms = self.NAV_RMS.get() ## Navigation RMS
            G_rms = self.GYRO_RMS.get() ## Gyro RMS
            P_rms = self.PITCH_RMS.get() ## Pitch RMS
            R_rms = self.ROLL_RMS.get() ## Roll RMS
            GPSH_rms = self.GPSH_RMS.get() ## GPS Height RMS
            DH_rms = self.DH_RMS.get() ## Delayed Heave RMS


            ## Creating Import_Aux.bat with POSMV Parameters
            A_Format = ('APP_POSMV')## Import Fromate - Applanix POSMV

            with open("Import_Aux_POSMV.bat", "w") as Import:
                    Import.write('@ECHO OFF' + '\n')
                    Import.write('@ECHO Importing POS' + '\n')
                    Import.write('cd '+ Caris + '\n')

                    Import.write('carisbatch --run ImportHIPSFromAuxiliary --input-format ' +
                                A_Format +  ' --input-crs ' + CRS + ' --maximum-gap ' +
                                Maximum_Gap + ' --reference-week ' +  str(REFWEEK) + ' --allow-partial')
                    if Nav==1:
                        Import.write(' --navigation ')
                    if self.POS_GYRO.get()==1:
                         Import.write(' --gyro ' + Gyro)
                    if  self.POS_PITCH.get==1:
                        Import.write( ' --pitch '+ Pitch)
                    if self.POS_ROLL.get()==1:
                        Import.write(' --roll ' + Roll)
                    if self.POS_GPSH.get()==1:
                        Import.write(' --gps-height ' + GPS_h)
                    if self.POS_DH.get()==1:
                        Import.write(' --delayed-heave ' + D_h)
                    if self.POS_NRMS.get()==1:
                        Import.write(' --navigation-rms ' + N_rms)
                    if self.POS_GRMS.get()==1:
                        Import.write(' --gyro-rms ' + G_rms)
                    if self.POS_PRMS.get()==1:
                        Import.write(' --pitch-rms ' + P_rms)
                    if self.POS_RRMS.get()==1:
                        Import.write(' --roll-rms ' + R_rms)
                    if self.POS_GPSHRMS.get()==1:
                        Import.write(' --gps-height-rms ' + GPSH_rms)
                    if self.POS_DHRMS.get()==1:
                        Import.write(' --delayed-heave-rms ' + DH_rms)

                    for file in L_POSF:
                        Import.write(' "' + AUX_F + '/' + file + '"')
                    Import.write(r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' +
                                 Vessel + ';Day=' + str(Year) + '-' + str(JD))
                    Import.write(' > ' + Out + '/' + JD + '/2.Import_POSMV' + JD + '_' + Year + '.txt' + '\n')

            ## Import Auxillairy Data using Caris Batch Through cmd
            p = S.Popen(['Import_Aux_POSMV.bat'])
            p.communicate()

        ## Getting Import SBET and RMS Parameters
        if self.A_T.get()==2:
            SBETF = self.AUX_F3.get()
            RMSF = self.AUX_F2.get()

            Nav = self.NAV.get() ## Navigation
            Gyro = self.GYRO.get() ## Gyro
            Pitch = self.PITCH.get() ## Pitch
            Roll = self.ROLL.get() ## Roll
            GPS_h = self.GPS_H.get() ## GPS Height
            D_h = self.D_H.get() ## Delayed
            N_rms = self.NAV_RMS.get() ## Navigation RMS
            G_rms = self.GYRO_RMS.get() ## Gyro RMS
            P_rms = self.PITCH_RMS.get() ## Pitch RMS
            R_rms = self.ROLL_RMS.get() ## Roll RMS
            GPSH_rms = self.GPSH_RMS.get() ## GPS Height RMS
            DH_rms = self.DH_RMS.get() ## Delayed Heave RMS

            ## Creating Import_Aux.bat with SBET Parameters
            A_Format = ('APP_SBET')
            A_Format2 = ('APP_RMS')
            with open("Import_Aux_SBET_RMS.bat", "w") as Import:
                    Import.write('@ECHO OFF' + '\n')
                    Import.write('@ECHO Importing SBET and RMS' + '\n')
                    Import.write('cd '+ Caris + '\n')

                    Import.write('carisbatch --run ImportHIPSFromAuxiliary --input-format ' +
                                 A_Format +  ' --input-crs ' + CRS + ' --maximum-gap ' +
                                 Maximum_Gap + ' --reference-week ' +  str(REFWEEK) + ' --allow-partial')
                    if Nav==1:
                        Import.write(' --navigation ')
                    if self.POS_GYRO.get()==1:
                         Import.write(' --gyro ' + Gyro)
                    if  self.POS_PITCH.get==1:
                        Import.write( ' --pitch '+ Pitch)
                    if self.POS_ROLL.get()==1:
                        Import.write(' --roll ' + Roll)
                    if self.POS_GPSH.get()==1:
                        Import.write(' --gps-height ' + GPS_h)

                    Import.write(' ' + SBETF)
                    Import.write(r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips')
                    Import.write(' > ' + Out + '/' + JD + '/2.Import_RMS' + JD + '_' + Year + '.txt' + '\n')

            ## Creating RMS Parameters
                    Import.write('carisbatch --run ImportHIPSFromAuxiliary --input-format ' +
                                 A_Format2)
                    if self.POS_NRMS.get()==1:
                        Import.write(' --navigation-rms ' + N_rms)
                    if self.POS_GRMS.get()==1:
                        Import.write(' --gyro-rms ' + G_rms)
                    if self.POS_PRMS.get()==1:
                        Import.write(' --pitch-rms ' + P_rms)
                    if self.POS_RRMS.get()==1:
                        Import.write(' --roll-rms ' + R_rms)
                    if self.POS_GPSHRMS.get()==1:
                        Import.write(' --gps-height-rms ' + GPSH_rms)

                    Import.write(' ' + RMSF)
                    Import.write(r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips')
                    Import.write(' > ' + Out + '/' + JD + '/2.Import_RMS' + JD + '_' + Year + '.txt' + '\n')

            p = S.Popen(['Import_Aux_SBET_RMS.bat'])
            p.communicate()


    def Load_Hips_Project_Par(self):
        """Loads default or user saved parameters for Caris HIPS Proccesing."""

        chdir(owd)##Application Dir

        ###Reading defaults or user saved inputs for HIPS Proccessing
        Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
        RAW = Parameters.iloc[0,0]
        HDCS = Parameters.iloc[0,1]
        PROJECT = Parameters.iloc[0,2]
        VESSEL = Parameters.iloc[0,3]
        CRS = Parameters.iloc[0,4]
        JULIAN_DAY = Parameters.iloc[0,5]
        YEAR = Parameters.iloc[0,6]
        CON_NAV = int(Parameters.iloc[0,7])
        OUT = Parameters.iloc[0,8]

        ##Setting defaults from Parameter file for HIPS Proccessing
        self.RAW_F.set(RAW)
        self.HDCS_D.set(HDCS)
        self.PROJECT_N.set(PROJECT)
        self.VESSEL_N.set(VESSEL)
        self.CRS_O.set(CRS)
        self.JULIAN_D.set(JULIAN_DAY)
        self.YEAR.set(YEAR)
        self.CONVERT_N.set(CON_NAV)
        self.OUT_F.set(OUT)

        ##ToolTips for HIPS Project Options
        tip_RAW = ToolTip(self.RAW_f, (self.RAW_F.get()))
        tip_HDCS = ToolTip(self.HDCS_d, (self.HDCS_D.get()))
        tip_Vessel = ToolTip(self.VESSEL_n, (self.VESSEL_N.get()))
        tip_Out = ToolTip(self.OUT_f, (self.OUT_F.get()))


##    def Copy_HIPS(self):
##
##        
##        CopyToHips_ops = LabelFrame(frame11, text="Copy to Hips", foreground="blue")
##        CopyToHips_ops.grid(row=0, column=0, padx=1, sticky=W)
##
##        self.INHIPS = StringVar()
##        self.Inhips = Entry(CopyToHips_ops, width=38, textvariable=self.INHIPS)
##        self.Inhips_text = Label(CopyToHips_ops, text="Input Project")
##        self.Inhips_text.grid(row=1, column=0, sticky=W)
##        self.Inhips.grid(row=1, column=1, sticky=W)
##        self.ButtonInhips = Button(CopyToHips_ops, text="...", height=0,
##                              command=self.Search_TrackLines)
##        self.ButtonInhips.grid(row=1, column=2, sticky=W, padx=2)
##
##        self.OUTHIPS = StringVar()
##        self.Outhips = Entry(CopyToHips_ops, width=38, textvariable=self.OUTHIPS)
##        self.Outhips_text = Label(CopyToHips_ops, text="Destination\nProject")
##        self.Outhips_text.grid(row=2, column=0, sticky=W)
##        self.Outhips.grid(row=2, column=1, sticky=W)
##        self.ButtonOuthips = Button(CopyToHips_ops, text="...", height=0,
##                              command=self.Search_TrackLines)
##        self.ButtonOuthips.grid(row=2, column=2, sticky=W, padx=2)
##    
##        Copy_ToHips_TL = LabelFrame(frame11, text="Trackline Selection", foreground="blue")
##        Copy_ToHips_TL.grid(row=1, column=0, padx=1, sticky=W)
##
##        self.TL_Dir = StringVar()
##        self.tl_dir = Entry(Copy_ToHips_TL, width=38, textvariable=self.TL_Dir)
##        self.tl_dir_text = Label(Copy_ToHips_TL, text="TrackLine Directory")
##        self.tl_dir_text.grid(row=1, column=0, sticky=W)
##        self.tl_dir.grid(row=1, column=1, sticky=W)
##        self.ButtonCTH = Button(Copy_ToHips_TL, text="...", height=0,
##                              command=self.Search_TrackLines)
##        self.ButtonCTH.grid(row=1, column=2, sticky=W, padx=2)
##
##        self.listbox = Listbox(Copy_ToHips_TL, height=25, width=38, selectmode=EXTENDED)
##        self.listbox.grid(row=2, column=1, sticky=W)

        

##        self.STlat_text = Label(Copy_ToHips, text="Check boxes \n to backup Projects")
##        self.STlat_text.grid(row=0, column=0, sticky=W)

##        self.CMAIN = IntVar()
##        self.CMain = Checkbutton(Copy_ToHips, onvalue=1, offvalue=0, variable=self.CMAIN, text= "Copy Main Project")
##                                    #command=self.TIDES)
##        self.CMain.grid(row=1, column=0, sticky=W)
##
##        self.CSIDE = IntVar()
##        self.CSide = Checkbutton(Copy_ToHips, onvalue=1, offvalue=0, variable=self.CSIDE, text= "Copy Side Project")
##                                    #command=self.TIDES)
##        self.CSide.grid(row=1, column=1, sticky=W)
##
##        self.Button_Copyto = Button(Copy_ToHips, text="Run Copyto", height=0)
##                               #command=self.PosPAC_Processing)
##        self.Button_Copyto.grid(row=2, column=0, sticky=W, padx=2)

##    def CopytoHIPS(self):
##
##        TL_Dir = self.TL_Dir.get()
##        TL_Dir_list = listdir(TL_Dir)
##        Tracklines = [self.listbox.get(idx) for idx in self.listbox.curselection()]
##        InputHIPS = self.INHIPS.get()
##        OutputHIPS = self.OUTHIPS.get()
##
##        if CRS=='EPSG:7912@2010':
##            CRS2 = 'EPSG:5937'
##        else:
##            CRS2 = CRS
##
##            with open("CreateHIPSFile.bat", "w") as Import:
##                Import.write('@ECHO OFF' + '\n')
##                Import.write('@ECHO Creating HIPS File' + '\n')
##                Import.write('cd '+ Caris + '\n')
##                Import.write('carisbatch --run CreateHIPSFile --output-crs ' +
##                             CRS2 +
##                             r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips' +
##                             '> ' + Out + '/' + 'CreateHIPSFile.txt')
##
##            p = S.Popen(['CreateHIPSFile.bat'])
##            p.communicate()
##
##
##        with open("CopytoHips.bat", "w") as Import:
##            Import.write('@ECHO OFF' + '\n')
##            Import.write('@ECHO Copying Tracklines' + '\n')
##            Import.write('cd '+ Caris + '\n')
##            Import.write('carisbatch --run CopyHIPSToHIPS "'file:///' + C:/Users/LegerMi/Documents/CARIS/HIPS and SIPS/Copyt2' +
##                         '/Copyto.hips?Line=')
##            i = 1
##            
##            for line in tracklines:
##                if i == lengthTracklines:
##                    Import.write(line + '" ')
##                else:
##                    Import.write(line + '&')
##                i = i + 1
##
##            Import.write('C:/Users/LegerMi/Desktop/DA_PDAP_2020/9XXXXXXX_PDAP_2020/Processed_Data/Kestrel_EM2040C/Kestrel_EM2040C.hips')

        


        
    def POSPAC_Par(self):

        w_E = 45

        pos_pac = LabelFrame(frame10, text="POSPAC Proccesing", foreground="blue")
        pos_pac.grid(row=0, column=0, padx=1, sticky=W)

        self.POSDIR = StringVar()
        self.POSDir = Entry(pos_pac, width=w_E, textvariable=self.POSDIR)
        self.POSDir_text = Label(pos_pac, text="POS File Folder")
        self.POSDir_text.grid(row=0, column=0, sticky=W)
        self.POSDir.grid(row=0, column=1, sticky=W)
        self.Buttonpospac = Button(pos_pac, text="...", height=0,
                              command=self.Search_Aux_Data)
        self.Buttonpospac.grid(row=0, column=2, sticky=W, padx=2)

        self.STID = StringVar()
        self.STid = Entry(pos_pac, width=20, textvariable=self.STID)
        self.STid_text = Label(pos_pac, text="GNSS Station Name")
        self.STid_text.grid(row=2, column=0, sticky=W)
        self.STid.grid(row=2, column=1, sticky=W)

        self.STLAT = StringVar()
        self.STlat= Entry(pos_pac, width=20, textvariable=self.STLAT)
        self.STlat_text = Label(pos_pac, text="PPP Lat D-M-S")
        self.STlat_text.grid(row=3, column=0, sticky=W)
        self.STlat.grid(row=3, column=1, sticky=W)

        self.STLONG = StringVar()
        self.STlong= Entry(pos_pac, width=20, textvariable=self.STLONG)
        self.STlong_text = Label(pos_pac, text="PPP Long D-M-S")
        self.STlong_text.grid(row=4, column=0, sticky=W)
        self.STlong.grid(row=4, column=1, sticky=W)

        self.STH = StringVar()
        self.STh= Entry(pos_pac, width=20, textvariable=self.STH)
        self.STh_text = Label(pos_pac, text="PPP Height Metres")
        self.STh_text.grid(row=5, column=0, sticky=W)
        self.STh.grid(row=5, column=1, sticky=W)

        self.GNSSFile = StringVar()
        self.GnssFile= Entry(pos_pac, width=w_E, textvariable=self.GNSSFile)
        self.GnssFile_text = Label(pos_pac, text="GNSS Receiver Observation File")
        self.GnssFile_text.grid(row=1, column=0, sticky=W)
        self.GnssFile.grid(row=1, column=1, sticky=W)
        self.Buttongnss = Button(pos_pac, text="...", height=0,
                              command=self.Search_GNSS_Obs)
        self.Buttongnss.grid(row=1, column=2, sticky=W, padx=2)

        self.REF_FRAME = StringVar()
        ref_frame_op = ['NAD83 (CSRS):Ellipsoid:GRS 1980:Epoch:2010']
                      #'ITRF 2014: Ellipsoid: ?: Epoch: 2010']

        self.ref_frame_op = ttk.Combobox(pos_pac, values=ref_frame_op, width=42, textvariable=self.REF_FRAME)
        self.ref_frame_text = Label(pos_pac, text="Corrections/Base Station Reference Frame")
        self.ref_frame_text.grid(row=6, column=0, sticky=W)
        self.ref_frame_op.grid(row=6, column=1, sticky=W+E, padx=0)

        self.Button_POSPAC = Button(pos_pac, text="Run POSPAC", height=0,
                               command=self.PosPAC_Processing)
        self.Button_POSPAC.grid(row=7, column=0, sticky=W, padx=2)



        Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
        PDir = Parameters.iloc[13,0]
        OBSf = Parameters.iloc[13,1]
        STid = Parameters.iloc[13,2]
        RefFrame = Parameters.iloc[13,3]


        self.POSDIR.set(PDir)
        self.GNSSFile.set(OBSf)
        self.STID.set(STid)
        self.REF_FRAME.set(RefFrame)


    def PosPAC_Processing(self):

        POSDir = self.POSDIR.get()
        StationFile = self.GNSSFile.get()
        POSFiles = listdir(POSDir)
        StationID = self.STID.get()
        LONGDMS = self.STLONG.get()
        LATDMS = self.STLAT.get()
        LATDD = DMS_to_DD(LATDMS)
        LONGDD = DMS_to_DD(LONGDMS)
        LATRAD = DD_to_Rads(LATDD)
        LONGRAD = DD_to_Rads(LONGDD)
        HeightM = self.STH.get()

        REF_FRAME = self.REF_FRAME.get()
        Output_ref_Datum = REF_FRAME.split(":")[0]
        print(Output_ref_Datum)
        Output_Ellipsoid = REF_FRAME.split(":")[2]
        if Output_ref_Datum == ('NAD83 (CSRS)'):
            Station_ref_frame = Output_ref_Datum.replace('NAD83 (CSRS)', 'NAD83_CSRS')
            Station_Ellipsoid =  Output_Ellipsoid.replace('GRS 1980', 'GRS_1980')
##        elif Output_ref_Datum.get() == ('ITRF 2014'):
##            Station_ref_frame = Output_ref_Datum.replace('ITRF_2014')
##            Station_Ellipsoid =  Output_ref_Datum.replace('')
        Epoch = REF_FRAME.split(":")[4]


        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]
        JD = self.JULIAN_D.get()
        Year = self.Year.get()
        Out = self.OUT_F.get()
        OutFile = JD + '_' + HIPSFILE + '_SBET'

        POSPACBatch = ('POSPACBatch.txt')
        tree = ET.parse(POSPACBatch)

        Name = ('JD' + JD + '_' + HIPSFILE)
        FirstPosFile = POSFiles[0]
        LastPosFile = POSFiles[-1]

        POSDir = POSDir.replace('/','\\')

        ##Project Settings
        tree.find('.//Name').text = Name
        tree.find('.//FirstPosFile').text = (POSDir + '\\' + FirstPosFile)
        tree.find('.//LastPosFile').text = (POSDir + '\\' + LastPosFile)

        ## GNSS Station Settings
        tree.find('.//StationID').text = StationID
        tree.find('.//DataFile').text = StationFile
        tree.find('.//Frame').text = Station_ref_frame
        tree.find('.//Bsci/Ellipsoid').text = Station_Ellipsoid
        tree.find('.//Bsci/Epoch').text = Epoch
        st_coords = tree.findall('.//Posllh//double')
        st_coords[0].text = str(LATRAD)
        st_coords[1].text = str(-1 * (LONGRAD))
        st_coords[2].text = str(HeightM)

        ## Export Settings
        tree.find('.//OutputFile').text = OutFile
        tree.find('.//Export//Ellipsoid').text = Output_Ellipsoid
        tree.find('.//Export//Datum').text = Output_ref_Datum
        tree.find('.//Export//TargetEpoch').text = Epoch

        ## Write settings to Pospac Batch
        tree.write(POSPACBatch)

        ## tree.find('.//AntennaManufacturer).text = AntennaMan
        ## tree.find('.//AntennaType).text = AntennaType
        ## tree.find('.//GnssMode).text = GnssMode


        data = data2 = " "
        with open('POSPACBatch.txt') as P:
            data2 = P.read()
        with open('1_Header.txt') as H:
            data = H.read()
        data += "\n"
        data += data2
        with open ('temp.txt', 'w') as Cp:
            Cp.write(data)
        with open('temp.txt') as C, open(Name + '.posbat', 'w') as Final:
            lines = C.readlines()
            Final.writelines(lines[0:2])
            Final.writelines(lines[3:])
        remove('temp.txt')


        with open("POSPAC.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Running POSPAC' + '\n')
                Import.write('cd '+ PosPac + '\n')
                Import.write('PospacBatch -b ' + owd + '/' + Name + '.posbat > ' + Out + '/' + JD + '/' +
                             '9.POSPACProccessing_' + JD + '_' + Year + '.txt')
                
##        p = S.Popen(['POSPAC.bat'])## Doesnt work on DFO Imaged Machines
##        p.communicate()


    def general_hips_options(self):
        """Sets user inputs for Caris Project Parameters"""

        chdir(owd)##Application Directory
        w_E = 45

        ##Creating HIPS Processing User Input Options
        hips_op = LabelFrame(frame1, text="Hips Project", foreground="blue")
        hips_op.grid(row=0, column=0, padx=1, sticky=W)

        ##RAW Sonar Data (.all,.gsf,.xtf)
        self.RAW_F = StringVar()
        self.RAW_f = Entry(hips_op, width=w_E, textvariable=self.RAW_F,)
        self.RAW_text = Label(hips_op, text="Raw Files")
        self.RAW_text.grid(row=0, column=0, sticky=W)
        self.RAW_f.grid(row=0, column=1, sticky=W)
        self.Button1 = Button(hips_op, text="...", height=0,
                              command=self.Search_RAW_Data)
        self.Button1.grid(row=0, column=2, sticky=W, padx=2)


        ##HDCS_Data Folder Location
        self.HDCS_D = StringVar()
        self.HDCS_d = Entry(hips_op, width=w_E, textvariable=self.HDCS_D)
        self.HDCS_text = Label(hips_op, text="Proccessing folder")
        self.HDCS_text.grid(row=1, column=0, sticky=W)
        self.HDCS_d.grid(row=1, column=1, sticky=W)
        self.Button2 = Button(hips_op, text="...", height=0,
                              command=self.Search_HDCS_Data)
        self.Button2.grid(row=1, column=2, sticky=W, padx=2)

        ## Output Folder for Script Output
        self.OUT_F = StringVar()
        self.OUT_f = Entry(hips_op, width=w_E, textvariable=self.OUT_F)
        self.OUT_f_text = Label(hips_op, text="Output Folder")
        self.OUT_f_text.grid(row=2, column=0, sticky=W)
        self.OUT_f.grid(row=2, column=1, sticky=W)
        self.Button2 = Button(hips_op, text="...", height=0,
                              command=self.Search_OUTPUT)
        self.Button2.grid(row=2, column=2, sticky=W, padx=2)

        ## CHS Project Number
        self.PROJECT_N = StringVar()
        self.PROJECT_n = Entry(hips_op, width=w_E, textvariable=self.PROJECT_N)
        self.PROJECT_text = Label(hips_op, text="Project Name")
        self.PROJECT_text.grid(row=3, column=0, sticky=W)
        self.PROJECT_n.grid(row=3, column=1, sticky=W)

        ## Vessel File Name (inlcuing .hvf)
        self.VESSEL_N = StringVar()
        self.VESSEL_n = Entry(hips_op, width=w_E, textvariable=self.VESSEL_N)
        self.VESSEL_text = Label(hips_op, text="Select Vessel File")
        self.VESSEL_text.grid(row=4, column=0, sticky=W)
        self.VESSEL_n.grid(row=4, column=1, sticky=W)
        self.Button3 = Button(hips_op, text="...", height=0,
                              command=self.Search_VesselFile)
        self.Button3.grid(row=4, column=2, sticky=W, padx=2)

        ##Project and Hips Data CSRS
        self.CRS_O = StringVar()
        crs_op = ['NAD83(CSRS)/UTM Zone 19N: EPSG:2960@2010',
                  'NAD83(CSRS)/UTM Zone 20N: EPSG:2961@2010',
                  'NAD83(CSRS)/UTM Zone 21N: EPSG:2962@2010',
                  'WGS84/EPSG Canada Polar Stereographic: EPSG:5937@2010',
                  'WGS84/UTM Zone 19N: EPSG:32619@2010',
                  'WGS84/UTM Zone 20N: EPSG:32620@2010',
                  'WGS84/UTM Zone 21N: EPSG:32621@2010']

        self.CRS_op = ttk.Combobox(hips_op, values=crs_op, width=42, textvariable=self.CRS_O)
        self.CRS_text = Label(hips_op, text="Choose CRS")
        self.CRS_text.grid(row=5, column=0, sticky=W)
        self.CRS_op.grid(row=5, column=1, sticky=W+E, padx=0)

        ##Julian Day of RAW
        self.JULIAN_D = StringVar()
        self.JULIAN_d = Entry(hips_op, width=5, textvariable=self.JULIAN_D)
        self.JULIAN_text = Label(hips_op, text="Julian Day")
        self.JULIAN_text.grid(row=6, column=0, sticky=W)
        self.JULIAN_d.grid(row=7, column=0, sticky=W)

        ##Year of RAW
        self.YEAR = StringVar()
        self.Year = Entry(hips_op, width=5, textvariable=self.YEAR)
        self.Year_text = Label(hips_op, text="Year")
        self.Year_text.grid(row=6, column=1, sticky=W)
        self.Year.grid(row=7, column=1, sticky=W)

        ##Covert Navigation
        self.CONVERT_N = IntVar()
        self.CONVERT_n = Checkbutton(hips_op, variable=self.CONVERT_N, text= "Convert Navigation",
                                     state='disabled')
        self.CONVERT_n.grid(row=8, column=0, sticky=W)

        ##Intial Run to Create HIPS File
        self.IntialRun = IntVar()
        self.Intial_Run = Checkbutton(hips_op, variable=self.IntialRun, text= "Intial Run")
        self.Intial_Run.grid(row=8, column=1, sticky=W)

        ##Creating Radio Box for Sensor Data Import
        self.Sensor_type = LabelFrame(frame1, text="Sensor Data", foreground="blue")
        self.Sensor_type.grid(row=1, column=0, padx=1, sticky=W)


        self.STYPE = IntVar()
        self.Stype = Checkbutton(self.Sensor_type, onvalue=1, offvalue=0, variable=self.STYPE, text= "Import RAW Files",
                                    command=self.RAW_Sensor)
        self.Stype.grid(row=0, column=0, sticky=W)


        self.S_T=IntVar()

        self.rb1 = Radiobutton(self.Sensor_type, text= "KONGSBERG\nALL", variable=self.S_T,
                    value=1, command=
                    self.Load_RAW_Par, state='disabled').grid(row=1, column=0, sticky= W)
        self.rb2 = Radiobutton(self.Sensor_type, text= "R2 SONIC\nGSF", variable=self.S_T,
                    value=2, command=
                    self.Load_RAW_Par, state='disabled').grid(row=1, column=2, sticky=W)
        self.rb3 = Radiobutton(self.Sensor_type, text= "TRITON\nXTF", variable=self.S_T,
                    value=3, command=
                    self.Load_RAW_Par,state='disabled').grid(row=1, column=3, sticky=W)
        self.rb4 = Radiobutton(self.Sensor_type, text= "Teledyne\nS7K", variable=self.S_T,
                    value=4, command=
                    self.Load_RAW_Par,state='disabled').grid(row=1, column=4, sticky=W)
        self.rb5 = Radiobutton(self.Sensor_type, text= "KONGSBERG\nKMALL", variable=self.S_T,
                    value=5, command=
                    self.Load_RAW_Par, state='disabled').grid(row=1, column=5, sticky= W)

        ##Creating Check Box for Applanix Data Import
        self.Applanix_Data = LabelFrame(frame1, text="Applanix Data", foreground="blue")
        self.Applanix_Data.grid(row=2, column=0, padx=1, sticky=W)

        self.ATYPE = IntVar()
        self.Atype = Checkbutton(self.Applanix_Data, onvalue=1, offvalue=0, variable=self.ATYPE, text= "Import POS",
                                    command=self.Applanix)
        self.Atype.grid(row=0, column=0, sticky=W)

        self.A_T=IntVar()
        self.rb4 = Radiobutton(self.Applanix_Data, text= "POSMV", variable=self.A_T,
                    value=1, command=
                    self.Load_Auxiliary_Par, state='disabled').grid(row=1, column=0, sticky= W, padx=1)
        self.rb5 = Radiobutton(self.Applanix_Data, text= "SBET & RMS", variable=self.A_T,
                    value=2, command=
                    self.Load_Auxiliary_Par, state='disabled').grid(row=1, column=2, sticky=W, padx=1)

        ##Creating Radio Box for Tide Type
        self.GEO_REF = LabelFrame(frame1, text="Geo-Referencing", foreground="blue")
        self.GEO_REF.grid(row=3, column=0, padx=1, sticky=W)

        self.TTYPE = IntVar()
        self.Ttype = Checkbutton(self.GEO_REF, onvalue=1, offvalue=0, variable=self.TTYPE, text= "Apply Tide",
                                    command=self.TIDES)
        self.Ttype.grid(row=0, column=0, sticky=W)

        self.T_T=IntVar()
        self.rb6 = Radiobutton(self.GEO_REF, text= "GPS Tide", variable=self.T_T,
                    value=1, command=
                    self.Load_Tide_Par, state='disabled').grid(row=1, column=0, sticky= W, padx=1)
        self.rb7 = Radiobutton(self.GEO_REF, text= "Observed/Predicted", variable=self.T_T,
                    value=2, command=
                    self.Load_Tide_Par, state='disabled').grid(row=1, column=1, sticky=W, padx=1)

        ##Compute TPU
        self.COMP_TPU = IntVar()
        self.Comp_TPU = Checkbutton(self.GEO_REF, onvalue=1, offvalue=0, variable=self.COMP_TPU, text= "Compute\nTPU",
                                    command=self.Load_TPU_Par)
        self.Comp_TPU.grid(row=2, column=0, sticky=W)

        ##Apply SVP
        self.APPLY_SVP = IntVar()
        self.Apply_SVP = Checkbutton(self.GEO_REF, onvalue=1, offvalue=0, variable=self.APPLY_SVP, text= "Apply SVP",
                                    command=self.Load_SVP_Par)
        self.Apply_SVP.grid(row=2, column=1, sticky=W, padx=1)

        ##Create/Add to HIPS Grid
        self.H_GRID = LabelFrame(frame1, text="HIPS Coverage (Surface Creation)", foreground="blue")
        self.H_GRID.grid(row=5, column=0, padx=1, sticky=W)

        self.GRIDS = IntVar()
        self.Grids = Checkbutton(self.H_GRID, onvalue=1, offvalue=0, variable=self.GRIDS, text= "Create Grids",
                                    command=self.SURFACE)
        self.Grids.grid(row=0, column=0, sticky=W)

        self.GRID = IntVar()
        self.rb8 = Radiobutton(self.H_GRID, text= "Create HIPS Grid", variable=self.GRID,
                    value=1, command=
                    self.Load_GRID_Par, state='disabled').grid(row=1, column=0, sticky= W, padx=1)
        self.rb9 = Radiobutton(self.H_GRID, text= "Create/Add to HIPS Grid", variable=self.GRID,
                    value=2, command=
                    self.Load_GRID_Par, state='disabled').grid(row=1, column=2, sticky=W, padx=1)

        ##Merge Tracklines
        self.MERGE_TRACK = IntVar()
        self.Merge_Track = Checkbutton(self.GEO_REF, onvalue=1, offvalue=0, variable=self.MERGE_TRACK, text= "Merge Tracklines",
                                       command=self.Loads_MergeTrack)
        self.Merge_Track.grid(row=2, column=2, sticky=W, padx=1)
##        ##Noise Classifier
##        N_C = LabelFrame(frame1, text="Sonar Noise Classifier", foreground="blue")
##        N_C.grid(row=6, column=0, padx=1, sticky=W)
##
##        self.Noise_C = IntVar()
##        self.Noise_c = Checkbutton(N_C, onvalue=1, offvalue=0, variable=self.Noise_C, text= "Classify Noise",
##                                   command=self.Load_CARIS_MIRA)
##        self.Noise_c.grid(row=1, column=0, sticky=W)

##        ##Create BS
##        B_S = LabelFrame(frame1, text="Create BackScatter Mosiac", foreground="blue")
##        B_S.grid(row=7, column=0, padx=1, sticky=W)
##
##        self.BACK_S= IntVar()
##        self.Back_S = Checkbutton(B_S, onvalue=1, offvalue=0, variable=self.BACK_S, text= "Create BackScatter")
##        self.Back_S.grid(row=1, column=0, sticky=W)

    def SURFACE(self):
        
        if self.GRIDS.get()==1:
            self.GRID = IntVar()
            self.rb8 = Radiobutton(self.H_GRID, text= "Create HIPS Grid", variable=self.GRID,
                    value=1, command=
                    self.Load_GRID_Par).grid(row=1, column=0, sticky= W, padx=1)
            self.rb9 = Radiobutton(self.H_GRID, text= "Create/Add to HIPS Grid", variable=self.GRID,
                    value=2, command=
                    self.Load_GRID_Par).grid(row=1, column=2, sticky=W, padx=1)
        else:
            self.GRID = IntVar()
            self.rb8 = Radiobutton(self.H_GRID, text= "Create HIPS Grid", variable=self.GRID,
                    value=1, command=
                    self.Load_GRID_Par, state='disabled').grid(row=1, column=0, sticky= W, padx=1)
            self.rb9 = Radiobutton(self.H_GRID, text= "Create/Add to HIPS Grid", variable=self.GRID,
                    value=2, command=
                    self.Load_GRID_Par, state='disabled').grid(row=1, column=2, sticky=W, padx=1)

            try:
                ##Forget Create Grid
                self.CREATEGRID_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Create Grid 2
                self.CREATEGRID2_op.grid_forget()
            except AttributeError:
                pass

      
    def TIDES(self):

        if self.TTYPE.get()==1:
            self.T_T=IntVar()
            self.rb6 = Radiobutton(self.GEO_REF, text= "GPS Tide", variable=self.T_T,
                        value=1, command=
                        self.Load_Tide_Par).grid(row=1, column=0, sticky= W, padx=1)
            self.rb7 = Radiobutton(self.GEO_REF, text= "Observed/Predicted", variable=self.T_T,
                        value=2, command=
                        self.Load_Tide_Par).grid(row=1, column=1, sticky=W, padx=1)
        else:
            
            self.T_T=IntVar()
            self.rb6 = Radiobutton(self.GEO_REF, text= "GPS Tide", variable=self.T_T,
                        value=1, command=
                        self.Load_Tide_Par, state='disabled').grid(row=1, column=0, sticky= W, padx=1)
            self.rb7 = Radiobutton(self.GEO_REF, text= "Observed/Predicted", variable=self.T_T,
                        value=2, command=
                        self.Load_Tide_Par, state='disabled').grid(row=1, column=1, sticky=W, padx=1)
            try:
                ## Forget the Observed/Predicted Tides
                self.OPTide_op.grid_forget()
            except AttributeError:
                pass

            try:
                ## Forget the GPS Tide Options
                self.GPSTide_op.grid_forget()
            except AttributeError:
                pass

    
    def Applanix(self):

        if self.ATYPE.get()==1:
            
            self.A_T=IntVar()
        
            self.rb4 = Radiobutton(self.Applanix_Data, text= "POSMV", variable=self.A_T,
                        value=1, command=
                        self.Load_Auxiliary_Par).grid(row=1, column=0, sticky= W, padx=1)
            self.rb5 = Radiobutton(self.Applanix_Data, text= "SBET & RMS", variable=self.A_T,
                        value=2, command=
                        self.Load_Auxiliary_Par).grid(row=1, column=2, sticky=W, padx=1)
        else:
            self.A_T=IntVar()
            self.rb4 = Radiobutton(self.Applanix_Data, text= "POSMV", variable=self.A_T,
                    value=1, command=
                    self.Load_Auxiliary_Par, state='disabled').grid(row=1, column=0, sticky= W, padx=1)
            self.rb5 = Radiobutton(self.Applanix_Data, text= "SBET & RMS", variable=self.A_T,
                    value=2, command=
                    self.Load_Auxiliary_Par, state='disabled').grid(row=1, column=2, sticky=W, padx=1)

            try:
                ## Forget the SBET & RMS Options
                self.SBET_RMS_op.grid_forget()
                self.AUX_f2.grid_forget()
                self.Button_AUX2.grid_forget()
                self.AUX_text2.grid_forget()
                
            except AttributeError:
                pass

            try:
                ## Forget the POSMV Options
                self.POSMV_op.grid_forget()

            except AttributeError:
                pass
            


    def RAW_Sensor(self):

        if self.STYPE.get()==1:
            
            self.S_T=IntVar()
            self.rb1 = Radiobutton(self.Sensor_type, text= "KONGSBERG\nALL", variable=self.S_T,
                        value=1, command=
                        self.Load_RAW_Par).grid(row=1, column=0, sticky= W)
            self.rb2 = Radiobutton(self.Sensor_type, text= "R2 SONIC\nGSF", variable=self.S_T,
                        value=2, command=
                        self.Load_RAW_Par).grid(row=1, column=2, sticky=W)
            self.rb3 = Radiobutton(self.Sensor_type, text= "TRITON\nXTF", variable=self.S_T,
                        value=3, command=
                        self.Load_RAW_Par).grid(row=1, column=3, sticky=W)
            self.rb4 = Radiobutton(self.Sensor_type, text= "Teledyne\nS7K", variable=self.S_T,
                        value=4, command=
                        self.Load_RAW_Par).grid(row=1, column=4, sticky=W)
            self.rb5 = Radiobutton(self.Sensor_type, text= "KONGSBERG\nKMALL", variable=self.S_T,
                        value=5, command=
                        self.Load_RAW_Par).grid(row=1, column=5, sticky=W)
           

        else:
            self.S_T=IntVar()

            self.rb1 = Radiobutton(self.Sensor_type, text= "KONGSBERG\nALL", variable=self.S_T,
                        value=1, command=
                        self.Load_RAW_Par, state='disabled').grid(row=1, column=0, sticky= W)
            self.rb2 = Radiobutton(self.Sensor_type, text= "R2 SONIC\nGSF", variable=self.S_T,
                        value=2, command=
                        self.Load_RAW_Par, state='disabled').grid(row=1, column=2, sticky=W)
            self.rb3 = Radiobutton(self.Sensor_type, text= "TRITON\nXTF", variable=self.S_T,
                        value=3, command=
                        self.Load_RAW_Par,state='disabled').grid(row=1, column=3, sticky=W)
            self.rb4 = Radiobutton(self.Sensor_type, text= "Teledyne\nS7K", variable=self.S_T,
                        value=34, command=
                        self.Load_RAW_Par,state='disabled').grid(row=1, column=4, sticky=W)
            self.rb5 = Radiobutton(self.Sensor_type, text= "KONGSBERG\nKMALL", variable=self.S_T,
                        value=34, command=
                        self.Load_RAW_Par,state='disabled').grid(row=1, column=5, sticky=W)
            

            try:
                ## Forget R2 Sonic options
                self.R2_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget XTF 0ptions
                 self.X_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg 0ptions
                 self.K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Teledyne S7K 0ptions
                 self.S7K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget KMALL Options
                self.KMALL_op.grid_forget()
            except:
                pass
            

        
    def Load_CARIS_MIRA(self):

        chdir(owd) ##Application Dir

        if self.Noise_C.get()==1:

            ##Creating IMPORT TO HIPS Options
            self.Mira_op = LabelFrame(frame8, text="Sonar Noise Classifier", foreground="blue")
            self.Mira_op.grid(row=0, column=0, padx=1, sticky=W)


            self.LEVEL_D = StringVar()
            self.LEVEL_d = Entry(self.Mira_op, width=10, textvariable=self.LEVEL_D, state='disabled')
            self.LEVEL_d_text = Label(self.Mira_op, text="Level of Detail")
            self.LEVEL_d_text.grid(row=1, column=0, sticky=W)
            self.LEVEL_d.grid(row=1, column=1, sticky=W, padx=1)

            self.FILTER_THRES = StringVar()
            self.FILTER_Thres = Entry(self.Mira_op, width=10, textvariable=self.FILTER_THRES, state='disabled')
            self.FILTER_Thres_text = Label(self.Mira_op, text="Noise Confidence Filter")
            self.FILTER_Thres_text.grid(row=2, column=0, sticky=W)
            self.FILTER_Thres.grid(row=2, column=1, sticky=W, padx=1)

            self.Mira_URL = StringVar()
            self.Mira_url = Entry(self.Mira_op, width=10, textvariable=self.Mira_URL, state='disabled')
            self.Mira_url_text = Label(self.Mira_op, text="MIRA URL")
            self.Mira_url_text.grid(row=3, column=0, sticky=W)
            self.Mira_url.grid(row=3, column=1, sticky=W, padx=1)

        else:
            try:
                ## Forget CARIS Classifier options
                self.Mira_op.grid_forget()
            except AttributeError:
                pass


    def SONAR_CLASS(self): ## Not running Code

        D = self.LEVEL_D.get()
        Thresh = self.FILTER_THRES.get()
        URL = self.MIRA_URL.get()

        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        HDCS_Folder = self.HDCS_D.get()
        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()


        with open("Sonar_Classifier.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Running CARIS MIRA' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run ClassifyHIPSNoise --level-of-detail ' + D +
                             ' --mira-url ' + url)
                if MIRA_Filter == 1:
                    Import.write(' --noise-confidence-filter-threshold ' + Thresh )

                Import.write(' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                             ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                             '1.Import_To_Hips_ALL_' + JD + '_' + Year + '.txt' + '\n')


    def Load_RAW_Par(self):

        chdir(owd) ##Application Dir

        if self.S_T.get()==1:

            ##Creating IMPORT TO HIPS Options
            self.K_op = LabelFrame(frame2, text="Import Kongsberg (.all) to HIPS", foreground="blue")
            self.K_op.grid(row=0, column=0, padx=1, sticky=W)

            ##Navigation Device
            self.Nav_D = StringVar()
            self.Nav_d = Entry(self.K_op, width=10, textvariable=self.Nav_D, state='disabled')
            self.Nav_d_text = Label(self.K_op, text="Navigation Device")
            self.Nav_d_text.grid(row=1, column=0, sticky=W)
            self.Nav_d.grid(row=1, column=1, sticky=W, padx=1)

            ##GPS Height Device
            self.GPSH_D = StringVar()
            self.GPS_h = Entry(self.K_op, width=10, textvariable=self.GPSH_D, state='disabled')
            self.GPS_h_text = Label(self.K_op, text="GPS Height Device")
            self.GPS_h_text.grid(row=2, column=0, sticky=W)
            self.GPS_h.grid(row=2, column=1, sticky=W, padx=1)

            ##GPS Height Device
            self.Heave_D = StringVar()
            self.Heave_d = Entry(self.K_op, width=10, textvariable=self.Heave_D, state='disabled')
            self.Heave_d_text = Label(self.K_op, text="Heave Device")
            self.Heave_d_text.grid(row=3, column=0, sticky=W)
            self.Heave_d.grid(row=3, column=1, sticky=W, padx=1)

            ##Heading Device
            self.Heading_D = StringVar()
            self.Heading_d = Entry(self.K_op, width=10, textvariable=self.Heading_D, state='disabled')
            self.Heading_d_text = Label(self.K_op, text="Heading Device")
            self.Heading_d_text.grid(row=4, column=0, sticky=W)
            self.Heading_d.grid(row=4, column=1, sticky=W, padx=1)

            ##GPS Time Stamps
            self.GPS_T = StringVar()
            self.GPS_t = Entry(self.K_op, width=10, textvariable=self.GPS_T, state='disabled')
            self.GPS_t_text = Label(self.K_op, text="GPS Time Stamps")
            self.GPS_t_text.grid(row=5, column=0, sticky=W)
            self.GPS_t.grid(row=5, column=1, sticky=W, padx=1)

            ##Pitch Device
            self.Pitch_D = StringVar()
            self.Pitch_d = Entry(self.K_op, width=10, textvariable=self.Pitch_D, state='disabled')
            self.Pitch_d_text = Label(self.K_op, text="Pitch Device")
            self.Pitch_d_text.grid(row=6, column=0, sticky=W)
            self.Pitch_d.grid(row=6, column=1, sticky=W, padx=1)

            ##Roll Device
            self.Roll_D = StringVar()
            self.Roll_d = Entry(self.K_op, width=10, textvariable=self.Roll_D, state='disabled')
            self.Roll_d_text = Label(self.K_op, text="Roll Device")
            self.Roll_d_text.grid(row=7, column=0, sticky=W)
            self.Roll_d.grid(row=7, column=1, sticky=W, padx=1)

            ##Sound Speed Device
            self.SSP_D = StringVar()
            self.SSP_d = Entry(self.K_op, width=10, textvariable=self.SSP_D, state='disabled')
            self.SSP_d_text = Label(self.K_op, text="Sound Speed Device")
            self.SSP_d_text.grid(row=8, column=0, sticky=W)
            self.SSP_d.grid(row=8, column=1, sticky=W, padx=1)

            ##Reading the defaults or user saved inputs for Kongsberg (.all)
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            NAV_D = Parameters.iloc[3,0]
            GPSH_D = Parameters.iloc[3,1]
            Heave_D = Parameters.iloc[3,2]
            Heading_D = Parameters.iloc[3,3]
            GPS_T = Parameters.iloc[3,4]
            Pitch_D = Parameters.iloc[3,5]
            Roll_D = Parameters.iloc[3,6]
            SSP_D = Parameters.iloc[3,7]

            self.Nav_D.set(NAV_D)
            self.GPSH_D.set(GPSH_D)
            self.Heave_D.set(Heave_D)
            self.Heading_D.set(Heading_D)
            self.GPS_T.set(GPS_T)
            self.Pitch_D.set(Pitch_D)
            self.Roll_D.set(Roll_D)
            self.SSP_D.set(SSP_D)

            self.Caris_RAW_Tooltips()

            try:
                ## Forget R2 Sonic options
                self.R2_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget XTF 0ptions
                self.X_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget S7K 0ptions
                self.S7K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg .kmall 0ptions
                self.KMALL_op.grid_forget()
            except AttributeError:
                pass

        elif self.S_T.get()==2:

            self.R2_op = LabelFrame(frame2, text="Import R2 Sonic (.gsf) to HIPS", foreground="blue")
            self.R2_op.grid(row=1, column=0, padx=1, sticky=W)

            self.D_S = StringVar()
            self.D_s = Entry(self.R2_op, width=10, textvariable=self.D_S, state='disabled')
            self.D_s_text = Label(self.R2_op, text="Depth Source")
            self.D_s_text.grid(row=1, column=0, sticky=W)
            self.D_s.grid(row=1, column=1, sticky=W,  padx=1)

            self.IN_OFF = IntVar()
            self.IN_off = Checkbutton(self.R2_op, variable=self.IN_OFF, text= "Include Offline", state='disabled')
            self.IN_off.grid(row=2, column=0, sticky=W)

            self.REJ_OFF = IntVar()
            self.REJ_off = Checkbutton(self.R2_op, variable=self.REJ_OFF, text= "Reject Offline", state='disabled')
            self.REJ_off.grid(row=3, column=0, sticky=W)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            DS = Parameters.iloc[4,0]
            IN_O = Parameters.iloc[4,1]
            R_O = Parameters.iloc[4,2]

            self.D_S.set(DS)
            self.IN_OFF.set(IN_O)
            self.REJ_OFF.set(R_O)

            self.Caris_RAW_Tooltips()

            try:
                ## Forget Konsberg options
                self.K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget XTF 0ptions
                self.X_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget S7K 0ptions
                self.S7K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg .kmall 0ptions
                self.KMALL_op.grid_forget()
            except AttributeError:
                pass

        elif self.S_T.get()==3:

            self.X_op = LabelFrame(frame2, text="Import ILRIS (.xtf) to HIPS", foreground="blue")
            self.X_op.grid(row=2, column=0, padx=1, sticky=W)

            self.Nav_DX = StringVar()
            self.Nav_dX = Entry(self.X_op, width=15, textvariable=self.Nav_DX, state='disabled')
            self.Nav_dX_text = Label(self.X_op, text="Navigation Device")
            self.Nav_dX_text.grid(row=1, column=0, sticky=W)
            self.Nav_dX.grid(row=1, column=1, sticky=W, padx=1)

            self.GPSH_DX = StringVar()
            self.GPS_hX = Entry(self.X_op, width=15, textvariable=self.GPSH_DX, state='disabled')
            self.GPS_hX_text = Label(self.X_op, text="GPS Height Device")
            self.GPS_hX_text.grid(row=2, column=0, sticky=W)
            self.GPS_hX.grid(row=2, column=1, sticky=W, padx=1)

            self.M_D = StringVar()
            self.M_d = Entry(self.X_op, width=10, textvariable=self.M_D, state='disabled')
            self.M_d_text = Label(self.X_op, text="Motion Device")
            self.M_d_text.grid(row=3, column=0, sticky=W)
            self.M_d.grid(row=3, column=1, sticky=W, padx=1)

            self.C_B = StringVar()
            self.C_b = Entry(self.X_op, width=15, textvariable=self.C_B, state='disabled')
            self.C_b_text = Label(self.X_op, text="Convert Bathymetry")
            self.C_b_text.grid(row=4, column=0, sticky=W)
            self.C_b.grid(row=4, column=1, sticky=W, padx=1)

            self.Heading_DX = StringVar()
            self.Heading_dX = Entry(self.X_op, width=15, textvariable=self.Heading_DX, state='disabled')
            self.Heading_dX_text = Label(self.X_op, text="Heading Device")
            self.Heading_dX_text.grid(row=5, column=0, sticky=W)
            self.Heading_dX.grid(row=5, column=1, sticky=W, padx=1)

            self.CONV_SS = StringVar()
            self.CONV_ss = Entry(self.X_op, width=15, textvariable=self.CONV_SS, state='disabled')
            self.CONV_ss_text = Label(self.X_op, text="Convert Side Scan")
            self.CONV_ss_text.grid(row=6, column=0, stick=W)
            self.CONV_ss.grid(row=6, column=1, sticky=W, padx=1)

            self.SSWF = StringVar()
            self.sswf = Entry(self.X_op, width=15, textvariable=self.SSWF, state='disabled')
            self.sswf_text = Label(self.X_op, text="Side Scan Weighting Factor")
            self.sswf_text.grid(row=7, column=0, stick=W)
            self.sswf.grid(row=7, column=1, sticky=W, padx=1)

            self.SS_NAV = StringVar()
            self.ss_nav = Entry(self.X_op, width=15, textvariable=self.SS_NAV, state='disabled')
            self.ss_nav_text = Label(self.X_op, text="Side Scan Navigation Device")
            self.ss_nav_text.grid(row=8, column=0, stick=W)
            self.ss_nav.grid(row=8, column=1, sticky=W, padx=1)

            self.SS_HEAD = StringVar()
            self.ss_head = Entry(self.X_op, width=15, textvariable=self.SS_HEAD, state='disabled')
            self.ss_head_text = Label(self.X_op, text="Side Scan Heading Device")
            self.ss_head_text.grid(row=9, column=0, stick=W)
            self.ss_head.grid(row=9, column=1, sticky=W, padx=1)

            self.TIME_S = StringVar()
            self.time_s = Entry(self.X_op, width=15, textvariable=self.TIME_S, state='disabled')
            self.time_s_text = Label(self.X_op, text="Time Stamps")
            self.time_s_text.grid(row=10, column=0, stick=W)
            self.time_s.grid(row=10, column=1, sticky=W, padx=1)

            ##Reading the defaults or user saved inputs for KongsBerg (.all)
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            NAV_DX = Parameters.iloc[5,0]
            GPSH_DX = Parameters.iloc[5,1]
            M_D = Parameters.iloc[5,2]
            C_B = Parameters.iloc[5,3]
            Heading_DX = Parameters.iloc[5,4]
            CONV_SS = Parameters.iloc[5,5]
            SSWF = Parameters.iloc[5,6]
            SS_NAV = Parameters.iloc[5,7]
            SS_HEAD = Parameters.iloc[5,8]
            TIME_S = Parameters.iloc[5,9]

            ##Setting the defaults or user saved inputs for KongsBerg (.all)
            self.Nav_DX.set(NAV_DX)
            self.GPSH_DX.set(GPSH_DX)
            self.M_D.set(M_D)
            self.C_B.set(C_B)
            self.Heading_DX.set(Heading_DX)
            self.CONV_SS.set(CONV_SS)
            self.SSWF.set(SSWF)
            self.SS_NAV.set(SS_NAV)
            self.SS_HEAD.set(SS_HEAD)
            self.TIME_S.set(TIME_S)

            try:
                ## Forget R2 Sonic options
                self.R2_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg .all 0ptions
                self.K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget S7K 0ptions
                self.S7K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg .kmall 0ptions
                self.KMALL_op.grid_forget()
            except AttributeError:
                pass

        elif self.S_T.get()==4:
            self.S7K_op = LabelFrame(frame2, text="Import Teledyne S7K to HIPS", foreground="blue")
            self.S7K_op.grid(row=3, column=0, padx=1, sticky=W)

            self.CB = StringVar()
            self.cb = Entry(self.S7K_op, width=15, textvariable=self.CB, state='disabled')
            self.cb_text = Label(self.S7K_op, text="Convert Bathymetry")
            self.cb_text.grid(row=1, column=0, sticky=W)
            self.cb.grid(row=1, column=1, sticky=W, padx=1)

            self.NAV_D = StringVar()
            self.nav_d = Entry(self.S7K_op, width=15, textvariable=self.NAV_D, state='disabled')
            self.nav_d_text = Label(self.S7K_op, text="Navigation Device")
            self.nav_d_text.grid(row=2, column=0, sticky=W)
            self.nav_d.grid(row=2, column=1, sticky=W, padx=1)

            self.HEAD_D = StringVar()
            self.head_d = Entry(self.S7K_op, width=15, textvariable=self.HEAD_D, state='disabled')
            self.head_d_text = Label(self.S7K_op, text="Heading Device")
            self.head_d_text.grid(row=3, column=0, sticky=W)
            self.head_d.grid(row=3, column=1, sticky=W, padx=1)

            self.MOTION_D = StringVar()
            self.motion_d = Entry(self.S7K_op, width=15, textvariable=self.MOTION_D, state='disabled')
            self.motion_d_text = Label(self.S7K_op, text="Motion Device")
            self.motion_d_text.grid(row=4, column=0, sticky=W)
            self.motion_d.grid(row=4, column=1, sticky=W, padx=1)

            self.SWATH_D = StringVar()
            self.swath_d = Entry(self.S7K_op, width=15, textvariable=self.SWATH_D, state='disabled')
            self.swath_d_text = Label(self.S7K_op, text="Swath Device")
            self.swath_d_text.grid(row=5, column=0, sticky=W)
            self.swath_d.grid(row=5, column=1, sticky=W, padx=1)

            ##Reading the defaults or user saved inputs for Teledyne (.S7K)
            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            C_B = Parameters.iloc[15,0]
            N_D = Parameters.iloc[15,1]
            H_D = Parameters.iloc[15,2]
            M_D = Parameters.iloc[15,3]
            S_D = Parameters.iloc[15,4]

            ##Setting the defaults or user saved inputs for Teledyne (.S7K)
            self.CB.set(C_B)
            self.NAV_D.set(N_D)
            self.HEAD_D.set(H_D)
            self.MOTION_D.set(M_D)
            self.SWATH_D.set(S_D)

            try:
                ## Forget R2 Sonic options
                self.R2_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg 0ptions
                self.K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget XTF 0ptions
                self.X_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg .kmall 0ptions
                self.KMALL_op.grid_forget()
            except AttributeError:
                pass


        if self.S_T.get()== 5:

            ##Creating IMPORT TO HIPS Options
            self.KMALL_op = LabelFrame(frame2, text="Import Kongsberg (.kmall) to HIPS", foreground="blue")
            self.KMALL_op.grid(row=0, column=0, padx=1, sticky=W)

            ##Navigation Device
            self.Nav_D = StringVar()
            self.Nav_d = Entry(self.KMALL_op, width=10, textvariable=self.Nav_D, state='disabled')
            self.Nav_d_text = Label(self.KMALL_op, text="Navigation Device")
            self.Nav_d_text.grid(row=1, column=0, sticky=W)
            self.Nav_d.grid(row=1, column=1, sticky=W, padx=1)

            ##GPS Height Device
            self.GPSH_D = StringVar()
            self.GPS_h = Entry(self.KMALL_op, width=10, textvariable=self.GPSH_D, state='disabled')
            self.GPS_h_text = Label(self.KMALL_op, text="GPS Height Device")
            self.GPS_h_text.grid(row=2, column=0, sticky=W)
            self.GPS_h.grid(row=2, column=1, sticky=W, padx=1)

            ##Heave Device
            self.Heave_D = StringVar()
            self.Heave_d = Entry(self.KMALL_op, width=10, textvariable=self.Heave_D, state='disabled')
            self.Heave_d_text = Label(self.KMALL_op, text="Heave Device")
            self.Heave_d_text.grid(row=3, column=0, sticky=W)
            self.Heave_d.grid(row=3, column=1, sticky=W, padx=1)

            ##Heading Device
            self.Heading_D = StringVar()
            self.Heading_d = Entry(self.KMALL_op, width=10, textvariable=self.Heading_D, state='disabled')
            self.Heading_d_text = Label(self.KMALL_op, text="Heading Device")
            self.Heading_d_text.grid(row=4, column=0, sticky=W)
            self.Heading_d.grid(row=4, column=1, sticky=W, padx=1)

            ##Pitch Device
            self.Pitch_D = StringVar()
            self.Pitch_d = Entry(self.KMALL_op, width=10, textvariable=self.Pitch_D, state='disabled')
            self.Pitch_d_text = Label(self.KMALL_op, text="Pitch Device")
            self.Pitch_d_text.grid(row=6, column=0, sticky=W)
            self.Pitch_d.grid(row=6, column=1, sticky=W, padx=1)

            ##Roll Device
            self.Roll_D = StringVar()
            self.Roll_d = Entry(self.KMALL_op, width=10, textvariable=self.Roll_D, state='disabled')
            self.Roll_d_text = Label(self.KMALL_op, text="Roll Device")
            self.Roll_d_text.grid(row=7, column=0, sticky=W)
            self.Roll_d.grid(row=7, column=1, sticky=W, padx=1)

            ##Delayed Heave Device
            self.DelHeave_D = StringVar()
            self.DelHeave_d = Entry(self.KMALL_op, width=10, textvariable=self.Heave_D, state='disabled')
            self.DelHeave_d_text = Label(self.KMALL_op, text="Delayed Heave Device")
            self.DelHeave_d_text.grid(row=8, column=0, sticky=W)
            self.DelHeave_d.grid(row=8, column=1, sticky=W, padx=1)

            ##GPS Time Stamps
            self.GPS_T = StringVar()
            self.GPS_t = Entry(self.KMALL_op, width=10, textvariable=self.GPS_T, state='disabled')
            self.GPS_t_text = Label(self.KMALL_op, text="GPS Time Stamps")
            self.GPS_t_text.grid(row=9, column=0, sticky=W)
            self.GPS_t.grid(row=9, column=1, sticky=W, padx=1)

            

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            NAV_D = Parameters.iloc[16,0]
            GPSH_D = Parameters.iloc[16,1]
            Heave_D = Parameters.iloc[16,2]
            Heading_D = Parameters.iloc[16,3]
            Pitch_D = Parameters.iloc[16,4]
            Roll_D = Parameters.iloc[16,5]
            DelHeave_D = Parameters.iloc[16,6]
            GPS_T = Parameters.iloc[16,7]

            self.Nav_D.set(NAV_D)
            self.GPSH_D.set(GPSH_D)
            self.Heave_D.set(Heave_D)
            self.Heading_D.set(Heading_D)
            self.Pitch_D.set(Pitch_D)
            self.Roll_D.set(Roll_D)
            self.DelHeave_D.set(DelHeave_D)
            self.GPS_T.set(GPS_T)

            try:
                ## Forget R2 Sonic options
                self.R2_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget Kongsberg 0ptions
                self.K_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget XTF 0ptions
                self.X_op.grid_forget()
            except AttributeError:
                pass
            try:
                ## Forget S7K 0ptions
                self.S7K_op.grid_forget()
            except AttributeError:
                pass


    def IMPORT_TO_HIPS(self):


        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        HDCS_Folder = self.HDCS_D.get()
        RAW_F = self.RAW_F.get()
        L_R = listdir(RAW_F)
        crs = self.CRS_O.get()
        CRS = crs.partition(": ")[2]
        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()

        if self.IntialRun.get()==1:

            with open("CreateHIPSFile.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Creating HIPS File' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run CreateHIPSFile --output-crs ' +
                             CRS +
                             r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips' +
                             '> ' + Out + '/' + 'CreateHIPSFile.txt')

            p = S.Popen(['CreateHIPSFile.bat'])
            p.communicate()


        if self.S_T.get()==1:
            nav_d = self.Nav_D.get()
            gpsh_d = self.GPSH_D.get()
            heading_d = self.Heading_D.get()
            gps_t = self.GPS_T.get()
            pitch_d = self.Pitch_D.get()
            roll_d = self.Roll_D.get()
            ssp_d = self.SSP_D.get()
            H_Format = ('KONGSBERG')
            with open("Import_To_Hips.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing RAW KongsBerg .all' + '\n')
                Import.write('cd '+ Caris + '\n')
                raw_f = []
                for file in L_R:
                    if file.endswith(".all"):
                        ## Rename ALL files to not have spaces
                        r = file.replace(" ", "")
                        if(r != file):
                            File = (RAW_F + '/' + file)
                            Rename = (RAW_F + '/' + r)
                            rename(File, Rename)
                            raw_f.append(Rename)
                        else:
                           raw_f.append(RAW_F + '/' + file)

                Import.write('carisbatch --run ImportToHIPS --input-format ' +
                             H_Format + ' --convert-navigation ' + '--input-crs ' + CRS +
                             ' --vessel-file ' + Vessel_F +
                             ' --navigation-device ' + nav_d + ' --gps-height-device ' + gpsh_d +
                             ' --heading-device ' + heading_d +  ' --gps-timestamps ' + gps_t +
                             ' --pitch-device ' + pitch_d + ' --roll-device ' + roll_d +
                             ' --ssp-device ' + ssp_d + ' ')
                for file in raw_f:
                    Import.write(file + ' ')
                Import.write(r'file:///' + HDCS_Folder + '/' +  HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                             ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                             '1.Import_To_Hips_ALL_' + JD + '_' + Year + '.txt' + '\n')

        elif self.S_T.get()==5:
            nav_d = self.Nav_D.get()
            gpsh_d = self.GPSH_D.get()
            heave_d = self.Heave_D.get()
            heading_d = self.Heading_D.get()
            pitch_d = self.Pitch_D.get()
            roll_d = self.Roll_D.get()
            delheave_d = self.DelHeave_D.get()
            gps_t = self.GPS_T.get()
            H_Format = ('KONGSBERGKMALL')
            with open("Import_To_Hips.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing RAW KongsBerg .kmall' + '\n')
                Import.write('cd '+ Caris + '\n')
                raw_f = []
                for file in L_R:
                    if file.endswith(".kmall"):
                        ## Rename KMALL files to not have spaces
                        r = file.replace(" ", "")
                        if(r != file):
                            File = (RAW_F + '/' + file)
                            Rename = (RAW_F + '/' + r)
                            rename(File, Rename)
                            raw_f.append(Rename)
                        else:
                           raw_f.append(RAW_F + '/' + file)

                Import.write('carisbatch --run ImportToHIPS --input-format ' +
                             H_Format + ' --convert-navigation ' + '--input-crs ' + CRS +
                             ' --vessel-file ' + Vessel_F +
                             ' --navigation-device ' + nav_d + ' --gps-height-device ' + gpsh_d +
                             ' --heading-device ' + heading_d +  ' --heave-device ' + heave_d +
                             ' --pitch-device ' + pitch_d + ' --roll-device ' + roll_d +
                             ' --delayed-heave-device ' + delheave_d +
                             ' --gps-timestamps ' + gps_t + ' ')
                for file in raw_f:
                    Import.write(file + ' ')
                Import.write(r'file:///' + HDCS_Folder + '/' +  HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                             ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                             '1.Import_To_Hips_ALL_' + JD + '_' + Year + '.txt' + '\n')

        elif self.S_T.get()==2:

            depth_s = self.D_S.get()
            H_Format = ('GSF')
            with open("Import_To_Hips.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing RAW R2 Sonic .gsf' + '\n')
                Import.write('cd '+ Caris + '\n')
                raw_f = []
                for file in L_R:
                    if file.endswith(".gsf"):
                        ## Rename GSf files to not have spaces
                        r = file.replace(" ", "")
                        if(r != file):
                            File = (RAW_F + '/' + file)
                            Rename = (RAW_F + '/' + r)
                            rename(File, Rename)
                            raw_f.append(Rename)
                        else:
                           raw_f.append(RAW_F + '/' + file)

                Import.write('carisbatch --run ImportToHIPS --input-format ' +
                             H_Format + ' --input-crs ' + CRS +
                             ' --vessel-file ' + Vessel_F +
                             ' --depth-source ' + depth_s + ' --include-offline ' + ' ')
                for file in raw_f:
                    Import.write(file + ' ')
                Import.write(r'file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                             ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                             '1.Import_To_Hips_GSF_' + JD + '_' + Year + '.txt' + '\n')

        elif self.S_T.get()==3:

            nav_dx = self.Nav_DX.get()
            gpsh_dx = self.GPSH_DX.get()
            m_d = self.M_D.get()
            c_b = self.C_B.get()
            heading_dx = self.Heading_DX.get()
            conv_ss = self.CONV_SS.get()
            sswf = self.SSWF.get()
            ss_nav = self.SS_NAV.get()
            ss_head = self.SS_HEAD.get()
            time_s = self.TIME_S.get()
            H_Format = ('XTF')
            with open("Import_To_Hips.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing RAW Triton .xtf' + '\n')
                Import.write('cd '+ Caris + '\n')
                raw_f = []
                for file in L_R:
                    if file.endswith(".xtf"):
                        ##Rename XTF files to not have spaces
                        r = file.replace(" ", "")
                        if(r != file):
                            File = (RAW_F + '/' + file)
                            Rename = (RAW_F + '/' + r)
                            rename(File, Rename)
                            raw_f.append(Rename)
                        else:
                           raw_f.append(RAW_F + '/' + file)

                Import.write('carisbatch --run ImportToHIPS --input-format ' +
                             H_Format + ' --input-crs ' + CRS +
                             ' --vessel-file ' + Vessel_F +
                             ' --navigation-device ' + nav_dx + ' --gps-height-device ' + gpsh_dx +
                             ' --heading-device ' + heading_dx +
                             ' --motion-device ' + m_d + ' --convert-bathymetry ' + c_b +
                             ' --convert-side-scan ' + conv_ss + ' --ss-weighting-factor ' + sswf +
                             ' --ss-navigation-device ' + ss_nav + ' --ss-heading-device ' + ss_head +
                             ' --timestamps ' + time_s + ' ')
                for file in raw_f:
                    Import.write(file + ' ')
                Import.write(r'file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                            '1.Import_To_Hips_XTF_' + JD + '_' + Year + '.txt' + '\n' 'Pause')

        elif self.S_T.get()==4:

            CB = self.CB.get()
            ND = self.NAV_D.get()
            HD = self.HEAD_D.get()
            MD = self.MOTION_D.get()
            SD = self.SWATH_D.get()
            H_Format = ('TELEDYNE_7K')
            with open("Import_To_Hips.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing RAW Teledyne .S7K' + '\n')
                Import.write('cd '+ Caris + '\n')
                raw_f = []
                for file in L_R:
                    if file.endswith(".s7k"):
                        ## Rename S7K files to not have spaces
                        r = file.replace(" ", "")
                        if(r != file):
                            File = (RAW_F + '/' + file)
                            Rename = (RAW_F + '/' + r)
                            rename(File, Rename)
                            raw_f.append(Rename)
                        else:
                           raw_f.append(RAW_F + '/' + file)
                Import.write('carisbatch --run ImportToHIPS --input-format ' +
                             H_Format + ' --input-crs ' + CRS +
                             ' --convert-bathymetry ' + CB + ' --navigation-device ' +
                             ND + ' --heading-device ' + HD + ' --motion-device ' + MD +
                             ' --swath-device ' + SD + ' ')
                for file in raw_f:
                    Import.write(file + ' ')
                Import.write(r'file:///' + HDCS_Folder + '/' +  HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                             ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                             '1.Import_To_Hips_S7K_' + JD + '_' + Year + '.txt' + '\n')

        p = S.Popen(['Import_To_Hips.bat'])
        p.communicate()


    def Load_Tide_Par(self):


        if self.T_T.get()==1:
            self.GPSTide_op = LabelFrame(frame4, text="Import GPS Tides", foreground="blue")
            self.GPSTide_op.grid(row=0, column=0, sticky=W)

            self.C_GPS_ADJ = IntVar()
            self.C_GPS_Adj = Checkbutton(self.GPSTide_op, variable=self.C_GPS_ADJ, text= "Compute GPS \n Vertical Adjustment \n Errors"
                                         , state='disabled')
            self.C_GPS_Adj.grid(row=7, column=0, sticky=W)

            self.M_F = StringVar()
            self.M_f = Entry(self.GPSTide_op, width=25, textvariable=self.M_F)
            self.M_f_text = Label(self.GPSTide_op, text="Model File")
            self.M_f_text.grid(row=0, column=0, sticky=W)
            self.M_f.grid(row=0, column=1, sticky=W)
            self.ButtonMF = Button(self.GPSTide_op, text="...", height=0,
                              command=self.Search_Model_File)
            self.ButtonMF.grid(row=0, column=2, sticky=W, padx=2)

            self.INFO_F = StringVar()
            self.Info_f = Entry(self.GPSTide_op, width=25, textvariable=self.INFO_F)
            self.Info_f_text = Label(self.GPSTide_op, text="Info File")
            self.Info_f_text.grid(row=1, column=0, sticky=W)
            self.Info_f.grid(row=1, column=1, sticky=W)
            self.ButtonMF1 = Button(self.GPSTide_op, text="...", height=0,
                              command=self.Search_Info_File)
            self.ButtonMF1.grid(row=1, column=2, sticky=W, padx=2)

            self.INFO_CRS = StringVar()
            infocrs_op = ['EPSG:4617@2010', 'EPSG:7912@2010']
            self.CRS_inf = ttk.Combobox(self.GPSTide_op, values=infocrs_op, width=15, textvariable=self.INFO_CRS)
            self.CRS_inf_text = Label(self.GPSTide_op, text="Model CRS")
            self.CRS_inf_text.grid(row=3, column=0, sticky=W)
            self.CRS_inf.grid(row=3, column=1, sticky=W)

            self.SD_OFF = StringVar()
            self.SD_Off = Entry(self.GPSTide_op, width=15, textvariable=self.SD_OFF, state='disabled')
            self.SD_Off_text = Label(self.GPSTide_op, text="Sounding Datum \n Offset")
            self.SD_Off_text.grid(row=4, column=0, sticky=W)
            self.SD_Off.grid(row=4, column=1, sticky=W)


            self.W_L = StringVar()
            WL_op = ['VESSEL',
                  'REALTIME',
                  'NONE']

            self.W_l = ttk.Combobox(self.GPSTide_op, values=WL_op, width=15, textvariable=self.W_L)
            self.W_l_text = Label(self.GPSTide_op, text="Waterline")
            self.W_l_text.grid(row=5, column=0, sticky=W)
            self.W_l.grid(row=5, column=1, sticky=W+E, padx=0)


            #self.W_l = Entry(self.GPSTide_op, width=15, textvariable=self.W_L, state='disabled')

            self.H_MERGED = StringVar()
            self.H_Merged = Entry(self.GPSTide_op, width=15, textvariable=self.H_MERGED, state='disabled')
            self.H_Merged_text = Label(self.GPSTide_op, text="Heave Type")
            self.H_Merged_text.grid(row=6, column=0, sticky=W)
            self.H_Merged.grid(row=6, column=1, sticky=W, padx=1)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            CGPSVA = Parameters.iloc[7,0]
            SDO = Parameters.iloc[7,1]
            MF = Parameters.iloc[7,2]
            INFO = Parameters.iloc[7,3]
            CRS_INFO = Parameters.iloc[7,4]
            WL = Parameters.iloc[7,5]
            HEAVE_M = Parameters.iloc[8,0]

            self.M_F.set(MF)
            self.INFO_F.set(INFO)
            self.INFO_CRS.set(CRS_INFO)
            self.SD_OFF.set(SDO)
            self.C_GPS_ADJ.set(CGPSVA)
            self.W_L.set(WL)
            self.H_MERGED.set(HEAVE_M)

            ##ToolTips for GPS Tides
            tip_MODEL = ToolTip(self.M_f, (self.M_F.get()))
            tip_INFO = ToolTip(self.Info_f, (self.INFO_F.get()))

            try:
                ## Forget the Observed/Predicted Tides
                self.OPTide_op.grid_forget()
            except AttributeError:
                pass

        elif self.T_T.get()==2:
            self.OPTide_op = LabelFrame(frame4, text="Import Observed/Pedicted Tides", foreground="blue")
            self.OPTide_op.grid(row=1, column=0, sticky=W)

            self.T_F = StringVar()
            self.T_f = Entry(self.OPTide_op, width=15, textvariable=self.T_F)
            self.T_f_text = Label(self.OPTide_op, text="Tide File")
            self.T_f_text.grid(row=0, column=0, sticky=W)
            self.T_f.grid(row=0, column=1, sticky=W)
            self.ButtonTF = Button(self.OPTide_op, text="...", height=0,
                              command=self.Search_TIDE_File)
            self.ButtonTF.grid(row=0, column=2, sticky=W, padx=2)

            self.W_Ave = IntVar()
            self.W_ave = Checkbutton(self.OPTide_op, variable=self.W_Ave, text= "Weighted Average", state='disabled')
            self.W_ave.grid(row=1, column=0, sticky=W)

            self.COMP_Errors = IntVar()
            self.COMP_errors = Checkbutton(self.OPTide_op, variable=self.COMP_Errors, text= "Compute Errors", state='disabled')
            self.COMP_errors.grid(row=2, column=0, sticky=W)

            self.H_MERGED = StringVar()
            self.H_Merged = Entry(self.OPTide_op, width=15, textvariable=self.H_MERGED, state='disabled')
            self.H_Merged_text = Label(self.OPTide_op, text="Heave Type")
            self.H_Merged_text.grid(row=3, column=0, sticky=W)
            self.H_Merged.grid(row=3, column=1, sticky=W, padx=1)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            TF = Parameters.iloc[6,0]
            WAVE = Parameters.iloc[6,1]
            COMPE = Parameters.iloc[6,2]
            HEAVE_M = Parameters.iloc[8,0]

            self.T_F.set(TF)
            self.W_Ave.set(WAVE)
            self.COMP_Errors.set(COMPE)
            self.H_MERGED.set(HEAVE_M)

            ## ToolTips for Observed/ Predicted Tides
            tip_TIDEFILE = ToolTip(self.T_f, (self.T_F.get()))

            try:
                ## Forget the GPS Tide Options
                self.GPSTide_op.grid_forget()
            except AttributeError:
                pass

    def Loads_MergeTrack(self):

        
        if self.MERGE_TRACK.get()==1:
             
            self.MERGE_O = LabelFrame(frame8, text="Merge Tracklines", foreground="blue")
            self.MERGE_O.grid(row=1, column=0, sticky=N+W)

            self.H_MERGED = StringVar()
            self.H_Merged = Entry(self.MERGE_O, width=20, textvariable=self.H_MERGED, state='disabled')
            self.H_Merged_text = Label(self.MERGE_O, text="Heave Type")
            self.H_Merged_text.grid(row=0, column=0, sticky=W)
            self.H_Merged.grid(row=0, column=1, sticky=W, padx=1)

            self.VERT_REF = StringVar()
            vert_ref = ['NONE',
                        'GPS',
                        'TIDE']

            self.VREF_op = ttk.Combobox(self.MERGE_O, values=vert_ref, width=7, textvariable=self.VERT_REF)
            self.VREF_text = Label(self.MERGE_O, text="Choose Vertical Reference")
            self.VREF_text.grid(row=1, column=0, sticky=W)
            self.VREF_op.grid(row=1, column=1, sticky=W+E, padx=0)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            HEAVE_M = Parameters.iloc[8,0]
            VERT_R = Parameters.iloc[8,1]

            self.H_MERGED.set(HEAVE_M)
            self.VERT_REF.set(VERT_R)

        if self.MERGE_TRACK.get()==0:
            try:
                ## Forget Merge Options
                self.MERGE_O.grid_forget()
            except AttributeError:
                pass


    def Load_TPU_Par(self):

        self.TPU_msg = LabelFrame(frame5, text="TPU User Message", foreground="blue")
        self.TPU_msg.grid(row=0, column=1, padx=1, sticky=N+W)

        msg = 'Situation\nPOS - VESSEL\nSBET - REALTIME'

        self.User_Msg_TPU = Text(self.TPU_msg, width=15, height=3)
        self.User_Msg_TPU.insert(END, msg)
        self.User_Msg_TPU.config(state='disabled')
        self.User_Msg_TPU.grid(row=0, column=0, padx=1, sticky=W)

        if self.COMP_TPU.get()==1:

            self.Comptpu_op = LabelFrame(frame5, text="Compute TPU", foreground="blue")
            self.Comptpu_op.grid(row=1, column=0, sticky=N+W)

            TPU_op = ['VESSEL',
                  'REALTIME']

            self.TIDE_M = StringVar()
            self.TIDE_m = Entry(self.Comptpu_op, width=10, textvariable=self.TIDE_M, state='normal')
            self.TIDE_m_text = Label(self.Comptpu_op, text="Measured Tide")
            self.TIDE_m_text.grid(row=0, column=0, sticky=W)
            self.TIDE_m.grid(row=0, column=1, sticky=W, padx=1)

            self.SV_M = StringVar()
            self.SV_m = Entry(self.Comptpu_op, width=10, textvariable=self.SV_M, state='normal')
            self.SV_m_text = Label(self.Comptpu_op, text="Measured Sound Velocity")
            self.SV_m_text.grid(row=1, column=0, sticky=W)
            self.SV_m.grid(row=1, column=1, sticky=W, padx=1)

            self.SS_V = StringVar()
            self.SS_v = Entry(self.Comptpu_op, width=10, textvariable=self.SS_V, state='normal')
            self.SS_v_text = Label(self.Comptpu_op, text="Surface Sound Velocity")
            self.SS_v_text.grid(row=3, column=0, sticky=W)
            self.SS_v.grid(row=3, column=1, sticky=W, padx=1)

            self.S_N = StringVar()
            self.S_n = ttk.Combobox(self.Comptpu_op, values=TPU_op, width=15, textvariable=self.S_N)
            self.S_n_text = Label(self.Comptpu_op, text="Navigation Source")
            self.S_n_text.grid(row=4, column=0, sticky=W)
            self.S_n.grid(row=4, column=1, sticky=W, padx=1)

            self.S_G = StringVar()
            self.S_g = ttk.Combobox(self.Comptpu_op, values=TPU_op, width=15, textvariable=self.S_G)
            self.S_g_text = Label(self.Comptpu_op, text="GYRO Source")
            self.S_g_text.grid(row=5, column=0, sticky=W)
            self.S_g.grid(row=5, column=1, sticky=W, padx=1)

            self.S_S = StringVar()
            self.S_s = ttk.Combobox(self.Comptpu_op, values=TPU_op, width=15, textvariable=self.S_S)
            self.S_s_text = Label(self.Comptpu_op, text="Sonar Source")
            self.S_s_text.grid(row=6, column=0, sticky=W)
            self.S_s.grid(row=6, column=1, sticky=W, padx=1)

            self.S_P = StringVar()
            self.S_p = ttk.Combobox(self.Comptpu_op, values=TPU_op, width=15, textvariable=self.S_P)
            self.S_p_text = Label(self.Comptpu_op, text="Pitch Source")
            self.S_p_text.grid(row=7, column=0, sticky=W)
            self.S_p.grid(row=7, column=1, sticky=W, padx=1)

            self.S_R = StringVar()
            self.S_r = ttk.Combobox(self.Comptpu_op, values=TPU_op, width=15, textvariable=self.S_R)
            self.S_r_text = Label(self.Comptpu_op, text="Roll Source")
            self.S_r_text.grid(row=8, column=0, sticky=W)
            self.S_r.grid(row=8, column=1, sticky=W, padx=1)

            self.S_H = StringVar()
            self.S_h = ttk.Combobox(self.Comptpu_op, values=TPU_op, width=15, textvariable=self.S_H)
            self.S_h_text = Label(self.Comptpu_op, text="Heave Source")
            self.S_h_text.grid(row=9, column=0, sticky=W)
            self.S_h.grid(row=9, column=1, sticky=W, padx=1)

            self.S_Tide = StringVar()
            self.S_tide = Entry(self.Comptpu_op, width=10, textvariable=self.S_Tide, state='disabled')
            self.S_tide_text = Label(self.Comptpu_op, text="Tide Source")
            self.S_tide_text.grid(row=10, column=0, sticky=W)
            self.S_tide.grid(row=10, column=1, sticky=W, padx=1)

            self.Merge_op = LabelFrame(frame5, text="Merge", foreground="blue")
            self.Merge_op.grid(row=0, column=0, sticky=N+W)

            self.H_MERGED = StringVar()
            self.H_Merged = Entry(self.Merge_op, width=20, textvariable=self.H_MERGED, state='disabled')
            self.H_Merged_text = Label(self.Merge_op, text="Heave Type")
            self.H_Merged_text.grid(row=0, column=0, sticky=W)
            self.H_Merged.grid(row=0, column=1, sticky=W, padx=1)

            self.VERT_REF = StringVar()
            vert_ref = ['NONE',
                        'GPS',
                        'TIDE']

            self.VREF_op = ttk.Combobox(self.Merge_op, values=vert_ref, width=7, textvariable=self.VERT_REF)
            self.VREF_text = Label(self.Merge_op, text="Choose Vertical Reference")
            self.VREF_text.grid(row=1, column=0, sticky=W)
            self.VREF_op.grid(row=1, column=1, sticky=W+E, padx=0)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            MEAS_TIDE = Parameters.iloc[9,0]
            MEAS_SP = Parameters.iloc[9,1]
            SSP = Parameters.iloc[9,2]
            NAV_S = Parameters.iloc[9,3]
            SONAR_S = Parameters.iloc[9,4]
            GYRO_S = Parameters.iloc[9,5]
            PITCH_S = Parameters.iloc[9,6]
            ROLL_S = Parameters.iloc[9,7]
            HEAVE_S = Parameters.iloc[9,8]
            TIDE_S = Parameters.iloc[9,9]
            HEAVE_M = Parameters.iloc[8,0]
            VERT_R = Parameters.iloc[8,1]

            self.TIDE_M.set(MEAS_TIDE)
            self.SV_M.set(MEAS_SP)
            self.SS_V.set(SSP)
            self.S_N.set(NAV_S)
            self.S_S.set(SONAR_S)
            self.S_G.set(GYRO_S)
            self.S_P.set(PITCH_S)
            self.S_R.set(ROLL_S)
            self.S_H.set(HEAVE_S)
            self.S_Tide.set(TIDE_S)
            self.H_MERGED.set(HEAVE_M)
            self.VERT_REF.set(VERT_R)

        if self.COMP_TPU.get()==0:
            try:
                ## Forget TPU Options
                self.Comptpu_op.grid_forget()
                self.Merge_op.grid_forget()
            except AttributeError:
                pass


    def GEOREFERENCE_HIPS(self):

        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]
        HDCS_Folder = self.HDCS_D.get()

        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()

        if self.T_T.get()==1:
            Model_File = self.M_F.get()
            Info_File = self.INFO_F.get()
            Info_CRS = self.INFO_CRS.get()
            self.SD_OFF.get()
            self.C_GPS_ADJ.get()
            Water_Line = self.W_L.get()

            with open("Import_Tides.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing GPS Tides' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run GeoreferenceHIPSBathymetry --vertical-datum-reference GPS' +
                             ' --datum-model-file ' + Model_File)
                if (Model_File.endswith('.txt') or Model_File.endswith('.csv')
                    or Model_File.endswith('.xyz')):
                    Import.write(' --info-file ' + Info_File)
                    Import.write(' --input-crs ' + Info_CRS)
                else:
                    Import.write(' --datum-model-band DEPTH')
                if self.C_GPS_ADJ.get()==1:
                    Import.write(' --compute-gps-vertical-adjustment --GPS-Vertical-Components CUSTOM --GPS-Component-Waterline ' + Water_Line)
                    Import.write(' --GPS-Component-Dynamic-Heave DELAYED_HEAVE --heave-source DELAYED_HEAVE' +
                                 ' --output-components')
                Import.write(r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                            '3.Compute_GPSTIDE_' + JD + '_' + Year + '.txt' + '\n')

            p = S.check_call("Import_Tides.bat", stdin=None, stdout=None, stderr=None, shell=False)

        elif self.T_T.get()==2:
            Tide_File = self.T_F.get()
            W_Ave = self.W_Ave.get()
            COMP_Errors = self.COMP_Errors.get()

            with open("Import_Tides.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Importing Observed/Predicted Tides' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run GeoreferenceHIPSBathymetry --vertical-datum-reference TIDE'+
                             ' --tide-file ' + Tide_File)
                if W_Ave == 1:
                    Import.write(' --weighted-average ')
                if COMP_Errors == 1:
                    Import.write(' --compute-errors ')
                Import.write(' --heave-source DELAYED_HEAVE')
                Import.write(r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                            '3.Import_OBSERVEDTIDE_' + JD + '_' + Year + '.txt' + '\n')

            p = S.check_call("Import_Tides.bat", stdin=None, stdout=None, stderr=None, shell=False)

        if self.COMP_TPU.get()==1:
            Meas_Tide = self.TIDE_M.get()
            Meas_SV = self.SV_M.get()
            Surf_So = self.SS_V.get()
            Nav_Source = self.S_N.get()
            Gyro_Source = self.S_G.get()
            Sonar_Source = self.S_S.get()
            Pitch_Source = self.S_P.get()
            Roll_Source = self.S_R.get()
            Heave_Source = self.S_H.get()
            Tide_Source = self.S_Tide.get()
            Heave_M = self.H_MERGED.get()
            Vert_Ref = self.VERT_REF.get()

            with open("Compute_TPU.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Computing TPU' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run GeoreferenceHIPSBathymetry --vertical-datum-reference ' + str(Vert_Ref) +
                             ' --heave-source ' + str(Heave_M) + ' --compute-tpu ' +
                             '--tide-measured ' + str(Meas_Tide) + ' --sv-measured ' + str(Meas_SV) +
                             ' --sv-surface ' + str(Surf_So) + ' --source-sonar ' + Sonar_Source +
                             ' --source-navigation ' + Nav_Source +
                             ' --source-gyro ' + Gyro_Source +
                             ' --source-pitch ' + Pitch_Source +
                             ' --source-roll ' + Roll_Source + ' --source-heave ' + Heave_Source +
                             ' --source-tide ' + Tide_Source)
                Import.write(r' file:///' + HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + str(JD) + '/' +
                            '4.Compute_TPU_' + JD + '_' + Year + '.txt' + '\n')

            p = S.check_call("Compute_TPU.bat", stdin=None, stdout=None, stderr=None, shell=False)

        if self.APPLY_SVP.get()==1:
            SVP_F = self.SVP_F.get()
            PROFILE = self.PROFILE.get()
            ND_HOUR = self.ND_HOUR.get()
            Heave_M2 = self.H_MERGED2.get()
            Vert_Ref2 = self.VERT_REF2.get()

            with open("Import_SVP.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Applying SVP' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run GeoreferenceHIPSBathymetry --vertical-datum-reference ' + str(Vert_Ref2) +
                             ' --svp ' + SVP_F + ' --heave-source ' + Heave_M2 +
                             ' --compute-svc --profile-selection-method ' + PROFILE +
                             ' --nearest-distance-hours ' + ND_HOUR + ' --ssp')
                Import.write(r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                            '5.Import_SVP_' + JD + '_' + Year + '.txt' + '\n')

            p = S.check_call("Import_SVP.bat", stdin=None, stdout=None, stderr=None, shell=False)

        if self.MERGE_TRACK.get()==1:
            Heave_M = self.H_MERGED.get()
            Vert_Ref = self.VERT_REF.get()
            with open("Merge_Tracklines.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('@ECHO Merging Tracklines' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run GeoreferenceHIPSBathymetry --vertical-datum-reference ' + str(Vert_Ref) +
                             ' --heave-source ' + Heave_M)
                Import.write(r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' > ' + Out + '/' + JD + '/' +
                            '7.Merge_Tracklines_' + JD + '_' + Year + '.txt' + '\n')

            p = S.check_call("Merge_Tracklines.bat", stdin=None, stdout=None, stderr=None, shell=False)
            


    def Load_GRID_Par(self):

        chdir(owd)

        if self.GRID.get()==1:
            self.CREATEGRID_op = LabelFrame(frame7, text="Create HIPS Grid", foreground="blue")
            self.CREATEGRID_op.grid(row=0, column=0, sticky=W)

            self.RES = StringVar()
            self.res = Entry(self.CREATEGRID_op, width=5, textvariable=self.RES)
            self.res_text = Label(self.CREATEGRID_op, text="Surface Resolution")
            self.res_text.grid(row=0, column=0, sticky=W)
            self.res.grid(row=0, column=1, sticky=W)

            self.GRID_DIR = StringVar()
            self.GRID_dir = Entry(self.CREATEGRID_op, width=32, textvariable=self.GRID_DIR)
            self.GRID_dir_text = Label(self.CREATEGRID_op, text="Surface Dir")
            self.GRID_dir_text.grid(row=1, column=0, sticky=W)
            self.GRID_dir.grid(row=1, column=1, sticky=W)
            self.ButtonG_dir = Button(self.CREATEGRID_op, text="...", height=0,
                                   command=self.Search_Grid_Dir)
            self.ButtonG_dir.grid(row=1, column=2, sticky=W, padx=2)

            self.IHO_ORDER = StringVar()
            iho_op = ['S44_SPECIAL',
                      'S44_1A',
                      'S44_1B',
                      'S44_2']
            self.IHO_op = ttk.Combobox(self.CREATEGRID_op, values=iho_op, width=10, textvariable=self.IHO_ORDER)
            self.IHO_text = Label(self.CREATEGRID_op, text="Choose IHO Order")
            self.IHO_text.grid(row=2, column=0, sticky=W)
            self.IHO_op.grid(row=2, column=1, sticky=W+E, padx=0)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            RES = Parameters.iloc[10,0]
            IHOORDER = Parameters.iloc[10,1]
            GRIDDIR = Parameters.iloc[10,2]

            self.RES.set(RES)
            self.IHO_ORDER.set(IHOORDER)
            self.GRID_DIR.set(GRIDDIR)

            try:
                ## Forget Create Grid 2
                self.CREATEGRID2_op.grid_forget()
            except AttributeError:
                pass

        if self.GRID.get()==2:

            self.CREATEGRID2_op = LabelFrame(frame7, text="Create/Add to HIPS Grid", foreground="blue")
            self.CREATEGRID2_op.grid(row=1, column=0, padx=1, sticky=W)

            self.RES = StringVar()
            self.res = Entry(self.CREATEGRID2_op, width=5, textvariable=self.RES)
            self.res_text = Label(self.CREATEGRID2_op, text="Surface Resolution")
            self.res_text.grid(row=0, column=0, sticky=W)
            self.res.grid(row=0, column=1, sticky=W)

            self.GRID_DIR = StringVar()
            self.GRID_dir = Entry(self.CREATEGRID2_op, width=32, textvariable=self.GRID_DIR)
            self.GRID_dir_text = Label(self.CREATEGRID2_op, text="Surface Dir")
            self.GRID_dir_text.grid(row=1, column=0, sticky=W)
            self.GRID_dir.grid(row=1, column=1, sticky=W)
            self.ButtonG_dir = Button(self.CREATEGRID2_op, text="...", height=0,
                                   command=self.Search_Grid_Dir)
            self.ButtonG_dir.grid(row=1, column=2, sticky=W, padx=2)

            self.IHO_ORDER = StringVar()
            iho_op = ['S44_SPECIAL',
                      'S44_1A',
                      'S44_1B',
                      'S44_2']
            self.IHO_op = ttk.Combobox(self.CREATEGRID2_op, values=iho_op, width=10, textvariable=self.IHO_ORDER)
            self.IHO_text = Label(self.CREATEGRID2_op, text="Choose IHO Order")
            self.IHO_text.grid(row=3, column=0, sticky=W)
            self.IHO_op.grid(row=3, column=1, sticky=W+E, padx=0)


            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            RES = Parameters.iloc[10,0]
            IHOORDER = Parameters.iloc[10,1]
            GRIDDIR = Parameters.iloc[10,2]
            COVERAGE = Parameters.iloc[10,3]

            self.RES.set(RES)
            self.IHO_ORDER.set(IHOORDER)
            self.GRID_DIR.set(GRIDDIR)

            try:
                ##Forget Create Grid
                self.CREATEGRID_op.grid_forget()
            except AttributeError:
                pass

        try:
            ##Tool Tip for GRID Directory
            tip_GRID = ToolTip(self.GRID_dir, str(self.GRID_DIR.get()))
        except AttributeError:
            pass


    def Create_Addto_Hips_Grid(self):

        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        Res = self.RES.get()
        Res_no = Res.split('m')
        Res_List = ['2m', '5m', '10m', '20m', '30m', '50m', Res]
        Res_no = float(Res_no[0])
        if Res_no <= 1:
            Res_P = Res_List[0]
        elif Res_no > 1 and Res_no <= 5:
            Res_P = Res_List[1]
        elif Res_no > 5 and Res_no <= 10:
            Res_P = Res_List[2]
        elif Res_no > 10 and Res_no <= 20:
            Res_P = Res_List[3]
        elif Res_no > 20 and Res_no <= 30:
            Res_P = Res_List[4]
        elif ResRes_no[0] > 30 and ResRes_no[0] <= 50:
            Res_P = Res_List[5]
        else:
            Res_P = Res_List[6]

        IHO = self.IHO_ORDER.get()
        Dir_Grid = self.GRID_DIR.get()
        HDCS_Folder = self.HDCS_D.get()
        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()
        crs = self.CRS_O.get()
        CRS = crs.partition(": ")[2]
        if CRS=='EPSG:7912@2010':
            CRS2 = 'EPSG:5937'
        else:
            CRS2 = CRS
        VCRS = 'CUSTOM:69036444' ## CHS PACD Vertical Reference
        G_M = ('CUBE')
        E_R = ('GEOTIFF')
        


        if self.GRID.get()==1:

            with open("Create_Grid.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('@ECHO Creating Grids' + '\n')

                ##Create Surface to Add Future Days to
                Import.write('carisbatch --run CreateHIPSGrid --gridding-method ' + G_M +
                             ' --resolution ' + str(Res_P) + ' --Output-crs ' + CRS2 +
                             ' --output-vertical-crs ' + VCRS +
                             ' --compute-band SHOAL --compute-band DEEP --compute-band DENSITY' +
                             ' --compute-band MEAN --compute-band STD_DEV ' +
                             ' --include-flag ACCEPTED --include-flag EXAMINED --include-flag OUTSTANDING ' +
                             ' --keep-up-to-date ' + ' --iho-order ' + IHO +
                             ' --disambiguation-method DENSITY_LOCALE')
                Import.write(r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' ' +
                             Dir_Grid + '/' + Project_N + '.csar' +
                             ' > ' + Out + '/' + JD + '/' + '6.Create_JD_Grid_' + JD + '_' + Year + '.txt' + '\n')
                Import.write('carisbatch --run ExportRaster --output-format ' + E_R +
                             ' --include-band Depth ' +
                             Dir_Grid + '/' + Project_N + '.csar ' +
                             Dir_Grid + '/' + Project_N + '.geotiff' + '\n')

                ##Create Surface Containing only Daily Julian Day Data
                Import.write('carisbatch --run CreateHIPSGrid --gridding-method ' + G_M +
                             ' --resolution ' + str(Res) + ' --Output-crs ' + CRS2 +
                             ' --output-vertical-crs ' + VCRS +
                             ' --compute-band SHOAL --compute-band DEEP --compute-band DENSITY' +
                             ' --compute-band MEAN --compute-band STD_DEV ' +
                             ' --include-flag ACCEPTED --include-flag EXAMINED --include-flag OUTSTANDING ' +
                             ' --keep-up-to-date ' + ' --iho-order ' + IHO +
                             ' --disambiguation-method DENSITY_LOCALE')
                Import.write(r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' ' +
                             Dir_Grid + '/' + str(JD) + '_' + str(Year) + '.csar' +
                             ' > ' + Out + '/' + JD + '/' + '6.Create_Project_Grid_' + JD + '_' + Year + '.txt' + '\n')
                Import.write('carisbatch --run ExportRaster --output-format ' + E_R +
                             ' --include-band Depth ' +
                             Dir_Grid + '/' + str(JD) + '_' + str(Year) + '.csar ' +
                             Dir_Grid + '/' + str(JD) + '_' + str(Year) + '.geotiff')

        elif self.GRID.get()==2:

            with open("Create_Grid.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('@ECHO Creating and Adding to Grid' + '\n')

                ## Add Hips to Project Surface
                Import.write('carisbatch --run AddToHIPSGrid ')
                Import.write(r' file:///' + HDCS_Folder + '/'+ HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' ' +
                            Dir_Grid + '/' + Project_N + '.csar' +
                            ' > ' + Out + '/' + JD + '/' + '6.Add_to_Grid_' + JD + '_' + Year + '.txt' + '\n')

                ##Create Surface Containing only Daily Julian Day Data
                Import.write('carisbatch --run CreateHIPSGrid --gridding-method ' + G_M +
                             ' --resolution ' + str(Res) + ' --Output-crs ' + CRS2 +
                             ' --output-vertical-crs ' + VCRS +
                             ' --compute-band SHOAL --compute-band DEEP --compute-band DENSITY' +
                             ' --compute-band MEAN --compute-band STD_DEV ' +
                             ' --include-flag ACCEPTED --include-flag EXAMINED --include-flag OUTSTANDING ' +
                             ' --keep-up-to-date ' + ' --iho-order ' + IHO +
                             ' --disambiguation-method DENSITY_LOCALE')
                Import.write(r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' ' +
                             Dir_Grid + '/' + str(JD) + '_' + str(Year) + '.csar' +
                             ' > ' + Out + '/' + JD + '/' + '6.Create_JD_Grid_' + JD + '_' + Year + '.txt' + '\n')
                Import.write('carisbatch --run ExportRaster --output-format ' + E_R +
                         ' --include-band Depth ' +
                         Dir_Grid + '/' + Project_N + '.csar ' +
                         Dir_Grid + '/' + Project_N + '.geotiff' + '\n')
                Import.write('carisbatch --run ExportRaster --output-format ' + E_R +
                         ' --include-band Depth ' +
                         Dir_Grid + '/' + str(JD) + '_' + str(Year) + '.csar ' +
                         Dir_Grid + '/' + str(JD) + '_' + str(Year) + '.geotiff')

        p = S.check_call("Create_Grid.bat", stdin=None, stdout=None, stderr=None, shell=False)


    def Create_Backscatter(self): ## Not running Code

        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        Dir_Grid = self.GRID_DIR.get()
        HDCS_Folder = self.HDCS_D.get()
        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()
        crs = self.CRS_O.get()
        CRS = crs.partition(": ")[2]
        SIPS_RES = self.RES.get()

        if self.BACK_S.get()==1:

            with open("Create_BackScatter.bat", "w") as Import:
                Import.write('@ECHO OFF' + '\n')
                Import.write('cd '+ Caris + '\n')
                Import.write('carisbatch --run CreateSIPSMosaic --mosaic-engine SIPS_BACKSCATTER --resolution ' + SIPS_RES +
                             ' --blending WEIGHTED --output-crs ' + CRS)
                Import.write(r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' ' +
                             Dir_Grid + '/BackScatter' + str(JD) + '_' + str(Year) + '.csar' +
                             ' > ' + Out + '/' + JD + '/' + '7.Create_BS' + JD + '_' + Year + '.txt' + '\n')

    def Combine_Caris_Output(self):

        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()
        chdir(Out + '/' + JD)
        OutFiles = [f for f in listdir(Out + '/' + JD) if re.match(r'[0-9]+.*\.txt', f)]
        with open('Caris_Output_' + str(JD) + '_' + str(Year) + '.txt', 'w') as fout, fileinput.input(OutFiles) as fin:
            for line in fin:
                fout.write(line)
        startfile('Caris_Output_' + str(JD) + '_' + str(Year) + '.txt')


    def Sub_Rep(self):

        chdir(owd)

        self.sub_op = LabelFrame(frame9, text="Reports, QC and Submissions", foreground="blue")
        self.sub_op.grid(row=0, column=0, sticky=W)

        self.D_R = IntVar()
        self.D_r = Checkbutton(self.sub_op, onvalue=1, offvalue=0, variable=self.D_R, text= "Reports and QC",
                               command=self.Load_Daily_Reports)
        self.D_r.grid(row=0, column=0, sticky=W)

        self.Finalize = IntVar()
        self.Finalize_t = Checkbutton(self.sub_op, onvalue=1, offvalue=0, variable=self.Finalize, text="Finalization and Submission",
                                      command=self.Load_Finalization_Submission)
        self.Finalize_t.grid(row=0, column=1, sticky=W)

        self.HDC_F = IntVar()
        self.HDC_f = Checkbutton(self.sub_op, onvalue=1, offvalue=0, variable=self.HDC_F, text= "HDC ISO Submission Form",
                                    command=self.Load_ATL_SUB_ISO)
        self.HDC_f.grid(row=0, column=2, sticky=W)

        self.RUN_FF = IntVar()
        self.RUN_ff = Checkbutton(self.sub_op, onvalue=1, offvalue=0, variable=self.RUN_FF, text= "Flier Finder",
                                    command=self.Load_FlierFinder)
        self.RUN_ff.grid(row=0, column=3, sticky=W)

        self.BP = IntVar()
        self.Bp = Checkbutton(self.sub_op, onvalue=1, offvalue=0, variable=self.BP, text= "Bounding Polygon",
                                    command=self.Load_BoundingPoly)
        self.Bp.grid(row=0, column=4, sticky=W)

        


    def Load_Daily_Reports(self):

        chdir(owd)

        if self.D_R.get() == 1:

            self.Daily = LabelFrame(frame9, text="Create Daily Excel Reports", foreground ="blue")
            self.Daily.grid(row=1, column=0, sticky=W)


##            self.LINE_F = StringVar()
##            self.LINE_f = Entry(self.Daily, width=38, textvariable=self.LINE_F)
##            self.LINE_ftext = Label(self.Daily, text="Caris Line Report File")
##            self.LINE_ftext.grid(row=1, column=0, sticky=W)
##            self.LINE_f.grid(row=1, column=1, sticky=W)
##            self.ButtonLINE_f = Button(self.Daily, text="...", height=0,
##                                  command=self.Search_LINE_File)
##            self.ButtonLINE_f.grid(row=1, column=2, sticky=W, padx=2)

            self.SVPDir = StringVar()
            self.SVP_Dir = Entry(self.Daily, width=38, textvariable=self.SVPDir)
            self.SVP_Dirtext = Label(self.Daily, text="Julian Day SVP Directory")
            self.SVP_Dirtext.grid(row=1, column=0, sticky=W)
            self.SVP_Dir.grid(row=1, column=1, sticky=W)
            self.ButtonSVP_Dir = Button(self.Daily, text="...", height=0,
                                  command=self.Search_SVP)
            self.ButtonSVP_Dir.grid(row=1, column=2, sticky=W, padx=2)
            


            self.REP_F = StringVar()
            self.REP_f = Entry(self.Daily, width=38, textvariable=self.REP_F)
            self.REP_ftext = Label(self.Daily, text="Daily Report Spreadsheet")
            self.REP_ftext.grid(row=2, column=0, sticky=W)
            self.REP_f.grid(row=2, column=1, sticky=W)
            self.ButtonREP_f = Button(self.Daily, text="...", height=0,
                                  command=self.Search_SpreadSheet_File)
            self.ButtonREP_f.grid(row=2, column=2, sticky=W, padx=2)

            self.WREP_F = StringVar()
            self.WREP_f = Entry(self.Daily, width=38, textvariable=self.WREP_F)
            self.WREP_ftext = Label(self.Daily, text="Weekly Report Spreadsheet")
            self.WREP_ftext.grid(row=3, column=0, sticky=W)
            self.WREP_f.grid(row=3, column=1, sticky=W)
            self.ButtonWREP_f = Button(self.Daily, text="...", height=0,
                                  command=self.Search_SpreadSheet_File2)
            self.ButtonWREP_f.grid(row=3, column=2, sticky=W, padx=2)

            self.IHO_ORDER2 = StringVar()
            iho_op2 = ['EXCLUSIVE',
                       'SPECIAL',
                       '1A',
                       '1B',
                       '2',
                       '3']
            self.IHO2_op = ttk.Combobox(self.Daily, values=iho_op2, width=10, textvariable=self.IHO_ORDER2)
            self.IHO2_text = Label(self.Daily, text="Choose IHO Order")
            self.IHO2_text.grid(row=4, column=0, sticky=W)
            self.IHO2_op.grid(row=4, column=1, sticky=W+E, padx=0)

            self.WeekNO = StringVar()
            self.Weekno = Entry(self.Daily, width=20, textvariable=self.WeekNO)
            self.Weeknotext = Label(self.Daily, text="Week #, JD#-JD#")
            self.Weeknotext.grid(row=5, column=0, sticky=W)
            self.Weekno.grid(row=5, column=1, sticky=W)

            self.New_Sheet = IntVar()
            self.New_Sheet_B = Checkbutton(self.Daily, onvalue=1, offvalue=0, variable=self.New_Sheet, text= "Create a New Sheet")
            self.New_Sheet_B.grid(row=7, column=0, sticky=W)

            self.TPUQC=IntVar()
            Radiobutton(self.Daily, text= "HIPS Points QC", variable=self.TPUQC, value=2).grid(row=6, column=0, sticky= W, padx=1)
            Radiobutton(self.Daily, text= "Surface QC", variable=self.TPUQC, value=1).grid(row=6, column=1, sticky= W, padx=1)

            self.Button_Rep = Button(self.Daily, text="Create Daily, Weekly Reports", height=0,
                               command=self.Run_Daily_Report)
            self.Button_Rep.grid(row=8, column=0, sticky=W, padx=2)

            Parameters = pd.read_csv('Parameters.txt', delimiter=',', header=None)
            #Line = (Parameters.iloc[14,0])
            Daily = Parameters.iloc[14,1]
            Weekly = Parameters.iloc[14,2]
            Weekno = Parameters.iloc[14,4]
            IHOOrder = Parameters.iloc[14,3]
            QC = Parameters.iloc[14,5]

            ##Setting defaults from Parameter file for SBET
            #self.LINE_F.set(Line)
            self.REP_F.set(Daily)
            self.WREP_F.set(Weekly)
            self.WeekNO.set(Weekno)
            self.IHO_ORDER2.set(IHOOrder)
            self.TPUQC.set(QC)

        else:
            try:
                ##Forget Options
                self.Daily.grid_forget()
            except AttributeError:
                pass

        #tip_LINE_F = ToolTip(self.LINE_f, str(self.LINE_F.get()))
        tip_REP_F = ToolTip(self.REP_f, (self.REP_F.get()))


    def Load_BoundingPoly(self):

        if self.BP.get() == 1:
            self.BoundingP = LabelFrame(frame9, text="Create Bounding Polygon", foreground="blue")
            self.BoundingP.grid(row=5, column=0, sticky=W)

            self.VALSRC_F = StringVar()
            self.VALSRC_f = Entry(self.BoundingP, width=38, textvariable=self.VALSRC_F)
            self.VALSRC_text = Label(self.BoundingP, text="Surface Folder")
            self.VALSRC_text.grid(row=1, column=0, sticky=W)
            self.VALSRC_f.grid(row=1, column=1, sticky=W)
            self.ButtonSF = Button(self.BoundingP, text="...", height=0,
                                      command=self.Search_VALSRC_Folder)
            self.ButtonSF.grid(row=1, column=2, sticky=W, padx=2)

            self.Button_Final = Button(self.BoundingP, text="Create Bounding Polygons", height=0,
                               command=self.Create_BoundingPoly)
            self.Button_Final.grid(row=20, column=0, sticky=W, padx=2)

        else:
            try:
                ##Forget Options for Finalization
                self.BoundingP.grid_forget()
            except AttributeError:
                pass
        


    def Load_Finalization_Submission(self):

        if self.Finalize.get() == 1:
            self.Finalization = LabelFrame(frame9, text="Finalize and Submission", foreground="blue")
            self.Finalization.grid(row=2, column=0, sticky=W)

            self.ArcFinalize = IntVar()
            self.ArcFinalize_B = Checkbutton(self.Finalization, onvalue=1, offvalue=0, variable=self.ArcFinalize, text= "Arctic Proccessing")
            self.ArcFinalize_B.grid(row=2, column=0, sticky=W)

            self.VALSRC_F = StringVar()
            self.VALSRC_f = Entry(self.Finalization, width=38, textvariable=self.VALSRC_F)
            self.VALSRC_text = Label(self.Finalization, text="Surface Folder")
            self.VALSRC_text.grid(row=1, column=0, sticky=W)
            self.VALSRC_f.grid(row=1, column=1, sticky=W)
            self.ButtonSF = Button(self.Finalization, text="...", height=0,
                                      command=self.Search_VALSRC_Folder)
            self.ButtonSF.grid(row=1, column=2, sticky=W, padx=2)

            self.Convert11to4 = IntVar()
            self.Convert11to4_B = Checkbutton(self.Finalization, onvalue=1, offvalue=0, variable=self.Convert11to4, text= "Convert Surface to Base 4.4")
            self.Convert11to4_B.grid(row=2, column=1, sticky=W)

            self.Button_Final = Button(self.Finalization, text="Finalize Surfaces", height=0,
                               command=self.FinalizeQC)
            self.Button_Final.grid(row=20, column=0, sticky=W, padx=2)


            self.SURSTA = StringVar()
            self.Sursta = Entry(self.Finalization, width=20, textvariable=self.SURSTA)
            self.Surstatext = Label(self.Finalization, text="SURSTA (YYYYMMDD)")
            self.Surstatext.grid(row=3, column=0, sticky=W)
            self.Sursta.grid(row=3, column=1, sticky=W)

            self.SUREND = StringVar()
            self.Surend = Entry(self.Finalization, width=20, textvariable=self.SUREND)
            self.Surendtext = Label(self.Finalization, text="SUREND (YYYYMMDD)")
            self.Surendtext.grid(row=4, column=0, sticky=W)
            self.Surend.grid(row=4, column=1, sticky=W)

            self.POSACC = StringVar()
            self.Posacc = Entry(self.Finalization, width=20, textvariable=self.POSACC)
            self.Posacctext = Label(self.Finalization, text="POSACC (m)")
            self.Posacctext.grid(row=5, column=0, sticky=W)
            self.Posacc.grid(row=5, column=1, sticky=W)

            self.POSHDW= StringVar()
            pos_meth = ['APOSMV',
                        'APOSAV',]
            self.pos_hdw_op = ttk.Combobox(self.Finalization, values=pos_meth, width=1, textvariable=self.POSHDW)
            self.pos_hdw_text = Label(self.Finalization, text="POSHDW")
            self.pos_hdw_text.grid(row=6, column=0, sticky=W)
            self.pos_hdw_op.grid(row=6, column=1, sticky=W+E, padx=0)

            self.TECPOS= StringVar()
            tec_pos = ['RTK',
                       'PPKGPS',
                        'DGPS',
                        'WAAS']
            self.tec_pos_op = ttk.Combobox(self.Finalization, values=tec_pos, width=1, textvariable=self.TECPOS)
            self.tec_pos_text = Label(self.Finalization, text="TECPOS")
            self.tec_pos_text.grid(row=7, column=0, sticky=W)
            self.tec_pos_op.grid(row=7, column=1, sticky=W+E, padx=0)

            self.COLCMETH= StringVar()
            tec_pos = ['SIS',
                       'QINSY',
                        'HYPACK']
            self.colc_meth_op = ttk.Combobox(self.Finalization, values=tec_pos, width=1, textvariable=self.COLCMETH)
            self.colc_meth_op_text = Label(self.Finalization, text="Collection Method")
            self.colc_meth_op_text.grid(row=8, column=0, sticky=W)
            self.colc_meth_op.grid(row=8, column=1, sticky=W+E, padx=0)

            self.Finalization_msg = LabelFrame(frame9, text="Finalization Message", foreground="blue")
            self.Finalization_msg.grid(row=10, column=0, padx=1, sticky=N+W)

            msg = ('PLEASE READ ALL 5 MEASSAGES BEFORE RUNNING PROCCESS\n'
                   '1. Please check that the name of the project is correct and populated '
                   'as this will be used to '
                   'populated the '
                   'Project, Location, Vessel, and System '
                   'metadata fields in the VALSRC ISO Form.\n'
                   'boundary polygon projection.\n'
                   '\n2. Ensure to select the Apply Tide '
                   'option and the corrsiponding tide reduction used (GPS or Observed/Predicted '
                   'on the Caris Hips Proccessing '
                   'tab, and populate the '
                   'Tide Model or Tide file '
                   'option in the Apply Tides tab. This will be used to populate '
                   'the Tide reduction field in the VALSRC ISO Form.\n'
                   '\n3. For regular Atlantic proccessing all surfaces will be Finalized with the Depth band only, and '
                   'TVU graphs and Caris surface QC reports created in the QC folder.\n'
                   '\n4. For Arctic Proccessing the surfaces will be Finalized with the Depth and Uncertainty bands, '
                   'TVU graphs and Caris surface QC reports created in the QC folder, '
                   'surfaces will be transformed from EPSG:5937 into ESPG8999@2010 with PACD set as the vertical reference, '
                   'and the surfaces will be cut to the Arctic Tiles.\n'
                   '\n5. Lastly the bounding polygon for each surfaces will be created and added to the corrisponding bounding polygon band  '
                   'with all interior holes removed.-CRS Issue need further investigatetion take the Final shapefile and manually import\n'
                   'Atlantic: VALSRC_cvrage(A)_FinalBP.shp\n'
                   'Arctic: Warped_VALSRC#_Tile#_Extractcvrage(A)_FinalBP.shp'
                   ) ## User Reminder

            self.User_Msg = Text(self.Finalization_msg, width=55, height=17, wrap=WORD)
            self.User_Msg.insert(END, msg)
            self.User_Msg.config(state='disabled')
            self.User_Msg.grid(row=0, column=1, padx=1, sticky=W)


        else:
            try:
                ##Forget Options for Finalization
                self.Finalization.grid_forget()
                self.Finalization_msg.grid_forget()
            except AttributeError:
                pass


    def ISO_1001_07_F02(self):

        PH = self.split_Project_Name()
        Projectlist = PH[2]
        VF = self.VALSRC_F.get()
        V_F = VF + '/' + self.Finalized_Folder
        list_VF = listdir(V_F)
        self.QC_Folder = ('QC')

        if self.T_T.get()==1:
            Model_Path = self.M_F.get()
            Model_Pathsplit = path.split(Model_Path)
            Model_File = Model_Pathsplit[1]

        Projectno = Projectlist[0]
        Location = Projectlist[1]
        Vessel = Projectlist[3]
        SNDTYP = Projectlist[4]
        TRDCT = Model_File
        SURSTA = self.SURSTA.get()
        SUREND = self.SUREND.get()
        POSHDW = self.POSHDW.get()
        TECPOS = self.TECPOS.get()
        COLCMETH = self.COLCMETH.get()

        POSACC = str(self.POSACC.get() + 'm')
        CATZOC = None

        for V in list_VF:
            if V.endswith('.csar'):
                File_Name = V.replace(".csar", "")
                ASCII_Out = pd.read_csv(VF + '/' + self.QC_Folder + '/' + File_Name + '_FinalizedQC.txt', sep=' ', header=0, low_memory=False)
                ASCII_Out.columns = ["Lat", "Long", "Depth", "Depth TPU"]
                ASCII_Out["Weight"] = 1/(ASCII_Out['Depth'])
                ASCII_Out["W*TVU"] = ASCII_Out['Weight'] * ASCII_Out['Depth TPU']
                SOUACC = (str(round((ASCII_Out["W*TVU"].sum())/(ASCII_Out["Weight"].sum()),2)) + 'm')
                
                source_part = File_Name.split('_')
                VALSRC = source_part[0]

                raster = cov.Raster(V_F + '/' + V)
                Metadata = raster.iso19139_xml
                Csarxml = open(V_F + '/' + File_Name + '.xml', "w")
                Csarxml.write(Metadata)

                #Parse the Csar xml
                xmldoc = minidom.parse(V_F + '/' + str(File_Name) + '.xml')
                res_xml_loc = (xmldoc.getElementsByTagName('gco:Measure')[0])
                res_child_loc  =res_xml_loc.childNodes[0].nodeValue
                RES = str(res_child_loc)

                CRS_xml_loc = (xmldoc.getElementsByTagName('gco:CharacterString')[4])
                CRS_child_loc = CRS_xml_loc.childNodes[0].nodeValue
                CRS_T = CRS_child_loc.partition(',')[0]

                if CRS_T == 'PROJCS["NAD83(CSRS) / UTM zone 19N"':
                    HORDAT = 'NAD83(CSRS)/ UTM zone 19N@2010'
                elif CRS_T == 'PROJCS["NAD83(CSRS) / UTM zone 20N"':
                    HORDAT = 'NAD83(CSRS)/ UTM zone 20N@2010'
                elif CRS_T == 'PROJCS["NAD83(CSRS) / UTM zone 21N"':
                    HORDAT = 'NAD83(CSRS)/ UTM zone 21N@2010'
                    
                if CRS_T == 'PROJCS["WGS 84 / UTM zone 19N"':
                    HORDAT = 'WGS84/ UTM zone 19N@2010'
                elif CRS_T == 'PROJCS["WGS 84 / UTM zone 20N"':
                    HORDAT = 'WGS84/ UTM zone 20N@2010'
                elif CRS_T == 'PROJCS["WGS 84 / UTM zone 21N"':
                    HORDAT = 'WGS84/ UTM zone 21N@2010'
                    
                elif CRS_T == 'GEOGCS["ITRF2014"':
                    HORDAT = 'ITRF2014@2010'
                elif CRS_T == 'GEOGCS["ITRF2008"':
                    HORDAT = 'ITRF2008@2010'
    
                elif CRS_T == 'PROJCS["WGS 84 / EPSG Canada Polar Stereographic"':
                    HORDAT = 'ITRF2008@2010'

                META = [Location,
                        None,
                        'Validated',
                        SURSTA,
                        SUREND,
                        VALSRC,
                        Vessel,
                        'Canadian Hydrographic Service',
                        '1 : XXX',
                        'Multibeam',
                        'Controlled Survey',
                        POSHDW,
                        POSACC,
                        TECPOS,
                        HORDAT,
                        SNDTYP,
                        'Enter manually or use dropdown list in cell below',
                        'MULTBM - Multibeam',
                        SOUACC,
                        'PATCHT',
                        'POSPAC, HIPS',
                        'Enter manually or use dropdown list in cell below',
                        COLCMETH,
                        'N/A',
                        RES,
                        None,
                        TRDCT,
                        None,
                        'N/A',
                        CATZOC,
                        'Open/Restricted',
                        VALSRC]

                MetaData = openpyxl.load_workbook('1001-07-A-F02_Meta_Data_Template.xlsx')
                worksheet = MetaData.get_sheet_by_name('301-F03')
                worksheet['A4'] = ('Project # ' + Projectno + ' CHS Atlantic')
                worksheet['A5'] = ('METADATA FOR SOURCE #  ' + VALSRC)

                C = 8
                for MetaD in META:
                    worksheet['C' + str(C)] = MetaD
                    C = C + 1
                MetaData.save( V_F + '/'  + '1001-07-A-F02_' + str(File_Name) + '_Meta_Data.xlsx')


    def Load_ATL_SUB_ISO(self):

        if self.HDC_F.get() == 1:
            self.Dir_Form = LabelFrame(frame9, text="Create ATL ISO 1001 07 AF01", foreground="blue")
            self.Dir_Form.grid(row=3, column=0, sticky=W)

            self.SIG = StringVar()

            self.Sig = Entry(self.Dir_Form, width=5, textvariable=self.SIG)
            self.Sig_text = Label(self.Dir_Form, text="Intials")
            self.Sig_text.grid(row=1, column=0, sticky=W)
            self.Sig.grid(row=1, column=1, sticky=W)

            self.SUB_D = StringVar()
            self.SUB_d = Entry(self.Dir_Form, width=38, textvariable=self.SUB_D)
            self.SUB_dtext = Label(self.Dir_Form, text="Submission Directory")
            self.SUB_dtext.grid(row=2, column=0, sticky=W)
            self.SUB_d.grid(row=2, column=1, sticky=W)

            self.Buttonsub = Button(self.Dir_Form, text="...", height=0,
                                  command=self.Search_Sub_dir_file)
            self.Buttonsub.grid(row=2, column=2, sticky=W, padx=2)

            self.ButtonC = Button(self.Dir_Form, text="Fill HDC ISO Sub Form", height=0,
                                  command=self.ISO_1001_07_A_F01)
            self.ButtonC.grid(row=3, column=0, sticky=W, padx=2)

        else:
            try:
                ##Forget Options for ISO 1001 07 A F01 Form
                self.Dir_Form.grid_forget()
            except AttributeError:
                pass


    def Load_FlierFinder(self):

        if self.RUN_FF.get() == 1:

            self.FlierFinder = LabelFrame(frame9, text="QCTools Flier Finder", foreground="blue")
            self.FlierFinder.grid(row=4, column=0, sticky=W)

            self.DTM_DIR = StringVar()
            self.DTM_dir = Entry(self.FlierFinder, width=38, textvariable=self.DTM_DIR)
            self.DTM_dir_text = Label(self.FlierFinder, text="Surface Folder")
            self.DTM_dir_text.grid(row=2, column=0, sticky=W)
            self.DTM_dir.grid(row=2, column=1, sticky=W)
            self.ButtonSF = Button(self.FlierFinder, text="...", height=0,
                                      command=self.Search_DTMFolder)
            self.ButtonSF.grid(row=2, column=2, sticky=W, padx=2)

            self.FHEIGHT = StringVar()
            heights = ['AUTO',
                        '1','2','3','4','5','6','7','8','9','10']
            self.FHEIGHT_op = ttk.Combobox(self.FlierFinder, values=heights, textvariable=self.FHEIGHT)
            self.FHEIGHT_text = Label(self.FlierFinder, text="Flier Height (m)")
            self.FHEIGHT_text.grid(row=1, column=0, sticky=W)
            self.FHEIGHT_op.grid(row=1, column=1, sticky=W+E, padx=0)

            self.ButtonFF = Button(self.FlierFinder, text="Run Flier Finder", height=0,
                                  command=self.Find_Fliers)
            self.ButtonFF.grid(row=3, column=0, sticky=W, padx=2)

        else:
            try:
                ##Forget Options for QCTools Flier Finder
                self.FlierFinder.grid_forget()
            except AttributeError:
                pass

        
    def FinalizeQC(self):

        V_F = self.VALSRC_F.get()
        list_VF = listdir(V_F)
        self.Finalized_Folder = ('Finalized_Surfaces')
        self.QC_Folder = ('QC')
        chdir(V_F)
        if path.exists(self.Finalized_Folder):
            pass
        else:
            mkdir(self.Finalized_Folder)
        if path.exists(self.QC_Folder):
            pass
        else:
            mkdir(self.QC_Folder)
        chdir(owd)

        for V in list_VF:
            if V.endswith(".csar"):
                self.V = V.replace(".csar", "")
                self.Finalize_QC()
        
        self.Finalize_Surfaces()
        self.ISO_1001_07_F02()

        if self.ArcFinalize.get() == 1:
            self.Cut_Folder = (V_F + '/' + self.Finalized_Folder + '/CutSurfaces')
            if path.exists(self.Cut_Folder):
                pass
            else:
                mkdir(self.Cut_Folder)
            self.Warped_Folder = (V_F + '/' + self.Finalized_Folder + '/WarpedSurfaces')
            if path.exists(self.Warped_Folder):
                pass
            else:
                mkdir(self.Warped_Folder)
            if path.exists(self.Cut_Folder +'/WKT_Tiles'):
                pass
            else:
                mkdir(self.Cut_Folder +'/WKT_Tiles')
            self.Warping(V_F)
            self.TileCut()

            Cutsursdir = self.Cut_Folder
            Cutsurs = listdir(Cutsursdir)
            for cutsur in Cutsurs:
                if cutsur.endswith(".csar"):
                    pass
                    self.Vectorize_Raster2(cutsur, Cutsursdir)
            dircut = listdir(Cutsursdir)
            
            for shp in dircut:
                if shp.endswith('.shp'):
                    Prj = 'EPSG:8999'
                    self.Remove_Holes(Cutsursdir + '/' + shp, Prj)
        else:
            Fsursdir = V_F + '/' + self.Finalized_Folder
            Fsurs = listdir(Fsursdir)
            for Fsur in Fsurs:
                if Fsur.endswith(".csar"):
                    self.Vectorize_Raster2(Fsur, Fsursdir)
            Fsurs2 = listdir(Fsursdir)
                    
        for shp in Fsurs2:
            if shp.endswith('.shp'):
                Prj = 'EPSG:4326'
                self.Remove_Holes(Fsursdir + '/' + shp, Prj)
        
        if self.Convert11to4.get() ==1:
            self.HIPS11_to_BASE4()


    def Create_BoundingPoly(self):

        V_F = self.VALSRC_F.get()
        list_VF = listdir(V_F)

        for V in list_VF:

            if V.endswith(".csar"):
                self.V = V.replace(".csar", "")
                self.Vectorize_Raster2(V, V_F)

        list_VF2 = listdir(V_F)
        for shp in list_VF2:
            if shp.endswith('.shp'):
                Prj = 'EPSG:4326'
                self.Remove_Holes(V_F + '/' + shp, Prj)

  
    def Vectorize_Raster2(self, surface, Out):

        
        Surface = surface.split('.')
        
        hob = (Out + '/' + Surface[0] + '.hob')
        FilePrefix = Surface[0]
        shp = (Out + '/' + Surface[0] + '_cvrage(A).shp')

        with open("Vectorize_Raster.bat", "w") as Area_C:
            Area_C.write('@ECHO OFF' + '\n')
            Area_C.write('@ECHO Vectorizing Surface' + '\n')
            Area_C.write('cd '+ Caris + '\n')
            Area_C.write('carisbatch --run  VectorizeRaster --input-band Depth --feature-catalogue "Bathy DataBASE"' +
                         ' --polygon-feature cvrage --mode COVERAGE ' + Out + '/' + surface + ' ' + hob + '\n')
            Area_C.write('carisbatch --run ExportFeaturesToShapefile --feature-catalogue "Bathy DataBASE"' +
                         ' --file-prefix ' + FilePrefix + ' ' + hob + ' ' + Out)

        p = S.check_call("Vectorize_Raster.bat", stdin=None, stdout=None, stderr=None, shell=False)


    def Remove_Holes(self, Shpf, Prj):


        Poly = shapefile.Reader(Shpf)
        SHP = Shpf.split('.shp')

        g = []

        for s in Poly.shapes():
            g.append(pygeoif.geometry.as_shape(s))
        
        m = pygeoif.MultiPolygon(g)
        P = wkt.loads(str(m))

        omega = cascaded_union([
        Polygon(component.exterior) for component in P])


        features = [i for i in range(len(P))]
        gdr = gpd.GeoDataFrame({'feature': features, 'geometry': omega}, crs=Prj)
        #gdr.to_crs("EPSG:2961")
        gdr.to_file(SHP[0] + '_FinalBP.shp')
        

##        BP = shapefile.Reader(SHP[0] + '_FinalBP.shp')
##        g2 = []
##        for s2 in BP.shapes():
##            g2.append(pygeoif.geometry.as_shape(s2))
##        
##        m2 = pygeoif.MultiPolygon(g2)
##        BP_wkt = m2.wkt


##        SurName = Shpf.split('cvrage(A).shp')
##        options = cov.Options(open_type=cov.OpenType.WRITE)
##        raster = cov.Raster(SurName[0] + '.csar', options=options)
##        BPWKT = TransformCoordinates(BP_wkt, raster)
##        raster.bounding_polygon = BPWKT


    def Warping(self, V_F):
 
        s_list = listdir(V_F + '/Finalized_Surfaces')

        with open("Warp.bat", "w") as Import:
            Import.write('@ECHO OFF' + '\n')
            Import.write('@ECHO Warp Surfaces' + '\n')
            Import.write('cd '+ BASE5 + '\n')
            chdir(V_F)
            for s in s_list:
                if s.endswith('.csar'):
                   
                    raster = cov.Raster(s)
                    Metadata = raster.iso19139_xml
                    Csarxml = open(s + '.xml', "w")
                    Csarxml.write(Metadata)

                    xmldoc = minidom.parse(s + '.xml')
                    res_xml_loc = (xmldoc.getElementsByTagName('gco:Measure')[0])
                    res_child_loc  =res_xml_loc.childNodes[0].nodeValue
                    RES = str(res_child_loc)

                    Import.write('carisbatch --run WarpRaster' +
                                 ' --output-crs EPSG:8999@2010 --input-band Depth BICUBIC NONE' +
                                 ' --output-vertical-crs CUSTOM:69036444' +
                                 ' --primary-band Depth NEAREST_NEIGHBOUR NONE' +
                                 ' --input-band Uncertainty NEAREST_NEIGHBOUR NONE' +
                                 ' --resolution ' + RES + 'm' + 
                                 ' --reprojection-method EXACT ' +
                                 V_F + '/' + self.Finalized_Folder + '/' + s + ' ' + self.Warped_Folder +'/' + 'Warped_' + s + '\n')
        chdir(owd)

        p = S.check_call('Warp.bat', stdin=None, stdout=None, stderr=None, shell=False)



    def TileCut(self):

        WGS_84 = ('GEOGCS["WGS 84",DATUM["World Geodetic System 1984",SPHEROID["WGS 84",6378137,298.2572235629972,AUTHORITY["EPSG","7030"]],' +
                  'AUTHORITY["EPSG","6326"]],PRIMEM["Greenwich",0,AUTHORITY["EPSG","8901"]],UNIT["degree (supplier to define representation)",' +
                  '0.0174532925199433,AUTHORITY["EPSG","9122"]],AUTHORITY["EPSG","4326"]]')
        ITRF_2008 =('GEOGCS["ITRF2008",DATUM["International Terrestrial Reference Frame 2008",SPHEROID["GRS 1980",6378137,298.2572221010041,' +
                    'AUTHORITY["EPSG","7019"]],TOWGS84[0,0,0,0,0,0,0],AUTHORITY["EPSG","1061"]],PRIMEM["Greenwich",0,AUTHORITY["EPSG","8901"]],' +
                    'UNIT["degree (supplier to define representation)",0.0174532925199433,AUTHORITY["EPSG","9122"]],' +
                    'EXTENSION["tx_authority","EPSG:7666"],AUTHORITY["EPSG","8999"]]')
        

        Surfaces = self.Warped_Folder
        Csar_Dir = self.Cut_Folder
        chdir(owd)
        Grid_Tiles = ('ArcticGrid/ARCTIC_CELLS_CUT_TOOL_10x10_V2.shp')
        
        sf1 = shapefile.Reader(Grid_Tiles)
        
        i = 0
        for s in sf1.shapes():

            rec = sf1.record(i)
            Tno = rec[8]
            #print(Tno)
            wkt = pygeoif.geometry.as_shape(s)
          
            f = open(self.Cut_Folder + '/WKT_Tiles' + '/Tile_' + str(Tno) + '.wkt', "w+")
            f.write(ITRF_2008 + '\n' + str(wkt))
            f.close()
            i = i + 1

        SURFACES = listdir(Surfaces)
        

        
        for Sur in SURFACES:
            if Sur.endswith('.csar'):
                print(Sur)

                self.Vectorize_Raster2(Sur,Surfaces)

                S2 = re.sub(".csar", "", Sur)
                
                match = (S2 + 'cvrage(A).shp')
                SURFACES = listdir(Surfaces)

                polygon = [j for j in SURFACES if match in j]
                

                shp_dataset_a = caris.open(file_name=str(Surfaces + '/') + str(polygon[0]), open_mode=caris.OpenMode.READ_WRITE)
                shp_dataset_b = caris.open(file_name=str(Grid_Tiles), open_mode=caris.OpenMode.READ_WRITE)

                all_BP = shp_dataset_a.query_all()

                Tiles = []
                for feature_a in all_BP:

                    feature_geom = feature_a.geometry
                    
                    
                    intersecting_features = shp_dataset_b.query("ARCTIC_CELLS_CUT_TOOL_10x10_V2", intersects=feature_geom)


                    for feature in intersecting_features:
                            n = str(feature['FID'])
                            if n not in Tiles:
                                Tiles.append(n)
                            Sur = re.sub(".csar", "", Sur)
                            print (str(feature['FID']))

                with open("Extract_" + Sur + ".bat", "w+") as E:
                        E.write('@ECHO OFF' + '\n')
                        E.write('C:' + '\n')
                        E.write('cd '+ str(Caris) + '\n')

                        for T in Tiles:
                                E.write('carisbatch --run ExtractCoverage --include-band ALL --extract-type INCLUSIVE ' +
                                        '--geometry-file ' + self.Cut_Folder + '/WKT_Tiles' + '/Tile_' + T + '.wkt ' + Surfaces + '/' + Sur + '.csar ' +
                                        Csar_Dir + '/' + Sur + '_' + T + '_Extract.csar' + '\n')


                p = S.check_call("Extract_" + Sur + ".bat", stdin=None, stdout=None, stderr=None, shell=False)

       
    def Finalize_QC(self):

        V_F = self.VALSRC_F.get()
        list_VF = listdir(V_F)
        IHO_orders = ['EXCLUSIVE', 'SPECIAL', '1A', '1B', '2', '3']
        t_95_1d = 1.96
        t_95_2d = 2.45

        with open("Coverage_to_ASCII.bat", "w") as ECTA:
            ECTA.write('@ECHO OFF' + '\n')
            ECTA.write('cd '+ Caris + '\n')
            ECTA.write('@ECHO Exporting Coverage to ACSII' + '\n')
            ECTA.write('carisbatch --run ExportCoverageToASCII' +
                         ' --include POSITION Lat 6 DD' +
                         ' --include POSITION Lon 6 DD' +
                         ' --include BAND Depth 2 m' +
                        ' --include BAND Uncertainty 2 m ' +
                        V_F + '/' + self.V +  '.csar ' +
                        ' ' + V_F + '/' + self.QC_Folder + '/' + self.V + '_FinalizedQC.txt')
        p = S.check_call("Coverage_to_ASCII.bat", stdin=None, stdout=None, stderr=None, shell=False)

        ASCII_Out = pd.read_csv(V_F + '/' + self.QC_Folder + '/' + self.V + '_FinalizedQC.txt', sep=' ', header=0, low_memory=False)
        ASCII_Out.columns = ["Lat", "Long", "Depth", "Depth TPU"]

        TVU_QC = pd.DataFrame(columns=['IHO Order', 'Percentage within Allowable', 'Number of Nodes Considered'])

        for orders in IHO_orders:
            TVU_D = []
            for Depth in ASCII_Out['Depth']:
                TPU2 = TPU(orders, Depth)
                TVU_D.append(TPU2[0])
            ASCII_Out['Allowable TVU'] = TVU_D
            ASCII_Out['Within Allowable TVU'] = np.where(ASCII_Out['Depth TPU'] <= ASCII_Out['Allowable TVU'],
                                                'yes', 'no')
            P_W_A_TVU = round(((len(ASCII_Out[ASCII_Out['Within Allowable TVU'] == 'yes'])/len(ASCII_Out)))*100,2)
            TVU_QC = TVU_QC.append({'IHO Order' : orders,  'Percentage within Allowable': P_W_A_TVU, 'Number of Nodes Considered' : len(ASCII_Out) }, ignore_index=True)

            Depth_mean = round(ASCII_Out['Depth'].mean(),3)

            Depth_TPU_max = round(ASCII_Out['Depth TPU'].max(),3)
            Depth_TPU_min = round(ASCII_Out['Depth TPU'].min(),3)
            Depth_TPU_mean = round(ASCII_Out['Depth TPU'].mean(),3)
            Depth_TPU_std = round(ASCII_Out['Depth TPU'].std(),3)
            Depth_95_p = round((Depth_TPU_mean + t_95_1d*Depth_TPU_std),3)
            Depth_95_n = round((Depth_TPU_mean - t_95_1d*Depth_TPU_std),3)

            TPU_v = TPU(orders, Depth_mean)
            fig, ax = plt.subplots(nrows=2)
            D_TPU = list(ASCII_Out.loc[:,'Depth TPU'])
            ax[0].hist(D_TPU, weights=np.ones(len(D_TPU)) / len(D_TPU), alpha=0.5)
            ax[0].axvline(Depth_TPU_max, 0, c='r', label = "MAX = " + str(Depth_TPU_max) + 'm')
            ax[0].axvline(Depth_TPU_mean, 0, c='g', label = "MEAN = " + str(Depth_TPU_mean) + 'm')
            ax[0].axvline(Depth_TPU_min, 0, c='c', label = "MIN = " + str(Depth_TPU_min) + 'm')
            ax[0].axvline(Depth_95_p, 0, c='m', label = "95% Level = " + str(Depth_95_p) + 'm')
            ax[0].axvline(TPU_v[0], 0, c='k', label = "CHS/IHO = " + str(TPU_v[0]) + 'm')
            ax[0].legend(loc='upper right')
            ax[0].set_title('Vertical Accuracy (Ave Depth ' + str(Depth_mean) + ' m)' + '(Order = ' + str(orders) + ')\n')
            ax[0].set_xlabel('Depth Accuracy (m)')
            ax[0].set_ylabel('Percentage (%)')

            maxd = round(ASCII_Out['Depth'].max(),0)+5
            Depths = np.arange(0,maxd,1)
            Contours = [0,2,5,10,15,20,30,50,100]
            ATPU=[]
            for D in Depths:
                ATPU.append(TPU(orders,D)[0])

            colors = {'yes': 'green', 'no':'red'}
            ax[1].scatter(ASCII_Out['Depth'],ASCII_Out['Depth TPU'],c=ASCII_Out['Within Allowable TVU'].map(colors))
            ax[1].plot(Depths, ATPU)
            ax[1].set_title('Vertical Accuracy Scatterplot) ' + '(Order = ' + str(orders) + ')\n Red - Outside Allowable Green - Within Allowable(' + str(P_W_A_TVU) + '%)')
            ax[1].set_xlabel('Depth (m)')
            ax[1].set_ylabel('Depth Accuracy (m)')
            plt.tight_layout()
            plt.savefig(V_F  + '/' + self.QC_Folder + '/' + self.V + '_' + orders + '_Accuracy.png', dpi=200)
        TVU_QC.to_csv(V_F  + '/' + self.QC_Folder + '/' + self.V + '_TVUQC.csv', sep=',' , header=True, index=False)


    def Finalize_Surfaces(self):

        V_F = self.VALSRC_F.get()
        list_VF = listdir(V_F)


        with open("Finalized.bat", 'w') as Final:
            Final.write('@ECHO OFF' + '\n')
            Final.write('@ECHO Finalizing Surfaces' + '\n')
            Final.write('cd '+ Caris + '\n')

            if self.ArcFinalize.get() == 1:
                for VALSRCno in list_VF:
                    if VALSRCno.endswith(".csar"):
                        VALSRCno = VALSRCno.replace(".csar", "")
                        Final.write('carisbatch --run FinalizeRaster --include-band  Depth --include-band Uncertainty '  +
                                    '--apply-designated --uncertainty-source UNCERT ' +
                                    V_F +  '/' + VALSRCno + '.csar ' +
                                    V_F + '/' + self.Finalized_Folder + '/' + VALSRCno + '.csar' +'\n')
            else:
                for VALSRCno in list_VF:
                    if VALSRCno.endswith(".csar"):
                        VALSRCno = VALSRCno.replace(".csar", "")
                        Final.write('carisbatch --run FinalizeRaster --include-band  Depth '  +
                                    '--apply-designated --uncertainty-source UNCERT ' +
                                    V_F +  '/' + VALSRCno + '.csar ' +
                                    V_F + '/' + self.Finalized_Folder + '/' + VALSRCno + '.csar' +'\n')

        p = S.check_call("Finalized.bat", stdin=None, stdout=None, stderr=None, shell=False)


    def HIPS11_to_BASE4(self):
        """ This Function runs processing steps based on user inputs"""
        BASE4 = ('C:/Program Files/CARIS/BASE Editor/4.4/bin')
        HIPS11 =  ('C:/Program Files/CARIS/HIPS and SIPS/11.3/bin')


        Gridding_Method = 'SHOAL'
        V_F = (str(self.VALSRC_F.get()) + '/' + self.Finalized_Folder)
        list_VF = listdir(V_F)


        for file in list_VF:
            if file.endswith(".csar"):
                File_Name = file.replace(".csar", "")

##                raster = Raster(V_F + '/' + file)
##                Metadata = raster.iso19139_xml
##                Csarxml = open(V_F + '/' + File_Name + '.xml', "w")
##                Csarxml.write(Metadata)

                #Parse the Csar xml
                xmldoc = minidom.parse(V_F + '/' + str(File_Name) + '.xml')
                res_xml_loc = (xmldoc.getElementsByTagName('gco:Measure')[0])
                res_child_loc  =res_xml_loc.childNodes[0].nodeValue
                Gridding_Resolution = str(res_child_loc)

                CRS_xml_loc = (xmldoc.getElementsByTagName('gco:CharacterString')[4])
                CRS_child_loc = CRS_xml_loc.childNodes[0].nodeValue
                CRS_T = CRS_child_loc.partition(',')[0]

                if CRS_T == 'PROJCS["NAD83(CSRS) / UTM zone 19N"':
                    CRS = 'EPSG:2960'
                elif CRS_T == 'PROJCS["NAD83(CSRS) / UTM zone 20N"':
                    CRS = 'EPSG:2961'
                elif CRS_T == 'PROJCS["NAD83(CSRS) / UTM zone 21N"':
                    CRS = 'EPSG:2962'
                elif CRS_T == 'GEOGCS["ITRF2014"':
                    CRS = 'EPSG:8999'

                with open("HIPS11_to_BASE4.bat", "w") as Export:
                        Export.write('@ECHO OFF' + '\n')
                        Export.write('cd '+ Caris + '\n')
                        Export.write('@ECHO Exporting Coverage to ACSII' + '\n')
                        Export.write('carisbatch --run ExportCoverageToASCII' +
                                     ' --include POSITION X 3 m  --include POSITION Y 3 m' +
                                     ' --include BAND Depth 3 m ' +
                                     V_F + '/' + file +
                                     '  ' + V_F + '/'  + File_Name +
                                     '.txt' + '\n')
                        Export.write('cd '+ BASE4 + '\n')
                        Export.write('@ECHO Importing Coverage to Base 4' + '\n')
                        Export.write('carisbatch --run ImportPoints --input-format ASCII' +
                                     ' --gridding-method ' + str(Gridding_Method) +
                                     ' --resolution ' + Gridding_Resolution + 'm' +
                                     ' --include-band ALL' +
                                     ' --input-crs ' + CRS +
                                     ' --info-file ' + owd + '/Export_HIPS11_to_BASE4.info ' + V_F + '/' + File_Name  +
                                     '.txt ' + V_F + '/'  + File_Name +
                                     '_4.4.csar'+ '\n')

                p = S.check_call("HIPS11_to_BASE4.bat", stdin=None, stdout=None, stderr=None, shell=False)


    def ExporttoACSII(self):

        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        HDCS_Folder = self.HDCS_D.get()
        crs = self.CRS_O.get() ## Coordinate Ref System
        CRS = crs.partition(": ")[2]
        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Dir_Grid = self.GRID_DIR.get()
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()
        TPUQCFolder = str('TPUQC')
        HTA_Out = (Out + '/' + str(JD) + '/' + TPUQCFolder)
        chdir(Out + '/' + str(JD))
        if path.exists(TPUQCFolder):
            pass
        else:
            mkdir(TPUQCFolder)

        if self.TPUQC.get()==1:
            chdir(HTA_Out)
            if path.exists('Coverage'):
                pass
            else:
                mkdir('Coverage')
            chdir(owd)
            with open("Coverage_to_ASCII.bat", "w") as ECTA:
                ECTA.write('@ECHO OFF' + '\n')
                ECTA.write('cd '+ Caris + '\n')
                ECTA.write('@ECHO Exporting Coverage to ACSII' + '\n')
                ECTA.write('carisbatch --run ExportCoverageToASCII' +
                           ' --include POSITION Lat 6 DD' +
                           ' --include POSITION Lon 6 DD' +
                           ' --include BAND Depth 2 m' +
                           ' --include BAND Uncertainty 2 m ' +
                           Dir_Grid + '/' + JD + '_' + Year + '.csar ' +
                           ' ' + HTA_Out + '/Coverage/' + 'Coverage_' + JD + '_' + Year + '.txt')
            p = S.check_call("Coverage_to_ASCII.bat", stdin=None, stdout=None, stderr=None, shell=False)

        elif self.TPUQC.get()==2:
            chdir(HTA_Out)
            if path.exists('HIPS'):
                pass
            else:
                mkdir('HIPS')
            chdir(owd)

            with open("Hip_to_ASCII.bat", "w") as HTA:
                HTA.write('@ECHO OFF' + '\n')
                HTA.write('cd '+ Caris + '\n')
                HTA.write('@ECHO Exporting HIPS to ACSII' + '\n')
                HTA.write('carisbatch --run ExportHIPS --output-format ASCII --sample 0.5m MIN ALLDATA' +
                            ' --delimiter "," --overwrite --include-flag ACCEPTED --output-crs ' + CRS +
                            ' --include-flag EXAMINED --include-flag OUTSTANDING --include-header --coordinate-format LLDG_DD' +
                            ' --include-attribute DEPTH_PRO --include-attribute POSITION_TPU' +
                            ' --include-attribute Depth_TPU --coordinate-precision 7 --single-file Sampled_HIPS_Data_' + str(JD) + '.txt' +
                            r' file:///' + HDCS_Folder + '/'  + HIPSFILE + '/' + HIPSFILE + '.hips?Vessel=' + Vessel +
                            ';Day=' + str(Year) + '-' + str(JD) + ' ' + HTA_Out + '/HIPS' + '\n')

            p = S.check_call("Hip_to_ASCII.bat", stdin=None, stdout=None, stderr=None, shell=False)


    def Vectorize_Raster(self):

        Dir_Grid = self.GRID_DIR.get()
        Output = self.OUT_F.get()
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        chdir(Output + '/' + JD +'/TPUQC')
        if path.exists('Polygon_' + str(JD)):
            pass
        else:
            mkdir('Polygon_' + str(JD))
        chdir(owd)

        raster = (Dir_Grid + '/' + JD + '_' + Year + '.csar')
        hob = (Output + '/' + JD +'/TPUQC/' + 'Polygon_' + str(JD) + '/' + JD + '_' + Year + '.hob')
        shp = (Output + '/' + JD +'/TPUQC/' + 'Polygon_' + str(JD) + '/' + JD + '_' + Year + '.shp')

        with open("Vectorize_Raster.bat", "w") as Area_C:
            Area_C.write('@ECHO OFF' + '\n')
            Area_C.write('@ECHO Vectorizing Surface' + '\n')
            Area_C.write('cd '+ Caris + '\n')
            Area_C.write('carisbatch --run  VectorizeRaster --input-band Depth --feature-catalogue "Bathy DataBASE"' +
                         ' --polygon-feature cvrage --mode COVERAGE ' + raster + ' ' + hob + '\n')
            Area_C.write('@ECHO Adding Area Geometry' + '\n')
            Area_C.write('carisbatch --run AddGeometryAttributes --feature-catalogue "Bathy DataBASE"' +
                         ' --area ' + hob + ' ' + shp + '\n')

        p = S.check_call("Vectorize_Raster.bat", stdin=None, stdout=None, stderr=None, shell=False)


    def Run_Daily_Report(self):

        JD = self.JULIAN_D.get()
        Year = self.YEAR.get()

        JDS = str(self.GRID_DIR.get() + '/' + str(JD) + '_' + str(Year) +'.csar')

        chdir(str(self.OUT_F.get()))
        if path.exists(str(JD)):
            pass
        else:
            mkdir(str(JD)) ## Create a Julian Day dump folder for all Processing logs and CHSython Output

        chdir(owd)
        if path.isfile(JDS):
            pass
        else:
            self.Create_Addto_Hips_Grid()
            
        self.ExporttoACSII()
        self.Vectorize_Raster()
        self.Plotting()


    def Plotting(self):

        Dir_Grid = self.GRID_DIR.get()
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()
        ##Line_F = self.LINE_F.get()
        Report_F = self.REP_F.get()
        Weekly_Report = self.WREP_F.get()
        Name = self.WeekNO.get()
        New_Sheet = self.New_Sheet.get()
        order = self.IHO_ORDER2.get()
        TPUQCFolder = str('TPUQC')


        t_95_1d = 1.96
        t_95_2d = 2.45

        IHO_orders = ['EXCLUSIVE', 'SPECIAL', '1A', '1B', '2', '3']

        if self.TPUQC.get()==1:
            ASCII_Out = pd.read_csv(Out + '/' + JD + '/' + TPUQCFolder + '/Coverage/' + 'Coverage_' + JD + '_' + Year + '.txt', sep=' ', header=0, low_memory=False)
            ASCII_Out.columns = ["Lat", "Long", "Depth", "Depth TPU"]

            List_TVU = []
            for orders in IHO_orders:
                TVU_D = []
                for Depth in ASCII_Out['Depth']:
                    TPU2 = TPU(orders, Depth)
                    TVU_D.append(TPU2[0])
                ASCII_Out['Allowable TVU'] = TVU_D
                ASCII_Out['Within Allowable TVU'] = np.where(ASCII_Out['Depth TPU'] <= ASCII_Out['Allowable TVU'],
                                                   'yes', 'no')
                P_W_A_TVU = round(((len(ASCII_Out[ASCII_Out['Within Allowable TVU'] == 'yes'])/len(ASCII_Out)))*100,2)

                List_TVU.append(P_W_A_TVU)

                Depth_mean = round(ASCII_Out['Depth'].mean(),3)

                Depth_TPU_max = round(ASCII_Out['Depth TPU'].max(),3)
                Depth_TPU_min = round(ASCII_Out['Depth TPU'].min(),3)
                Depth_TPU_mean = round(ASCII_Out['Depth TPU'].mean(),3)
                Depth_TPU_std = round(ASCII_Out['Depth TPU'].std(),3)
                Depth_95_p = round((Depth_TPU_mean + t_95_1d*Depth_TPU_std),3)
                Depth_95_n = round((Depth_TPU_mean - t_95_1d*Depth_TPU_std),3)

                TPU_v = TPU(orders, Depth_mean)
                fig, ax = plt.subplots(nrows=2)
                D_TPU = list(ASCII_Out.loc[:,'Depth TPU'])
                ax[0].hist(D_TPU, weights=np.ones(len(D_TPU)) / len(D_TPU), alpha=0.5)
                ax[0].axvline(Depth_TPU_max, 0, c='r', label = "MAX = " + str(Depth_TPU_max) + 'm')
                ax[0].axvline(Depth_TPU_mean, 0, c='g', label = "MEAN = " + str(Depth_TPU_mean) + 'm')
                ax[0].axvline(Depth_TPU_min, 0, c='c', label = "MIN = " + str(Depth_TPU_min) + 'm')
                ax[0].axvline(Depth_95_p, 0, c='m', label = "95% Level = " + str(Depth_95_p) + 'm')
                #ax[0].axvline(Depth_95_n, 0, c='c', label = "95% Level = " + str(Depth_95_n) + 'm')
                ax[0].axvline(TPU_v[0], 0, c='k', label = "CHS/IHO = " + str(TPU_v[0]) + 'm')
                ax[0].legend(loc='upper right')
                ax[0].set_title('Vertical Accuracy (Ave Depth ' + str(Depth_mean) + ' m)' + '(Order = ' + str(orders) + ')\n')
                ax[0].set_xlabel('Depth Accuracy (m)')
                ax[0].set_ylabel('Percentage (%)')

                maxd = round(ASCII_Out['Depth'].max(),0)+5
                Depths = np.arange(0,maxd,1)
                Contours = [0,2,5,10,15,20,30,50,100]
                ATPU=[]
                for D in Depths:
                    ATPU.append(TPU(orders,D)[0])


                colors = {'yes': 'green', 'no':'red'}
                ax[1].scatter(ASCII_Out['Depth'],ASCII_Out['Depth TPU'],c=ASCII_Out['Within Allowable TVU'].map(colors))
                ax[1].plot(Depths, ATPU)
                #ax[1].legend(loc='upper right')
                ax[1].set_title('Vertical Accuracy Scatterplot) ' + '(Order = ' + str(orders) + ')\n Red - Outside Allowable Green - Within Allowable(' + str(P_W_A_TVU) + '%)')
                ax[1].set_xlabel('Depth (m)')
                ax[1].set_ylabel('Depth Accuracy (m)')
                plt.tight_layout()
                plt.savefig(str(Out) + '/' + str(JD) + '/' + TPUQCFolder  + '/Coverage/' + orders + '_Accuracy.png', dpi=200)


        elif self.TPUQC.get()==2:
            ASCII_Files = (Out + '/' + JD + '/' + TPUQCFolder + '/HIPS')
            HIPS_Files = listdir(ASCII_Files)
            Import = [pd.read_table(ASCII_Files + '/' + file, delimiter=',') for file in HIPS_Files]
            ASCII_Out = pd.concat(Import)
            ASCII_Out.columns = ['Lat', 'Long', 'Depth', 'Depth TPU', 'Position TPU']

            THU_D = []
            TVU_D = []
            for Depth in ASCII_Out['Depth']:
                TPU2 = TPU(order, Depth)
                TVU_D.append(TPU2[0])
                THU_D.append(TPU2[1])
            ASCII_Out['Allowable TVU'] = TVU_D
            ASCII_Out['Allowable THU'] = THU_D
            ASCII_Out['Within Allowable THU'] = np.where(ASCII_Out['Position TPU'] <= ASCII_Out['Allowable THU'],
                                               'yes', 'no')
            P_W_A_THU = round(((len(ASCII_Out[ASCII_Out['Within Allowable THU'] == 'yes'])/len(ASCII_Out)))*100,2)

            ASCII_Out['Within Allowable TVU'] = np.where(ASCII_Out['Depth TPU'] <= ASCII_Out['Allowable TVU'],
                                               'yes', 'no')

            P_W_A_TVU = round(((len(ASCII_Out[ASCII_Out['Within Allowable TVU'] == 'yes'])/len(ASCII_Out)))*100,2)

            Depth_mean = round(ASCII_Out['Depth'].mean(),3)

            POS_TPU_max = round(ASCII_Out['Position TPU'].max(),3)
            POS_TPU_min = round(ASCII_Out['Position TPU'].min(),3)
            POS_TPU_mean = round(ASCII_Out['Position TPU'].mean(),3)
            POS_TPU_std = round(ASCII_Out['Position TPU'].std(),3)
            POS_95_p = round((POS_TPU_mean + t_95_2d*POS_TPU_std),3)
            POS_95_n = round((POS_TPU_mean - t_95_2d*POS_TPU_std),3)

            Depth_TPU_max = round(ASCII_Out['Depth TPU'].max(),3)
            Depth_TPU_min = round(ASCII_Out['Depth TPU'].min(),3)
            Depth_TPU_mean = round(ASCII_Out['Depth TPU'].mean(),3)
            Depth_TPU_std = round(ASCII_Out['Depth TPU'].std(),3)
            Depth_95_p = round((Depth_TPU_mean + t_95_1d*Depth_TPU_std),3)
            Depth_95_n = round((Depth_TPU_mean - t_95_1d*Depth_TPU_std),3)

            TPU_v = TPU(order, Depth_mean)

            fig, ax = plt.subplots(nrows=2)
            D_TPU = list(ASCII_Out.loc[:,'Depth TPU'])
            ax[0].hist(D_TPU, weights=np.ones(len(D_TPU)) / len(D_TPU), alpha=0.5)
            ax[0].axvline(Depth_TPU_max, 0, c='r', label = "MAX = " + str(Depth_TPU_max) + 'm')
            ax[0].axvline(Depth_TPU_mean, 0, c='g', label = "MEAN = " + str(Depth_TPU_mean) + 'm')
            ax[0].axvline(Depth_TPU_min, 0, c='c', label = "MIN = " + str(Depth_TPU_min) + 'm')
            ax[0].axvline(Depth_95_p, 0, c='m', label = "95% Level = " + str(Depth_95_p) + 'm')
            #ax[0].axvline(Depth_95_n, 0, c='c', label = "95% Level = " + str(Depth_95_n) + 'm')
            ax[0].axvline(TPU_v[0], 0, c='k', label = "CHS/IHO = " + str(TPU_v[0]) + 'm')
            ax[0].legend(loc='upper right')
            ax[0].set_title('Vertical Accuracy (Ave Depth ' + str(Depth_mean) + ' m)' + '(Order = ' + str(order) + ')\n')
            ax[0].set_xlabel('Depth Accuracy (m)')
            ax[0].set_ylabel('Percentage (%)')

            POS_TPU = list(ASCII_Out.loc[:,'Position TPU'])
            ax[1].hist(POS_TPU, weights=np.ones(len(POS_TPU)) / len(POS_TPU), alpha=0.5)
            ax[1].axvline(POS_TPU_max, 0, c='r', label = "MAX = " + str(POS_TPU_max) + 'm')
            ax[1].axvline(POS_TPU_mean, 0, c='g', label = "MEAN = " + str(POS_TPU_mean) + 'm')
            ax[1].axvline(POS_TPU_min, 0, c='c', label = "MIN = " + str(POS_TPU_min) + 'm')
            ax[1].axvline(POS_95_p, 0, c='m', label = "95% Level = " + str(POS_95_p) + 'm')
            #ax[1].axvline(POS_95_n, 0, c='c', label = "95% Level = " + str(POS_95_n) + 'm')
            ax[1].axvline(TPU_v[1], 0, c='k', label = "CHS/IHO = " + str(TPU_v[1]) + 'm')
            ax[1].legend(loc='upper right')
            ax[1].set_title('Horizontal Accuracy' + '(Order = ' + str(order) + ')')
            ax[1].set_xlabel('Positional Accuracy (m)')
            ax[1].set_ylabel('Percentage (%)')
            plt.tight_layout()
            plt.savefig(str(Out) + '/' + JD + '/' + TPUQCFolder + '/HIPS/' + order + '_Accuracy.png', dpi=200)


##        total_t = Line_Report.loc[:,'Total Time']
##        length = Line_Report.loc[:,'Length']
##
##        F_total_t = []
##        for t in total_t:
##            s = str(t)
##            check = re.findall(r'\d\d\D\d\d\D\d\d\d',s)
##            if check==[]:
##                t = ('00:00:' + t)
##                F_total_t.append(t)
##            else:
##                t = ('00:' + t)
##                F_total_t.append(t)
##        line_count = len(F_total_t)
##        Line_Report.loc[:,'Total Time'] = F_total_t
##        total_sec = []
##        for t in F_total_t:
##            timeparts = [(s) for s in t.split(':')]
##            hour_s = int(timeparts[0])*3600
##            min_s = int(timeparts[1])*60
##            sec_s = float(timeparts[2])
##            total_sec.append((min_s + sec_s))
##        tts = sum(total_sec)
##        h = int(tts/3600)
##        m = int((tts - h*3600)/60)
##        s = round((tts - h*3600 - m*60),3)
##        Total_Survey_Time = (str(h) + ':' + str(m) + ':' + str(s))
##        Total_Survey_Length= sum(length)/1000

        SVPCount = self.SVP_Count()
        self.Line_Report()
        Line_Report = pd.read_csv(Out + '/' + JD + '/LineReport_' + JD + '.csv', delimiter=',')
        line_count = len(Line_Report)
        Total_Survey_Time = self.Total_Survey_Time
        Total_Survey_Length = self.Total_Survey_Length

        shp_file = (Out + '/' + JD +'/TPUQC/' + 'Polygon_' + str(JD) + '/' + JD + '_' + Year + 'cvrage(A).shp')
        Area_shp = shapefile.Reader(shp_file)
        shp_records = Area_shp.records()
        l_Areas = len(shp_records)
        Areas = []
        c = 0
        while c < l_Areas:
            rec = Area_shp.record(c)
            Areas.append(rec['AREA'])
            c = c + 1
        Total_Survey_Area = round(sum(Areas)/(1000*1000),3)


        myworkbook2 = openpyxl.load_workbook(Report_F)
        cws = myworkbook2.create_sheet(JD)
        worksheet = myworkbook2.get_sheet_by_name(JD)

        from openpyxl.utils.dataframe import dataframe_to_rows
        rows = dataframe_to_rows(Line_Report)
        for r_idx, row in enumerate(rows, 1):
            for c_idx, value in enumerate(row, 1):
                 worksheet.cell(row=r_idx, column=c_idx, value=value)

        ## Determine User Order
        if order == 'EXCLUSIVE':
            PWATVU = List_TVU[0]
        elif order == 'SPECIAL':
            PWATVU = List_TVU[1]
        elif order == '1A':
            PWATVU = List_TVU[2]
        elif order == '1B':
            PWATVU = List_TVU[3]    
        elif order == '2':
            PWATVU = List_TVU[4]
        elif order == '3':
            PWATVU = List_TVU[5]

        IHO_orders = ['EXCLUSIVE', 'SPECIAL', '1A', '1B', '2', '3']
        worksheet['A'+ str(line_count + 5)] = ('Summary')
        worksheet.merge_cells('A'+ str(line_count + 5) + ':B' + str(line_count + 5))
        worksheet['A'+ str(line_count + 6)] = ('Total Survey Time (hh:mm:ss.sss)')
        worksheet['A'+ str(line_count + 7)] = ('Total Length (km)')
        worksheet['A'+ str(line_count + 8)] = ('Total Area sqkm')
        worksheet['A'+ str(line_count + 9)] = ('Total Sound Velocity Casts')
        worksheet['A'+ str(line_count + 10)] = ('% TVU Values within ' + str(order))
        worksheet['A'+ str(line_count + 11)] = ('% THU Values within ' + str(order))
        worksheet['B'+ str(line_count + 6)] = (Total_Survey_Time)
        worksheet['B'+ str(line_count + 7)] = (Total_Survey_Length)
        worksheet['B'+ str(line_count + 8)] = (Total_Survey_Area)
        worksheet['B'+ str(line_count + 9)] = (SVPCount)
        worksheet['B'+ str(line_count + 10)] = str(PWATVU)

        
        if self.TPUQC.get()==2:
            worksheet['B'+ str(line_count + 11)] = str(P_W_A_THU)
        myworkbook2.save(Report_F)

        wb = openpyxl.load_workbook(Report_F)
        ws = wb.get_sheet_by_name(JD)
        if self.TPUQC.get()==1:
            img = openpyxl.drawing.image.Image(str(Out) + '/' + str(JD) + '/' + 'TPUQC' +  '/Coverage/' + order + '_Accuracy.png')
        elif self.TPUQC.get()==2:
            img = openpyxl.drawing.image.Image(str(Out) + '/' + str(JD) + '/' + 'TPUQC' + '/HIPS/' + order + '_Accuracy.png')
        ws.add_image(img, 'A' + str(line_count + 14))
        wb.save(Report_F)
        startfile(Report_F)

##        DayofWeek = date.today().strftime("%A")
##        T_Date = (str(date.today())).split('-')
##        if DayofWeek == 'Monday' and New_Sheet== 0:
##            Space = 8
##
##        elif DayofWeek == 'Tuesday' and New_Sheet== 0:
##            Space = 8*2
##
##        elif DayofWeek == 'Wednesday' and New_Sheet== 0:
##            Space = 8*3
##
##        elif DayofWeek == 'Thursday ' and New_Sheet== 0:
##            Space = 8*4
##
##        elif DayofWeek == 'Friday' and New_Sheet== 0:
##            Space = 8*5
##
##        elif DayofWeek == 'Saturday' and New_Sheet== 0:
##            Space = 8*6
##
##        elif DayofWeek == 'Sunday' or New_Sheet== 1:
##            Space = 0
##            cws = Weekly.create_sheet(Name)
        Weekly = openpyxl.load_workbook(Weekly_Report)
        worksheet = Weekly.get_sheet_by_name(Name)
        Space = len(worksheet['A']) + 1
        
        worksheet['A'+ str(1 + Space)] = ('Summary JD' + str(JD))
        worksheet.merge_cells('A'+ str(1) + ':B' + str(1))
        worksheet['A'+ str(2 + Space)] = ('Total Survey Time (hh:mm:ss.sss)')
        worksheet['A'+ str(3 + Space)] = ('Total Length (km)')
        worksheet['A'+ str(4 + Space)] = ('Total Area sqkm')
        worksheet['A'+ str(5 + Space)] = ('Total Sound Velocity Casts')
        worksheet['A'+ str(6 + Space)] = ('% TVU Values within ' + str(order))
        worksheet['A'+ str(7 + Space)] = ('% THU Values within ' + str(order))
        worksheet['B'+ str(2 + Space)] = (Total_Survey_Time)
        worksheet['B'+ str(3 + Space)] = (Total_Survey_Length)
        worksheet['B'+ str(4 + Space)] = (Total_Survey_Area)
        worksheet['B'+ str(5 + Space)] = (SVPCount)
        worksheet['B'+ str(6 + Space)] = str(PWATVU)
        if self.TPUQC.get()==2:
            worksheet['B'+ str(7 + Space)] = str(P_W_A_THU)
        Weekly.save(Weekly_Report)
        startfile(Weekly_Report)


    def SVP_Count(self):
        """"Count the total number of SVP cast coducted each Julian Day"""

        SVP_dir = self.SVPDir.get()
        SVPfiles = listdir(SVP_dir)

        svpcounter = 0
        for file in SVPfiles:
            if file.endswith((".asvp",".svp")):
                svpcounter = svpcounter + 1
        print(svpcounter)
        return(svpcounter)
                                  
        
    def Line_Report(self):

        HDCS_Folder = self.HDCS_D.get()
        Vessel_F = self.VESSEL_N.get()
        Vessel = path.basename(Vessel_F)
        Vessel = re.sub(".hvf","", Vessel)
        Year = (self.Year.get())
        JD = self.JULIAN_D.get()
        Out = self.OUT_F.get()
        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        hips = HIPSProject(HDCS_Folder + '/' + HIPSFILE + '/' + HIPSFILE + '.hips')
        lines = hips.get_lines()
        tot = 0

        Vessels = list(hips.get_vessels())
        vessel_n = []

        i = 1
        while i <= len(Vessels):
            v = list(Vessels[i-1])
            vessel_n.append(v[1])
            i = i + 1

        LR = pd.DataFrame()
        for line in lines:
            dict_new = line.attributes
            LR = LR.append(dict_new, ignore_index=True)
        print (LR)
        i2 = 1
        while i2 <= len(vessel_n):
            LR['Vessel Id'].mask(LR['Vessel Id'] == float(i2), vessel_n[i2-1], inplace=True)
            i2 = i2 + 1
        LR['Vessel'] = Vessel
        LR['Day'] = (Year + '-' + JD)


##        convert_dict = {'Georeferenced': int,
##                        'Tpu Computed' : int,
##                        'Gps Vertical Reference Available' : int,
##                        'Vertical Reference' : int,
##                        'Tide Available' : str,
##                        'Svp Corrected' : str,
##                        'Outdated' : str,
##                        'Raw Range' : str,
##                        'Data Confidence Computed' : str,
##                        'Del Dft Loaded' : str,
##                        'Raw Data Path': str
##                        }
##
##        LR = LR.astype(convert_dict)
        LR = LR[LR['Raw Data Path'].str.contains(('.+' +'JD' +  JD + '.+'), regex=True)]

        LR['Georeferenced'].mask(LR['Georeferenced'] == 1, 'Yes', inplace=True)
        LR['Georeferenced'].mask(LR['Raw Range'] == 'None', 'No', inplace=True)
        LR['Georeferenced'].mask(LR['Georeferenced'] == 0, 'No', inplace=True)
        
        LR['Tpu Computed'].mask(LR['Tpu Computed'] == 1, 'Yes', inplace=True)
        LR['Tpu Computed'].mask(LR['Tpu Computed'] == 'None', 'No', inplace=True)
        LR['Tpu Computed'].mask(LR['Tpu Computed'] == 0, 'No', inplace=True)

        LR['Gps Vertical Reference Available'].mask(LR['Gps Vertical Reference Available'] == 1, 'Yes', inplace=True)
        LR['Gps Vertical Reference Available'].mask(LR['Gps Vertical Reference Available'] == 'None', 'No', inplace=True)
        LR['Gps Vertical Reference Available'].mask(LR['Gps Vertical Reference Available'] == 0, 'No', inplace=True)

        LR['Tide Available'].mask((LR['Tide Available']) == 0, 'No', inplace=True)
        LR['Tide Available'].mask(LR['Tide Available'] == 'None', 'No', inplace=True)
        LR['Tide Available'].mask(LR['Tide Available'] == 1, 'Yes', inplace=True)

        LR['Svp Corrected'].mask((LR['Svp Corrected']) == 0, 'No', inplace=True)
        LR['Svp Corrected'].mask(LR['Svp Corrected'] == 'None', 'No', inplace=True)
        LR['Svp Corrected'].mask(LR['Svp Corrected'] == 1, 'Yes', inplace=True)

        LR['Outdated'].mask((LR['Outdated']) == 0, 'No', inplace=True)
        LR['Outdated'].mask(LR['Outdated'] == 'None', 'No', inplace=True)
        LR['Outdated'].mask(LR['Outdated'] == 1, 'Yes', inplace=True)

        LR['Raw Range'].mask((LR['Raw Range']) == 0, 'No', inplace=True)
        LR['Raw Range'].mask(LR['Raw Range'] == 'None', 'No', inplace=True)
        LR['Raw Range'].mask(LR['Raw Range'] == 1, 'Yes', inplace=True)

        LR['Data Confidence Computed'].mask((LR['Data Confidence Computed']) == 0, 'No', inplace=True)
        LR['Data Confidence Computed'].mask(LR['Data Confidence Computed'] == 'None', 'No', inplace=True)
        LR['Data Confidence Computed'].mask(LR['Data Confidence Computed'] == 1, 'Yes', inplace=True)

        LR['Del Dft Loaded'].mask((LR['Del Dft Loaded']) == 0, 'No', inplace=True)
        LR['Del Dft Loaded'].mask(LR['Del Dft Loaded'] == 'None', 'No', inplace=True)
        LR['Del Dft Loaded'].mask(LR['Del Dft Loaded'] == 1, 'Yes', inplace=True)
       
        LR['Vertical Reference'].mask(LR['Vertical Reference'] == 0, 'NONE', inplace=True)
        LR['Vertical Reference'].mask(LR['Vertical Reference'] == 1, 'TIDE', inplace=True)
        LR['Vertical Reference'].mask(LR['Vertical Reference'] == 2, 'GPS', inplace=True)

        
        LR['Total Time'] = (LR['Max Time'] - LR['Min Time'])/1000
        self.Total_Survey_Time = time.strftime("%H:%M:%S", time.gmtime(LR['Total Time'].sum()))


        LR['Max Time'] = LR['Max Time'].apply(lambda x: hips.convert_utc_time(x))
        LR['Min Time'] = LR['Min Time'].apply(lambda x: hips.convert_utc_time(x))
       
        LR['Total Time'] = LR['Total Time'].apply(lambda x: time.strftime("%H:%M:%S", time.gmtime(x)))


        LR = LR.drop( columns=['Concrete Object Id', 'Max Nav Time', 'Min Nav Time',
                               'Mod Time', 'Nav Examined', 'Nav Status', 'Nav Timestamps',
                               'Observed Depths Status', 'Procssed Depths Status',
                               'Resolution', 'Sources', 'Line Path', 'Towfish Nav Status',
                               'Tool Type'])

        LR = LR.sort_values(['Line Name'])
        print(LR)
        LR.to_csv(Out + '/' + JD + '/LineReport_' + JD + '.csv')

        self.Total_Survey_Length = (LR['Length'].sum())/1000


    def Create_Project_Dir(self):

        self.Search_dir()
        PH = self.split_Project_Name()
        CP = PH[0]
        HIPSFILE = PH[1]
        PL = self.PF
        chdir(PL)
        Projectname = CP
        mkdir(Projectname)

        P_F = path.join(PL,CP)
        chdir(path.join(PL,CP))
        PF_2021 = ['Processed_Data', 'BaseStation', 'MetaData', 'Inertial', 'Raw', 'SVP', 'Tide', 'VALSRC']
        for folder in PF_2021:
            full_path = P_F
            if path.exists(folder):
                return
            else:
                mkdir(folder)

        Pro_F = path.join(path.join(PL,CP), 'Processed_Data')
        PRF_2021 = [HIPSFILE, 'Surfaces', 'Vessel_Config']
        chdir(Pro_F)
        for folder in PRF_2021:
            full_path = Pro_F
            if path.exists(folder):
                return
            else:
                mkdir(folder)

        Meta_F = path.join(path.join(PL,CP), 'MetaData')
        MF_2021 = ['LogForms', 'Output', 'ISO_Documentation', 'Field_Notes']
        chdir(Meta_F)
        for folder in MF_2021:
            full_path = Meta_F
            if path.exists(folder):
                return
            else:
                mkdir(folder)

        startfile(P_F)


    def ISO_1001_07_A_F01(self):

        PH = self.split_Project_Name()
        Project_N = PH[0]
        HIPSFILE = PH[1]

        Sub_Filedir = self.SUB_D.get()
        D = '1001_07_A_F01_Template.docx'
        document = Document(D)
        Sig = self.SIG.get()
        Out = self.OUT_F.get()

        f_lst = listdir(Sub_Filedir)

        dir_lst = []
        for filename in listdir(Sub_Filedir):
            if path.isdir(path.join(Sub_Filedir, filename)):
                dir_lst.append(filename)

        Folders = []
        Files = []
        folder_size = []
        for f in dir_lst:
            fso = com.Dispatch("Scripting.FileSystemObject")
            folder = fso.GetFolder((Sub_Filedir + '/' + f) )
            F_size = folder.Size
            folder_size.append(F_size)
            F = ((Sub_Filedir + '/' + f))
            files = folders = 0
            for _, dirnames, filenames in walk(F):
                files += len(filenames)
                folders += len(dirnames)
            Folders.append(folders)
            Files.append(files)
            ##print ("{:,} files, {:,} folders".format(files, folders))

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        document.tables #a list of all tables in document
        document.tables[0].cell(1,1).text = (Project_N)
        c = 2
        i = 0
        t_d = DATES.datetime.today()
        for f in dir_lst:
            document.tables[0].cell(c,6).text = (str(dir_lst[i]))
            document.tables[0].cell(c,7).text = (str(format(folder_size[i], ',')) + ' bytes ' +
                                                 str(Files[i]) + ' Files ' +
                                                 str(Folders[i]) + ' Folders ')
            document.tables[0].cell(c,9).text = (str(Sig))
            document.tables[0].cell(c,8).text = (str(t_d))
            i = i + 1
            c = c + 1

        document.tables[1].cell(0,0).text = ('Submission verified by Hydrographer in Charge : ' + str(Sig))
        document.tables[1].cell(0,1).text = (str(t_d))

        doc = (Sub_Filedir + '/1001_07_A_F01_' + str(Project_N) + '.docx')
        document.save(doc)

        startfile(doc)

    def Find_Fliers(self):

        height = self.FHEIGHT.get()
        DTM_dir = self.DTM_DIR.get()
        DTMs = listdir(DTM_dir)

        if path.exists(DTM_dir + '/' + 'FlierFinder'):
            return
        else:
            mkdir(DTM_dir + '/' + 'FlierFinder')

        chdir(owd)

        with open("FindFliers.bat", "w") as Import:
            Import.write('@ECHO OFF' + '\n')
            Import.write('@ECHO Checking for Fliers' + '\n')
            Import.write('cd '+ QCTools + '\n')

            for dtm in DTMs:
                if dtm.ensdwith('*.csar'):
                    if height != 'AUTO':
                        Import.write('QCTools FindFliers -enforce_height' + height + '\n')
                    else:
                        Import.write('QCTools FindFliers ' + '\n')

                    Import.write('-check_laplacian ' + '-check_curv ' +  '-check_isolated ' + '-check_slivers ' +
                                 DTM_dir + '/' + dtm + ' ' + DTM_dir + '/' + 'FlierFinder')


        p = S.check_call("FindFliers.bat", stdin=None, stdout=None, stderr=None, shell=False)

        

    def Caris_RAW_Tooltips(self):

        RAW_Tooltip = pd.read_csv('Tool_Tips_RAW.txt', delimiter=';', header=None)

        R_T = []
        i = 0
        while i !=  len(RAW_Tooltip):
            rt = RAW_Tooltip.iloc[i,1]
            R_T.append(rt)
            i = i + 1

        if self.S_T.get()==1:
            RAW_tip1 = ToolTip(self.Nav_d_text, R_T[1])
            RAW_tip2 = ToolTip(self.GPS_h_text, R_T[2])
            RAW_tip3 = ToolTip(self.GPS_t_text, R_T[3])
            RAW_tip4 = ToolTip(self.Heading_d_text, R_T[4])
            RAW_tip5 = ToolTip(self.Heave_d_text, R_T[5])
            RAW_tip6 = ToolTip(self.Pitch_d_text, R_T[6])
            RAW_tip7 = ToolTip(self.Roll_d_text, R_T[7])
            RAW_tip8 = ToolTip(self.SSP_d_text, R_T[8])

        elif self.S_T.get()==2:
            RAW_tip1 = ToolTip(self.D_s_text, R_T[12])
            RAW_tip2 = ToolTip(self.IN_off, R_T[13])
            RAW_tip3 = ToolTip(self.REJ_off, R_T[14])

        elif self.S_T.get()==3:
            RAW_tip1 = ToolTip(self.D_s_text, R_T[12])
            RAW_tip2 = ToolTip(self.IN_off, R_T[13])
            RAW_tip3 = ToolTip(self.REJ_off, R_T[14])


    def Help(self):
        return


    def close(self):
        self.Exit = 'True'
        self.popup_SavePar()
        root.destroy()


def on_closing():
    if messagebox.askokcancel("Quit", "Did you save Parameters?"):
        root.destroy()

root = Tk()
root.title("CHS Pycessing Tool")
root.geometry("700x650")
menu = Menu(root)
root.config(menu=menu)
submenu = Menu(menu)
submenu2 = Menu(menu)
notebook = ttk.Notebook(root)
frame1 = ttk.Frame(notebook)
notebook.add(frame1, text="Caris Hips\nProcessing")
frame2 = ttk.Frame(notebook)
notebook.add(frame2, text="Import Sensor\nData")
frame3 = ttk.Frame(notebook)
notebook.add(frame3, text="Import Auxiliary\nData")
frame4 = ttk.Frame(notebook)
notebook.add(frame4, text="Apply\nTides")
frame5 = ttk.Frame(notebook)
notebook.add(frame5, text="Compute\nTPU")
frame6 = ttk.Frame(notebook)
notebook.add(frame6, text="Apply\nSVP")
notebook.grid(row=0, column=0)
frame7 = ttk.Frame(notebook)
notebook.add(frame7, text="HIPS\nGRID")
notebook.grid(row=0, column=0)
frame8 = ttk.Frame(notebook)
notebook.add(frame8, text="Merge\nTracklines")
notebook.grid(row=0, column=0)
##frame8 = ttk.Frame(notebook)
##notebook.add(frame8, text="CARIS\nMIRA")
##notebook.grid(row=0, column=0)
frame9 = ttk.Frame(notebook)
notebook.add(frame9, text="Reporting and \nData Submission")
notebook.grid(row=0, column=0)
frame10 = ttk.Frame(notebook)
notebook.add(frame10, text="POSPAC\nProcessing")
notebook.grid(row=0, column=0)
##frame11 = ttk.Frame(notebook)
##notebook.add(frame11, text="Copy\nHips")
##notebook.grid(row=0, column=0)
app = Application(root)
root.protocol("WM_DELETE_WINDOW", on_closing)
root.mainloop()
